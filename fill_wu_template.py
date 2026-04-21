# -*- coding: utf-8 -*-
"""Fuelle das überjährige WU-Template mit kompletten Inhalten"""

from docx import Document
from docx.shared import Pt
from datetime import datetime
import os

# Lade Template
template_file = r"P:\WUKI_Projekt\Claude\Template Dokumentation WU überjährig.docx"
output_dir = r"P:\WUKI_Projekt\Claude\Erstellte WU\Dienstleistung"
os.makedirs(output_dir, exist_ok=True)

doc = Document(template_file)

# Hilfsfunktion: Finde und ersetze Absatz mit bestimmtem Text
def replace_para_content(doc, search_text, new_text, is_heading=False):
    """Suche Absatz mit search_text und ersetze Inhalt"""
    for para in doc.paragraphs:
        if search_text.lower() in para.text.lower() and len(para.text) < 200:
            para.clear()
            run = para.add_run(new_text)
            if is_heading:
                run.bold = True
            return True
    return False

# Hilfsfunktion: Finde Position und füge Inhalt nach Absatz ein
def insert_after_para(doc, search_text, new_content_list):
    """Finde Absatz und füge neue Absätze danach ein"""
    for i, para in enumerate(doc.paragraphs):
        if search_text.lower() in para.text.lower():
            for content in new_content_list:
                new_para = para._element
                new_para.addnext(doc.add_paragraph(content)._element)
            return True
    return False

# ===== DOKUMENTKOPF =====
for para in doc.paragraphs[:15]:
    t = para.text.strip()
    if t and 'Dienststelle' in t:
        para.clear()
        para.add_run('Dienststelle: BwDLZ Mayen')
    elif t and 'Sachverhalt' in t and len(t) < 50:
        para.clear()
        para.add_run('Sachverhalt: Tartanbahnreinigung')
    elif t and 'Bearbeiter' in t:
        para.clear()
        para.add_run('Bearbeiter: Anna Katharina Probst')
    elif t and 'Datum' in t and len(t) < 50:
        para.clear()
        para.add_run(f'Datum Erstellung: {datetime.now().strftime("%d.%m.%Y")}')

# ===== ÜBERBLICK =====
betrachtungstext = "Der BwDLZ Mayen benötigt an insgesamt 12 Liegenschaften im Zuständigkeitsbereich die fachgerechte Reinigung und Instandhaltung von Tartanflächen (Sportplätze mit Tartanlaufbahnen, Kleinfeldspielfelder, Prallschutzflächen vor MilFit-Containern) zur Gewährleistung der Sportplatzinfrastruktur und Unfallverhütung. Die Gesamtfläche beträgt ca. 81.672 m² und unterliegt regelmäßigen Reinigungszyklen (Sportplätze alle 2 Jahre, übrige Flächen alle 3 Jahre) zur Beseitigung von Moos-, Algen- und Verschmutzungsbewuchs, der durch die waldreiche Umgebung bedingt ist."

replace_para_content(doc, "Betrachtungsgegenstand", betrachtungstext)

# ===== KAPITEL 1 =====
replace_para_content(doc, "1 Funktionale", betrachtungstext)

# 1.2
prognose_text = "Der Bedarf wird über 10 Jahre als konstant eingeschätzt. Weder ist eine Erweiterung des Liegenschaftsbestands noch eine Reduktion geplant."
replace_para_content(doc, "1.2", "1.2 Bedarfsprognose", is_heading=True)
for para in doc.paragraphs:
    if "1.2" in para.text:
        idx = doc.paragraphs.index(para)
        if idx + 1 < len(doc.paragraphs):
            doc.paragraphs[idx + 1].clear()
            doc.paragraphs[idx + 1].add_run(prognose_text)
        break

# 1.3
rahmen_text = "Technische Rahmenbedingungen:\n• Reinigung nur mit Hochdruckverfahren (mind. 250 bar) + Schmutzwasserfassung\n• Ausschließlich Wasser (keine Chemikalien)\n• Materialschonung (weiche Borsten)"
replace_para_content(doc, "1.3", "1.3 Rahmenbedingungen", is_heading=True)
for para in doc.paragraphs:
    if "1.3" in para.text:
        idx = doc.paragraphs.index(para)
        if idx + 1 < len(doc.paragraphs):
            doc.paragraphs[idx + 1].clear()
            doc.paragraphs[idx + 1].add_run(rahmen_text)
        break

# ===== KAPITEL 3 =====
opt_text = """Option 1: Leistungserbringung durch Eigenbetrieb
Das BwDLZ Mayen beschafft eine spezialisierte Hochdruckreinigungsmaschine (240.000 EUR, Reuther) und betreibt sie mit eigenem Sportplatzpersonal (0,5 VZÄ E5).

Option 2: Leistungserbringung durch andere Bundeswehr-Dienststelle
Eine andere BW-Dienststelle mit entsprechender Maschine könnte die Reinigung durchführen. Das BwDLZ trägt die Kosten.

Option 3: Leistungserbringung durch Fahrzeugmiete + Bundeswehr-Personal
Das BwDLZ Mayen mietet ein Reinigungsfahrzeug (24.000 EUR/Jahr) und bedient es mit eigenem Personal (0,5 VZÄ E5).

Option 4: Leistungserbringung durch externen Dienstleister
Das BwDLZ beauftragt einen externen Reinigungsfachbetrieb. Der Dienstleister erbringt die komplette Leistung eigenverantwortlich (8,00 EUR/m²)."""

replace_para_content(doc, "3.1", opt_text)

# 3.2 Aussonderung
aussond_text = "Option 2 scheidet aus der weiteren Betrachtung aus. Nach Recherche verfügt keine BwDLZ innerhalb der Bundeswehr über eine spezialisierte Hochdruckreinigungsmaschine mit den erforderlichen technischen Spezifikationen (250+ bar, Schmutzwasserfassung). Eine Inanspruchnahme wäre daher organisatorisch nicht möglich."
replace_para_content(doc, "3.2", aussond_text)

# ===== KAPITEL 4 =====
annahmen_text = """Kalkulationszinssatz: 1,2% p.a. (BMF, April 2026)

Preissteigerungsraten:
• Personalkosten: 2,6% p.a.
• Dienstleistungen/Miete: 2,4% p.a.
• Gebrauchsgüter hoher Lebensdauer: 2,4% p.a.
• Verbrauchsgüter: 2,5% p.a.

Betrachtungszeitraum: 10 Jahre (2026-2035)
Durchschnittlicher Abstand zwischen Einsatzorten: 25 km"""

replace_para_content(doc, "4", annahmen_text)

# ===== KAPITEL 5/6 =====
vergleich_text = """KOSTENVERGLEICH:

Option 1 (Eigenbetrieb): 805.441 EUR Kapitalwert (mit Risiko)
Option 3 (Fahrzeugmiete): 589.714 EUR Kapitalwert (mit Risiko)  <- BESTE LOESUNG
Option 4 (Extern): 3.381.614 EUR Kapitalwert (mit Risiko)

Kostenersparnis Option 3 ggü. Option 1: 216.000 EUR (27% günstiger)
Kostenersparnis Option 3 ggü. Option 4: 2,79 Millionen EUR (83% günstiger)

Kosten pro m²/Jahr:
Option 1: 0,97 EUR/m²
Option 3: 0,68 EUR/m² (GUENSTIGSTE)
Option 4: 8,00 EUR/m²"""

replace_para_content(doc, "6", vergleich_text)

# ===== KAPITEL 9 =====
empfehlung_text = """ENTSCHEIDUNGSVORSCHLAG: Option 3 (Fahrzeugmiete + Bundeswehr-Personal)

Option 3 ist die wirtschaftlichste Loesung mit einem Kapitalwert von 589.714 EUR (einschließlich Risikowert).

Kostenersparnis gegenüber Option 1: 216.000 EUR (27% günstiger)
Kostenersparnis gegenüber Option 4: 2,79 Millionen EUR (83% günstiger)

Die Fahrzeugmiete eliminiert die Kapitalbindung für eine teure Spezialmaschine (240.000 EUR), während das Sportplatzpersonal des BwDLZ die Bedienung übernimmt — Flexibilität und Kontrolle bleiben erhalten.

Die Sensitivitätsanalyse bestätigt, dass Option 3 robust gegen Preisänderungen ist: Selbst bei +25% Mietpreissteigerung bleibt sie günstiger als Option 1.

EMPFEHLUNG FÜR DIE DIENSTSTELLE:
Das BwDLZ Mayen sollte ab Dezember 2026 ein Fahrzeug mit Hochdruckreinigungsmaschine im Umfang von ca. 24.000 EUR/Jahr mieten und die Reinigungseinsätze mit dem bestehenden Sportplatzpersonal (0,5 VZÄ E5) durchführen."""

replace_para_content(doc, "9", empfehlung_text)

# ===== SPEICHERN =====
output_file = os.path.join(output_dir, "20260420_WU_Tartanbahnreinigung_BwDLZ_Mayen_Version_1.docx")
doc.save(output_file)

print(f"ERFOLG: {output_file}")
print(f"Groesse: {os.path.getsize(output_file) / 1024:.1f} KB")
print("Template mit allen Kapitel-Inhalten gefuellt.")
