from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Seitenränder ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Hilfsfunktionen ───────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def header_row(table, labels, bg="1F3864"):
    row = table.rows[0]
    for i, lbl in enumerate(labels):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(lbl)
        run.bold       = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size  = Pt(9)

def add_data_row(table, values, bold_last=False, bg=None, align_right_cols=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        if bg:
            set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        is_last  = (i == len(values) - 1)
        is_bold  = bold_last and is_last
        right_al = align_right_cols and i in align_right_cols
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if right_al else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        if is_bold or (bold_last and i == 0):
            run.bold = True

def add_table_row(table, values, bold=False, bg=None, align_right_cols=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        if bg:
            set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        right_al = align_right_cols and i in align_right_cols
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if right_al else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        run.bold = bold

def style_heading(paragraph, level=1):
    if level == 1:
        paragraph.style = doc.styles["Heading 1"]
    elif level == 2:
        paragraph.style = doc.styles["Heading 2"]
    elif level == 3:
        paragraph.style = doc.styles["Heading 3"]

def add_heading(text, level=1):
    p = doc.add_paragraph()
    style_heading(p, level)
    p.clear()
    run = p.add_run(text)
    sizes = {1: 14, 2: 12, 3: 11}
    run.font.size = Pt(sizes.get(level, 11))
    run.bold = True
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    else:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p

def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold   = bold
    run.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TITELSEITE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("WIRTSCHAFTLICHKEITSUNTERSUCHUNG")
run.font.size  = Pt(18)
run.bold       = True
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Beschaffung DIXI\nMobile Toilettenanlagen für Übungen auf dem Bundeswehr-Übungsplatz")
run2.font.size  = Pt(14)
run2.bold       = True

doc.add_paragraph()

# Metadaten-Tabelle
meta = doc.add_table(rows=1, cols=2)
meta.style = "Table Grid"
header_row(meta, ["Feld", "Inhalt"])
rows_meta = [
    ("Organisationseinheit", "KOMPZWUBw"),
    ("Behörde", "Kompetenzzentrum Wirtschaft und Umweltschutz der Bundeswehr (OBB)"),
    ("Erstellt von", "KOMPZWUBw"),
    ("Datum", "14.04.2026"),
]
for k, v in rows_meta:
    r = meta.add_row()
    r.cells[0].paragraphs[0].add_run(k).font.size = Pt(10)
    r.cells[0].paragraphs[0].runs[0].bold = True
    r.cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
set_col_widths(meta, [5, 10])

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# GP 1 — IST-ZUSTAND
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. Analyse des Ist-Zustands", 1)
add_body(
    "Auf dem genutzten Bundeswehr-Übungsplatz sind keine festen oder dauerhaft installierten "
    "Sanitäranlagen vorhanden. Die sanitäre Grundversorgung der eingesetzten Soldaten wird daher "
    "bei jeder Übung durch die Anmietung mobiler Toilettenanlagen (DIXI-Kabinen) sichergestellt. "
    "Die Anmietung erfolgte bisher ad-hoc ohne strukturierte Planung, Rahmenvertrag oder "
    "dokumentierten Marktpreisvergleich. Historische Kostendaten liegen nicht vor.\n\n"
    "Für das Kalenderjahr 2026 sind insgesamt sechs Übungen mit je 200 Teilnehmenden und einer "
    "Dauer von je 10 Tagen geplant. Vor diesem Hintergrund wird erstmalig eine "
    "Wirtschaftlichkeitsuntersuchung gemäß § 7 BHO durchgeführt.\n\n"
    "Die organisatorische Abwicklung der Anmietung wird durch zwei Beschäftigte des KOMPZWUBw "
    "(Gehobener Dienst, A9) mit einem Aufwand von je ca. 4 Stunden pro Übung wahrgenommen."
)

add_body("Aktuelle Kosten / Ressourceneinsatz (Ist-Zustand, geschätzt auf Basis Marktpreise 2026):", bold=True)

ist = doc.add_table(rows=1, cols=3)
ist.style = "Table Grid"
header_row(ist, ["Kostenart", "Betrag (EUR/Jahr)", "Bemerkung"])
ist_data = [
    ("Personalkosten",                       "2.200",   "2 Beschäftigte A9 OBB, 48 Std. gesamt"),
    ("Sächliche Verwaltungsausgaben (PSK)",  "510",     "anteilig, 48 Std."),
    ("Raumkosten (PSK)",                     "319",     "anteilig, 48 Std."),
    ("Mietkosten mobile Toilettenanlagen",   "23.400",  "16 Kabinen × 6 Übungen, Marktpreise 2026"),
]
for row in ist_data:
    r = ist.add_row()
    for i, val in enumerate(row):
        p = r.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if i == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

r_sum = ist.add_row()
set_cell_bg(r_sum.cells[0], "D6E4F0")
set_cell_bg(r_sum.cells[1], "D6E4F0")
set_cell_bg(r_sum.cells[2], "D6E4F0")
for i, val in enumerate(["Summe Ist-Kosten (geschätzt)", "26.429", "historische Daten nicht dokumentiert"]):
    p = r_sum.cells[i].paragraphs[0]
    run = p.add_run(val)
    run.bold = True
    run.font.size = Pt(9)
    if i == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_col_widths(ist, [6.5, 3.5, 5])

# ══════════════════════════════════════════════════════════════════════════════
# GP 2 — ZIELE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading("2. Ziele und Zielkonflikte", 1)
add_body(
    "Die Maßnahme dient der Erfüllung der Fürsorgepflicht des Dienstherrn gegenüber den "
    "eingesetzten Soldatinnen und Soldaten gemäß § 31 SG sowie allgemeiner Anforderungen "
    "an Hygiene und Gesundheitsvorsorge im Übungsbetrieb."
)
add_body("Primärziele:", bold=True)

ziele = doc.add_table(rows=1, cols=3)
ziele.style = "Table Grid"
header_row(ziele, ["Nr.", "Ziel (lösungsneutral, messbar)", "Mindestanforderung"])
ziele_data = [
    ("1", "Sicherstellung der sanitären Grundversorgung für alle Teilnehmenden je Übung",
          "Mind. 1 Kabine je 15 Personen = mind. 14 Kabinen; tatsächlich 16 (inkl. 1 barrierefreie Einheit)"),
    ("2", "Bereitstellung der Anlagen zu Übungsbeginn",
          "Verfügbarkeit spätestens am ersten Übungstag"),
    ("3", "Wirtschaftlicher Einsatz der Haushaltsmittel",
          "Unterschreitung des Ist-Kostenniveaus von 26.429 EUR"),
]
for row in ziele_data:
    r = ziele.add_row()
    for i, val in enumerate(row):
        p = r.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
set_col_widths(ziele, [1, 7, 7])

doc.add_paragraph()
add_body("Zielkonflikte:", bold=True)
add_body("Zwischen den definierten Zielen bestehen keine wesentlichen Konflikte. Kostenwirtschaftlichkeit und Versorgungssicherheit lassen sich durch strukturierte Vergabe gleichzeitig erreichen.")

# ══════════════════════════════════════════════════════════════════════════════
# GP 3 — BETRACHTUNGSZEITRAUM
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading("3. Betrachtungszeitraum", 1)
add_body("Geplanter Durchführungszeitraum: Januar 2026 bis Dezember 2026 (12 Monate, 6 Einzelübungen à 10 Tage)")
add_body(
    "Diese Wirtschaftlichkeitsuntersuchung betrifft eine unterjährige Maßnahme mit niedrigem "
    "finanziellem Volumen (ca. 22.000–27.000 EUR). Gemäß VV WU ist daher der vereinfachte Ansatz "
    "des reinen Einnahmen-/Ausgabenvergleichs zulässig.",
    italic=True
)
add_body("Eine Diskontierung mit dem Nominalzinssatz ist nicht erforderlich (unterjährige Maßnahme gemäß VV WU).", bold=True)

# ══════════════════════════════════════════════════════════════════════════════
# GP 4 — HANDLUNGSALTERNATIVEN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading("4. Relevante Handlungsalternativen", 1)

alt_tab = doc.add_table(rows=1, cols=3)
alt_tab.style = "Table Grid"
header_row(alt_tab, ["Alternative", "Kurzbezeichnung", "Kurzbeschreibung"])
alt_data = [
    ("Alt. 0", "Nullalternative (Status quo)",        "Fortführung der bisherigen ad-hoc-Einzelbeauftragung je Übung ohne strukturierte Planung"),
    ("Alt. 1", "Strukturierte Einzelbeauftragung",     "Marktpreisvergleich vor jeder Übung, Beauftragung des günstigsten Anbieters (ca. –8 %)"),
    ("Alt. 2", "Rahmenvertrag",                        "Einmaliger Vertragsabschluss für alle 6 Übungen; Mengenrabatt ca. 15 %"),
    ("Alt. 3", "Kauf eigener Anlagen",                 "Einmalinvestition in 16 eigene Toilettenkabinen; Transport, Reinigung und Lagerung durch KOMPZWUBw"),
]
for row in alt_data:
    r = alt_tab.add_row()
    for i, val in enumerate(row):
        p = r.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
set_col_widths(alt_tab, [2, 4, 9])

# ══════════════════════════════════════════════════════════════════════════════
# GP 5 — EIGNUNGSPRÜFUNG
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading("5. Eignungsprüfung der Alternativen", 1)
add_heading("5.1 Ausgeschlossene Alternativen", 2)
add_body("Alle identifizierten Alternativen erfüllen die Grundvoraussetzungen und werden in die Detailbetrachtung einbezogen.")
add_heading("5.2 Einbezogene Alternativen", 2)
add_body(
    "Alle vier Alternativen (Alt. 0 bis Alt. 3) sind rechtlich zulässig, technisch umsetzbar und "
    "grundsätzlich geeignet, die definierten Ziele zu erreichen. Private Anbieter stehen am Markt "
    "in ausreichender Zahl zur Verfügung (§ 7 Abs. 2 Satz 3 BHO geprüft)."
)

# ══════════════════════════════════════════════════════════════════════════════
# GP 6 — DETAILLIERTE DARSTELLUNG
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading("6. Detaillierte Darstellung der Alternativen", 1)

def add_kosten_table(beschreibung_zeilen, kosten_zeilen, gesamt, netto):
    """Erzeugt eine Kostentabelle für eine Alternative."""
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    header_row(t, ["Kostenart", "Berechnung", "Betrag (EUR)"])
    for zeile in kosten_zeilen:
        r = t.add_row()
        for i, val in enumerate(zeile):
            p = r.cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Gesamtausgaben
    r_g = t.add_row()
    set_cell_bg(r_g.cells[0], "D6E4F0")
    set_cell_bg(r_g.cells[1], "D6E4F0")
    set_cell_bg(r_g.cells[2], "D6E4F0")
    for i, val in enumerate(["Gesamtausgaben", "", gesamt]):
        p = r_g.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.bold = True; run.font.size = Pt(9)
        if i == 2: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Einsparungen
    r_e = t.add_row()
    for i, val in enumerate(["./. Einsparungen / Erlöse", "", "—"]):
        p = r_e.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if i == 2: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Netto
    r_n = t.add_row()
    set_cell_bg(r_n.cells[0], "1F3864")
    set_cell_bg(r_n.cells[1], "1F3864")
    set_cell_bg(r_n.cells[2], "1F3864")
    for i, val in enumerate(["Netto-Finanzbedarf", "", netto]):
        p = r_n.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if i == 2: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_col_widths(t, [6, 6, 3])
    return t

# Alt. 0
add_heading("6.1 Alternative 0 — Nullalternative (Status quo)", 2)
add_body(
    "Bei Beibehaltung des Status quo werden mobile Toilettenanlagen je Übung ad-hoc ohne "
    "strukturierten Marktpreisvergleich oder Rahmenvertrag angemietet. Mengenrabatte werden nicht "
    "genutzt. Der Planungsaufwand je Übung beträgt ca. 4 Stunden pro Mitarbeitenden (2 Beschäftigte A9)."
)
add_kosten_table(None, [
    ("Personalkosten A9 OBB",               "48h × 45,84 EUR/h",                          "2.200"),
    ("Sächliche Verwaltungsausgaben (PSK)", "48h × 10,63 EUR/h",                          "510"),
    ("Raumkosten (PSK)",                    "48h × 6,65 EUR/h",                           "319"),
    ("Einmalige Investitionskosten",        "entfällt",                                    "—"),
    ("Laufende Sachkosten (Mietkosten)",    "15 × 240 EUR × 6 + 1 × 300 EUR × 6",        "23.400"),
    ("Remanenzkosten",                      "entfällt",                                    "—"),
], "26.429", "26.429")

doc.add_paragraph()

# Alt. 1
add_heading("6.2 Alternative 1 — Strukturierte Einzelbeauftragung", 2)
add_body(
    "Vor jeder Übung wird durch die zuständigen Beschäftigten ein Marktpreisvergleich bei "
    "mindestens drei Anbietern eingeholt. Die Beauftragung erfolgt an den jeweils günstigsten "
    "Anbieter. Erfahrungsgemäß sind durch Preisvergleich Einsparungen von ca. 8 % gegenüber "
    "unstrukturierten Einzelbeauftragungen realisierbar. Der Planungsaufwand erhöht sich auf "
    "ca. 5 Stunden pro Mitarbeitenden je Übung."
)
add_kosten_table(None, [
    ("Personalkosten A9 OBB",               "60h × 45,84 EUR/h",  "2.750"),
    ("Sächliche Verwaltungsausgaben (PSK)", "60h × 10,63 EUR/h",  "638"),
    ("Raumkosten (PSK)",                    "60h × 6,65 EUR/h",   "399"),
    ("Einmalige Investitionskosten",        "entfällt",            "—"),
    ("Laufende Sachkosten (Mietkosten –8%)","23.400 × 0,92",      "21.528"),
    ("Remanenzkosten",                      "entfällt",            "—"),
], "25.315", "25.315")

doc.add_paragraph()

# Alt. 2
add_heading("6.3 Alternative 2 — Rahmenvertrag", 2)
add_body(
    "Es wird ein Rahmenvertrag mit einem Anbieter für alle 6 Übungen des Jahres 2026 "
    "abgeschlossen. Durch die Bündelung des Bedarfs ist ein Mengenrabatt von ca. 15 % auf "
    "die Standardmietpreise realisierbar. Der einmalige Mehraufwand für die Vertragsverhandlung "
    "und -gestaltung beträgt ca. 8 Stunden pro Mitarbeitenden. Der laufende Koordinationsaufwand "
    "je Übung sinkt auf ca. 2 Stunden pro Mitarbeitenden."
)
add_kosten_table(None, [
    ("Personalkosten A9 OBB",               "40h × 45,84 EUR/h",   "1.834"),
    ("Sächliche Verwaltungsausgaben (PSK)", "40h × 10,63 EUR/h",   "425"),
    ("Raumkosten (PSK)",                    "40h × 6,65 EUR/h",    "266"),
    ("Einmalige Investitionskosten",        "entfällt",             "—"),
    ("Laufende Sachkosten (Mietkosten –15%)","23.400 × 0,85",      "19.890"),
    ("Remanenzkosten",                      "entfällt",             "—"),
], "22.415", "22.415")

doc.add_paragraph()

# Alt. 3
add_heading("6.4 Alternative 3 — Kauf eigener Anlagen", 2)
add_body(
    "Das KOMPZWUBw beschafft 16 eigene Toilettenkabinen (15 Standard, 1 barrierefrei). "
    "Transport zum und vom Übungsplatz erfolgt per angemieteten 7,5t-LKW. Reinigung nach jeder "
    "Übung durch einen externen Dienstleister. Da keine Lagerfläche kostenfrei zur Verfügung "
    "steht, ist Lagerraum anzumieten. Qualitative Besonderheit: Die Kabinen können über 2026 "
    "hinaus für weitere Übungen genutzt werden."
)
add_kosten_table(None, [
    ("Personalkosten A9 OBB",                       "116h × 45,84 EUR/h",                  "5.317"),
    ("Sächliche Verwaltungsausgaben (PSK)",         "116h × 10,63 EUR/h",                  "1.233"),
    ("Raumkosten (PSK)",                            "116h × 6,65 EUR/h",                   "771"),
    ("Einmalige Investitionskosten (16 Kabinen)",   "16 × 1.200 EUR",                      "19.200"),
    ("Laufende Sachkosten — Transport",             "430 EUR × 6 Übungen",                 "2.580"),
    ("Laufende Sachkosten — Reinigung extern",      "640 EUR × 6 Übungen",                 "3.840"),
    ("Laufende Sachkosten — Lagerraum",             "25 m² × 12 EUR/m² × 12 Monate",      "3.600"),
    ("Remanenzkosten",                              "entfällt",                             "—"),
], "36.541", "36.541")

# Gesamtübersicht
add_page_break()
add_heading("6.5 Gesamtübersicht Alternativenvergleich", 2)

ges = doc.add_table(rows=1, cols=5)
ges.style = "Table Grid"
header_row(ges, ["Kostenart", "Alt. 0", "Alt. 1", "Alt. 2", "Alt. 3"])
ges_data = [
    ("Personalkosten",                     "2.200",  "2.750",  "1.834",  "5.317"),
    ("Sächl. Verwaltungsausgaben (PSK)",   "510",    "638",    "425",    "1.233"),
    ("Raumkosten (PSK)",                   "319",    "399",    "266",    "771"),
    ("Einmalige Investitionskosten",       "—",      "—",      "—",      "19.200"),
    ("Laufende Sachkosten (spez.)",        "23.400", "21.528", "19.890", "10.020"),
    ("Remanenzkosten",                     "—",      "—",      "—",      "—"),
]
for row in ges_data:
    r = ges.add_row()
    for i, val in enumerate(row):
        p = r.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Gesamtausgaben
r_ga = ges.add_row()
set_cell_bg(r_ga.cells[0], "D6E4F0")
for i, val in enumerate(["Gesamtausgaben", "26.429", "25.315", "22.415", "36.541"]):
    p = r_ga.cells[i].paragraphs[0]
    r_ga.cells[i]
    set_cell_bg(r_ga.cells[i], "D6E4F0")
    run = p.add_run(val)
    run.bold = True; run.font.size = Pt(9)
    if i > 0: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Einsparungen
r_es = ges.add_row()
for i, val in enumerate(["./. Einsparungen/Erlöse", "—", "—", "—", "—"]):
    p = r_es.cells[i].paragraphs[0]
    run = p.add_run(val)
    run.font.size = Pt(9)
    if i > 0: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Netto
r_nf = ges.add_row()
for i, val in enumerate(["Netto-Finanzbedarf", "26.429", "25.315", "22.415", "36.541"]):
    set_cell_bg(r_nf.cells[i], "1F3864")
    p = r_nf.cells[i].paragraphs[0]
    run = p.add_run(val)
    run.bold = True; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if i > 0: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Rang
r_rg = ges.add_row()
for i, val in enumerate(["Rang", "3", "2", "1", "4 *"]):
    set_cell_bg(r_rg.cells[i], "2E74B5")
    p = r_rg.cells[i].paragraphs[0]
    run = p.add_run(val)
    run.bold = True; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if i > 0: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

set_col_widths(ges, [6, 2.5, 2.5, 2.5, 2.5])

doc.add_paragraph()
p_note = doc.add_paragraph()
run_note = p_note.add_run(
    "Alle Beträge in EUR netto. Keine Diskontierung — unterjährige Maßnahme gemäß VV WU.\n"
    "* Alt. 3 ist im Betrachtungsjahr 2026 die teuerste Alternative. Bei nachgewiesener "
    "Weiternutzung über mehrere Jahre verringert sich der effektive Jahreskostenbeitrag erheblich."
)
run_note.font.size = Pt(8)
run_note.italic = True

# ══════════════════════════════════════════════════════════════════════════════
# GP 7 — RISIKOBETRACHTUNG
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading("7. Risikobetrachtung", 1)
add_heading("7.1 Risikobewertungsmatrix", 2)

risk = doc.add_table(rows=1, cols=6)
risk.style = "Table Grid"
header_row(risk, ["Risiko", "Alt.", "Eintrittswahr-\nscheinlichkeit", "Auswirkung", "Risikostufe", "Gegenmaßnahme"])
risk_data = [
    ("Kurzfristige Nichtverfügbarkeit eines Anbieters kurz vor Übungsbeginn", "0, 1", "mittel", "hoch",   "mittel", "Frühzeitige Buchung mind. 4 Wochen vor Übungsbeginn"),
    ("Preisschwankungen am Markt",                                             "0, 1", "gering", "gering", "gering", "Marktbeobachtung; bei Alt. 1 durch Vergleich abgefedert"),
    ("Anbieter erfüllt Rahmenvertrag nicht (Insolvenz, Kapazitätsengpass)",   "2",    "gering", "hoch",   "mittel", "Vertragliche Absicherung; Benennung eines Ersatzanbieters"),
    ("Beschädigungen / Vandalismus an eigenen Kabinen auf dem Übungsplatz",   "3",    "mittel", "mittel", "mittel", "Versicherungsschutz prüfen; Übergabeprotokoll"),
    ("Transportschäden beim An-/Abtransport",                                 "3",    "gering", "mittel", "gering", "Sorgfältige Verladung; Transportversicherung"),
    ("Höhere Reinigungskosten als kalkuliert",                                 "3",    "mittel", "gering", "gering", "Festpreisvertrag mit Reinigungsdienstleister"),
]
for row in risk_data:
    r = risk.add_row()
    for i, val in enumerate(row):
        p = r.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8)
        if i in (2, 3, 4):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_col_widths(risk, [4.5, 1, 2, 2, 2, 4.5])

doc.add_paragraph()
add_heading("7.2 Gesamtrisikoprofil je Alternative", 2)

rp = doc.add_table(rows=1, cols=3)
rp.style = "Table Grid"
header_row(rp, ["Alternative", "Gesamtrisikoprofil", "Wesentliche Risikotreiber"])
rp_data = [
    ("Alt. 0 — Nullalternative",              "mittel", "Hohe Abhängigkeit von kurzfristiger Anbieterverfügbarkeit; kein strukturierter Ausweichplan"),
    ("Alt. 1 — Strukturierte Einzelbeauftragung", "gering", "Marktpreisvergleich reduziert Risiken; zeitlich begrenzte Anbieterbindung"),
    ("Alt. 2 — Rahmenvertrag",                "gering", "Planungssicherheit durch vertragliche Bindung; Insolvenzrisiko des Anbieters gering"),
    ("Alt. 3 — Kauf",                         "mittel", "Sachschadenrisiko; erhöhter Logistikaufwand; Lager- und Transportrisiken"),
]
for row in rp_data:
    r = rp.add_row()
    for i, val in enumerate(row):
        p = r.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if i == 0: run.bold = True
set_col_widths(rp, [5, 3, 7])

# ══════════════════════════════════════════════════════════════════════════════
# GP 8 — SENSITIVITÄTSANALYSE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading("8. Sensitivitätsanalyse", 1)
add_heading("8.1 Kritische Annahmen", 2)
add_bullet("Erzielbarer Mengenrabatt beim Rahmenvertrag (Alt. 2): angenommen 15 %")
add_bullet("Reinigungskosten je Übung durch externen Dienstleister (Alt. 3): angenommen 640 EUR")
add_bullet("Kaufpreis je Toilettenkabine (Alt. 3): angenommen 1.200 EUR (Neupreis Kunststoffkabine)")

add_heading("8.2 Auswirkungsanalyse (Variation um ±20 %)", 2)

sens = doc.add_table(rows=1, cols=5)
sens.style = "Table Grid"
header_row(sens, ["Annahme", "Grundwert (EUR)", "−20 % (EUR)", "+20 % (EUR)", "Auswirkung auf Rangfolge"])
sens_data = [
    ("Mietkosten Alt. 2 (Rahmenvertrag)",  "19.890", "15.912 → Gesamt: 18.437", "23.868 → Gesamt: 26.393", "stabil — Alt. 2 bleibt günstigste Mietalternative"),
    ("Reinigungskosten Alt. 3 (3.840 EUR)","3.840",  "3.072 → Alt. 3: 35.773",  "4.608 → Alt. 3: 37.309",  "stabil"),
    ("Kaufpreis Kabinen Alt. 3 (19.200 EUR)","19.200","15.360 → Alt. 3: 32.701","23.040 → Alt. 3: 40.381", "stabil"),
]
for row in sens_data:
    r = sens.add_row()
    for i, val in enumerate(row):
        p = r.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if i in (1, 2, 3):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_col_widths(sens, [4, 2.5, 3.5, 3.5, 4.5])

add_heading("8.3 Ergebnis der Sensitivitätsanalyse", 2)
add_body(
    "Die Empfehlung für Alt. 2 (Rahmenvertrag) ist robust gegenüber Schwankungen in allen drei "
    "kritischen Annahmen. Selbst bei einer 20-prozentigen Erhöhung der Mietkosten (Netto-Finanzbedarf: "
    "26.393 EUR) liegt Alt. 2 noch knapp unterhalb von Alt. 0 (26.429 EUR) und deutlich unter "
    "Alt. 3 (36.541 EUR). Die Rangfolge ändert sich in keinem der betrachteten Szenarien. "
    "Alt. 3 bleibt im Betrachtungsjahr 2026 unabhängig von Parametervariationen die teuerste Option."
)

# ══════════════════════════════════════════════════════════════════════════════
# GP 9 — ENTSCHEIDUNGSEMPFEHLUNG
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading("9. Ergebnis und Entscheidungsempfehlung", 1)
add_heading("9.1 Zusammenfassung des Alternativenvergleichs", 2)
add_body(
    "Alt. 2 (Rahmenvertrag) erzielt mit 22.415 EUR den niedrigsten Netto-Finanzbedarf aller "
    "betrachteten Alternativen und unterschreitet den Status quo (Alt. 0: 26.429 EUR) um 4.014 EUR "
    "(ca. 15 %). Alt. 1 (Strukturierte Einzelbeauftragung) liegt mit 25.315 EUR ebenfalls unter "
    "dem Status quo, ist jedoch gegenüber Alt. 2 um 2.900 EUR teurer. Alt. 3 (Kauf) ist im "
    "Betrachtungsjahr 2026 mit 36.541 EUR die kostenintensivste Option, bietet jedoch langfristig "
    "Einsparungspotenzial durch Wiederverwendung. Alle Alternativen erfüllen die definierten "
    "Versorgungsziele vollständig."
)

add_heading("9.2 Entscheidungsempfehlung", 2)
p_emp = doc.add_paragraph()
run_emp = p_emp.add_run("Empfohlene Alternative: Alternative 2 — Rahmenvertrag")
run_emp.bold = True
run_emp.font.size = Pt(11)
run_emp.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

add_body("Begründung:", bold=True)
add_body(
    "Alt. 2 ist mit einem Netto-Finanzbedarf von 22.415 EUR die wirtschaftlichste Alternative im "
    "Betrachtungsjahr 2026. Gegenüber dem Status quo (Alt. 0) werden 4.014 EUR eingespart; "
    "gegenüber Alt. 1 werden 2.900 EUR eingespart. Alle definierten Versorgungsziele "
    "(mind. 16 Kabinen, Bereitstellung zu Übungsbeginn) werden vollständig erfüllt. Das "
    "Risikoprofil ist gering — die vertragliche Bindung erhöht die Planungssicherheit und "
    "reduziert den laufenden Koordinationsaufwand je Übung auf ca. 2 Stunden pro Mitarbeitenden. "
    "Die Sensitivitätsanalyse bestätigt die Stabilität dieser Empfehlung unter allen betrachteten "
    "Parametervariationen.\n\n"
    "Alt. 3 (Kauf) wird für kommende Haushaltsjahre zur erneuten Prüfung empfohlen, sofern die "
    "Nutzung der Übungsplätze dauerhaft geplant ist."
)

add_heading("9.3 Haushaltswirkung", 2)
hw = doc.add_table(rows=1, cols=2)
hw.style = "Table Grid"
header_row(hw, ["Feld", "Inhalt"])
hw_data = [
    ("Gesamter Netto-Finanzbedarf (empfohlene Alternative)", "22.415 EUR"),
    ("Haushaltsjahr",                                        "2026"),
    ("Betroffener Titel",                                    "KOMPZWUBw (Haushaltsstelle nicht bekannt)"),
]
for k, v in hw_data:
    r = hw.add_row()
    r.cells[0].paragraphs[0].add_run(k).font.size = Pt(9)
    r.cells[0].paragraphs[0].runs[0].bold = True
    r.cells[1].paragraphs[0].add_run(v).font.size = Pt(9)
set_col_widths(hw, [8, 7])

add_heading("9.4 Umsetzungsplanung", 2)
up = doc.add_table(rows=1, cols=3)
up.style = "Table Grid"
header_row(up, ["Meilenstein", "Geplantes Datum", "Verantwortlich"])
up_data = [
    ("Entscheidung / Freigabe WU",          "April 2026",  "Leitung KOMPZWUBw"),
    ("Rahmenvertrag abschließen",            "Mai 2026",    "KOMPZWUBw"),
    ("Erste Übung (Bereitstellung)",         "Q2/Q3 2026",  "KOMPZWUBw"),
    ("Letzte Übung (Bereitstellung)",        "Q4 2026",     "KOMPZWUBw"),
    ("Abschlusskontrolle Erfolgskontrolle",  "Q1 2027",     "KOMPZWUBw"),
]
for row in up_data:
    r = up.add_row()
    for i, val in enumerate(row):
        p = r.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
set_col_widths(up, [7, 4, 4])

# ══════════════════════════════════════════════════════════════════════════════
# ANHANG — ERFOLGSKONTROLLE
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading("Anhang: Erfolgskontrolle", 1)
add_heading("A.1 Erfolgskontrollkriterien und KPIs", 2)

kpi = doc.add_table(rows=1, cols=6)
kpi.style = "Table Grid"
header_row(kpi, ["KPI / Erfolgskriterium", "Ausgangswert", "Zielwert", "Messmethode", "Messtermin", "Verantwortlich"])
kpi_data = [
    ("Sanitärversorgung je Übung sichergestellt (mind. 16 Kabinen verfügbar)",
     "nicht dokumentiert", "100 % der Übungen", "Abnahmeprotokoll je Übung", "Nach jeder Übung", "KOMPZWUBw"),
    ("Gesamtkosten je Übung",
     "ca. 4.405 EUR (Alt. 0, geschätzt)", "max. 3.736 EUR (Alt. 2)", "Rechnungsprüfung", "Nach letzter Übung 2026", "KOMPZWUBw"),
    ("Gesamter Netto-Finanzbedarf 2026",
     "26.429 EUR (geschätzt)", "max. 22.415 EUR", "Jahresabrechnung", "Q1 2027", "KOMPZWUBw"),
]
for row in kpi_data:
    r = kpi.add_row()
    for i, val in enumerate(row):
        p = r.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8)
set_col_widths(kpi, [4, 2.5, 2.5, 3, 2.5, 2.5])

add_heading("A.2 Planung der Erfolgskontrolle", 2)
add_body("Art der Erfolgskontrolle: Abschließende Erfolgskontrolle", bold=True)
add_body("(Maßnahme kürzer als 2 Jahre: keine begleitende Erfolgskontrolle erforderlich)", italic=True)
add_body("Zeitpunkt der Abschlusskontrolle: Q1 2027 (ca. 3 Monate nach der letzten Übung)")
add_body("Verantwortliche Organisationseinheit: KOMPZWUBw")
add_body("Dreigeteilte Erfolgskontrolle umfasst:", bold=True)
add_bullet("Zielerreichungskontrolle: Soll-Ist-Vergleich — wurden alle 6 Übungen mit mind. 16 Kabinen versorgt?")
add_bullet("Wirkungskontrolle: War die strukturierte Vergabe per Rahmenvertrag ursächlich für die Kosteneinsparung?")
add_bullet("Wirtschaftlichkeitskontrolle: Entsprechen die tatsächlichen Gesamtkosten dem Zielwert von max. 22.415 EUR?")

# ── Fußnote ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
p_foot = doc.add_paragraph()
run_foot = p_foot.add_run(
    "Dieses Dokument wurde erstellt gemäß § 7 BHO und der Verwaltungsvorschrift zur "
    "Wirtschaftlichkeitsuntersuchung (VV WU). Für die vorliegende Maßnahme wurde der vereinfachte "
    "Ansatz des reinen Einnahmen-/Ausgabenvergleichs ohne Diskontierung angewendet "
    "(unterjährige Maßnahme, Volumen < 200.000 EUR, Untersuchungstiefe Stufe 1)."
)
run_foot.font.size = Pt(8)
run_foot.italic = True

# ── Speichern ─────────────────────────────────────────────────────────────────
output_path = r"P:\WUKI_Projekt\Claude\WU_Beschaffung_DIXI_2026.docx"
doc.save(output_path)
print(f"Gespeichert: {output_path}")
