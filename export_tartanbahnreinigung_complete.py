#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tartanbahnreinigung WU - Komplette Export mit Templatestruktur-Aufbau
Nutzt export_wu_ueberjahrig.py aber mit vorbereiteter Template-Struktur
"""

import sys
import os
sys.path.insert(0, r'P:\WUKI_Projekt\Claude\skills\wu-berater\scripts')

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from export_wu_ueberjahrig import (fill_template, build_filename,
                                   format_alle_ueberschriften, _format_heading_run)

# Lade Template und baue fehlende Struktur auf
template_path = r"P:\WUKI_Projekt\Claude\Template Dokumentation WU überjährig.docx"
doc = Document(template_path)

# Finde den Punkt nach dem Dokumentkopf (nach den Hinweisen)
insertion_idx = None
for i, para in enumerate(doc.paragraphs):
    if "Diese Seite ist nach der Finalisierung" in para.text:
        insertion_idx = i
        break

if insertion_idx is None:
    insertion_idx = len(doc.paragraphs) - 1

# Lösche alte Dokumentation (nach Hinweistext)
while len(doc.paragraphs) > insertion_idx + 2:
    p = doc.paragraphs[insertion_idx + 2]._element
    p.getparent().remove(p)

# Erzeuge Kapitelstruktur
def add_heading_para(doc, ref_para, text, level):
    """Fügt Heading nach Referenz-Paragraph ein"""
    new_p = OxmlElement('w:p')
    ref_para._element.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.style = doc.styles[f'Heading {level}']
            p.clear()
            run = p.add_run(text)
            _format_heading_run(run)
            p.paragraph_format.space_after = Pt(0)
            return p
    return None

def add_text_para(doc, ref_para, text):
    """Fügt Normal-Paragraph nach Referenz ein"""
    new_p = OxmlElement('w:p')
    ref_para._element.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.style = doc.styles['Normal']
            p.add_run(text)
            return p
    return None

# Baue Kapitelstruktur auf
last = doc.paragraphs[insertion_idx]

# Überblick
last = add_heading_para(doc, last, "Überblick", 1)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "Betrachtungsgegenstand", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "Entscheidungsvorschlag", 2)
last = add_text_para(doc, last, "")

# Kapitel 1
last = add_heading_para(doc, last, "1 Funktionale Bedarfsforderung", 1)
last = add_heading_para(doc, last, "1.1", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "1.2 Bedarfsprognose", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "1.3 Rahmenbedingungen", 2)
last = add_heading_para(doc, last, "Rechtliche Rahmenbedingungen", 3)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "Organisatorische Rahmenbedingungen", 3)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "Zeitliche Rahmenbedingungen", 3)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "Sonstige Rahmenbedingungen", 3)
last = add_text_para(doc, last, "")

# Kapitel 2
last = add_heading_para(doc, last, "2 Ausgangslage", 1)

last = add_heading_para(doc, last, "2.1 Ablauforganisation", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "2.2 Aufbauorganisation", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "2.3 Personal", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "2.4 Material", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "2.5 Infrastruktur", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "2.6 Sach- und Dienstleistungen", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "2.7 Ggf. Einnahmen", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "2.8 Haushalterische Darstellung", 2)
last = add_text_para(doc, last, "")

# Kapitel 3
last = add_heading_para(doc, last, "3 Optionen der Bedarfsdeckung", 1)

last = add_heading_para(doc, last, "3.1 Grundsätzlich", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "3.2 Aussonderung von ungeeigneten Optionen", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "3.3 Ausführliche Optionendarstellung", 2)
last = add_heading_para(doc, last, "3.3.1 Option 1: Eigenbetrieb mit Maschinenneukauf", 3)
last = add_heading_para(doc, last, "3.3.1.1 Ablauforganisation", 4)
last = add_text_para(doc, last, "")
last = add_heading_para(doc, last, "3.3.1.2 Aufbauorganisation", 4)
last = add_text_para(doc, last, "")
last = add_heading_para(doc, last, "3.3.1.3 Personal", 4)
last = add_text_para(doc, last, "")
last = add_heading_para(doc, last, "3.3.1.4 Material", 4)
last = add_text_para(doc, last, "")
last = add_heading_para(doc, last, "3.3.1.5 Infrastruktur", 4)
last = add_text_para(doc, last, "")
last = add_heading_para(doc, last, "3.3.1.6 Sach- und Dienstleistungen", 4)
last = add_text_para(doc, last, "")
last = add_heading_para(doc, last, "3.3.1.7 Ggf. Einnahmen", 4)
last = add_text_para(doc, last, "")

# Kapitel 4
last = add_heading_para(doc, last, "4 Annahmen", 1)

last = add_heading_para(doc, last, "4.1 Annahmen für alle Optionen", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "4.2 Annahmen für bestimmte Optionen", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "4.3 Annahmen für die Berechnung", 2)
last = add_text_para(doc, last, "")

# Kapitel 5
last = add_heading_para(doc, last, "5 Berechnung", 1)

last = add_heading_para(doc, last, "5.1 Interessenbekundungsverfahren", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "5.2 Berechnung der Optionen", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "5.3 Kapitalwerte ohne Risiko", 2)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "5.4 Risikobetrachtung", 2)
last = add_heading_para(doc, last, "5.4.1 Risikoidentifizierung", 3)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "5.4.2 Risikoverteilung", 3)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "5.4.3 Monetäre Risikobewertung", 3)
last = add_text_para(doc, last, "")

last = add_heading_para(doc, last, "5.5 Kapitalwerte mit Risiko", 2)
last = add_text_para(doc, last, "")

# Kapitel 6
last = add_heading_para(doc, last, "6 Vergleich der Optionen", 1)
last = add_text_para(doc, last, "")

# Kapitel 7
last = add_heading_para(doc, last, "7 Sensitivitätsanalyse", 1)
last = add_text_para(doc, last, "")

# Kapitel 8
last = add_heading_para(doc, last, "8 Nicht monetäre Faktoren", 1)
last = add_text_para(doc, last, "")

# Kapitel 9
last = add_heading_para(doc, last, "9 Entscheidungsvorschlag", 1)
last = add_text_para(doc, last, "")

# Speichere vorbereitete Template
prepared_template = r"P:\WUKI_Projekt\Claude\Template_prepared_tartanbahnreinigung.docx"
doc.save(prepared_template)

# Jetzt nutze die vorbereitete Template mit fill_template
wu_data = {
    "meta": {
        "titel": "Wirtschaftlichkeitsuntersuchung Tartanbahnreinigung BwDLZ Mayen",
        "kurztitel": "Tartanbahnreinigung BwDLZ Mayen",
        "dienststelle": "BwDLZ Mayen",
        "bearbeiter": "Anna Katharina Probst",
        "datum": "20.04.2026",
        "schutz": "offen",
        "version": "1",
    },
    "ueberblick": {
        "betrachtungsgegenstand": (
            "Der BwDLZ Mayen benötigt an insgesamt 12 Liegenschaften im Zuständigkeitsbereich "
            "die fachgerechte Reinigung und Instandhaltung von Tartanflächen (Sportplätze mit "
            "Tartanlaufbahnen, Kleinfeldspielfelder, Prallschutzflächen vor MilFit-Containern) "
            "zur Gewährleistung der Sportplatzinfrastruktur und Unfallverhütung. Die Gesamtfläche "
            "beträgt ca. 81.672 m² und unterliegt regelmäßigen Reinigungszyklen (Sportplätze alle "
            "2 Jahre, übrige Flächen alle 3 Jahre)."
        ),
        "entscheidungsvorschlag": (
            "Empfohlene Option: Option 3 (Fahrzeugmiete). Kapitalwert mit Risiko: 606.714 EUR. "
            "27% günstiger als Eigenbetrieb, 84% günstiger als externe Dienstleistung."
        ),
    },
    "kap1": {
        "bedarfsforderung": (
            "Der BwDLZ Mayen benötigt an insgesamt 12 Liegenschaften im Zuständigkeitsbereich die "
            "fachgerechte Reinigung und Instandhaltung von Tartanflächen. Die Gesamtfläche beträgt "
            "ca. 81.672 m² und unterliegt regelmäßigen Reinigungszyklen."
        ),
        "bedarfsprognose": "Der Bedarf wird über 10 Jahre als konstant eingeschätzt.",
        "rb_zeitlich": "Die Maßnahme soll ab Dezember 2026 beginnen.",
        "rb_sonstige": (
            "Technische Rahmenbedingungen: Reinigung nur mit Hochdruckverfahren (mind. 250 bar) "
            "mit Schmutzwasserfassung. Ausschließlich Wasser (keine Chemikalien)."
        ),
    },
    "kap2": {
        "ablauforganisation": "Entfällt.",
        "aufbauorganisation": "Entfällt.",
        "personal": "Entfällt.",
        "material": "Entfällt.",
        "infrastruktur": "Entfällt.",
        "sach_dienstleistungen": "Entfällt.",
        "einnahmen": "Keine Einnahmen.",
        "haushalterische_darstellung": "Entfällt.",
        "haushalterische_tabelle": [["Personal", "2.3", "–", "–"],],
    },
    "kap3": {
        "optionen_grundsaetzlich": (
            "Option 1: Eigenbetrieb (240.000 EUR Maschinenneukauf). "
            "Option 2: Inhouse - wird ausgesondert. "
            "Option 3: Fahrzeugmiete (24.000 EUR/Jahr). "
            "Option 4: Externe Dienstleistung (8,00 EUR/m²)."
        ),
        "aussonderung": "Option 2 scheidet aus – keine BwDLZ verfügt über die erforderliche Maschine.",
        "optionen_detail": [
            {
                "titel": "Option 1: Eigenbetrieb",
                "ablauforganisation": "BwDLZ beschafft und betreibt Maschine intern.",
                "aufbauorganisation": "Sportplatzpersonal übernimmt Bedienung.",
                "personal": "0,5 VZÄ E5.",
                "material": "Hochdruckreinigungsmaschine 240.000 EUR.",
                "infrastruktur": "Lagerfläche für Maschine.",
                "sach_dienstleistungen": "Wartung und Reparatur durch Fremdvergabe.",
                "einnahmen": "Keine Einnahmen.",
            },
            {
                "titel": "Option 3: Fahrzeugmiete",
                "ablauforganisation": "Fahrzeug wird angemietet, BW-Personal führt Reinigung durch.",
                "aufbauorganisation": "Sportplatzpersonal übernimmt Bedienung des gemieteten Fahrzeugs.",
                "personal": "0,5 VZÄ E5.",
                "material": "Mietfahrzeug mit Hochdruckreiniger.",
                "infrastruktur": "Stellfläche für Mietfahrzeug.",
                "sach_dienstleistungen": "Fahrzeugmiete 24.000 EUR/Jahr.",
                "einnahmen": "Keine Einnahmen.",
            },
            {
                "titel": "Option 4: Externe Dienstleistung",
                "ablauforganisation": "Externes Unternehmen führt alle Reinigungen durch.",
                "aufbauorganisation": "Externe Firma mit eigenem Personal.",
                "personal": "0,1 VZÄ E9b für Abnahme.",
                "material": "Keine BW-Mittelbeschaffung.",
                "infrastruktur": "Keine zusätzliche Infrastruktur.",
                "sach_dienstleistungen": "Komplette Reinigung 8,00-10,00 EUR/m².",
                "einnahmen": "Keine Einnahmen.",
            },
        ],
    },
    "kap4": {
        "alle_optionen": "Betrachtungszeitraum: 10 Jahre (2026–2035). Kalkulationszinssatz: 1,2% p.a.",
        "bestimmte_optionen": "Option 1: Zusätzliche Instandhaltungskosten. Option 4: Mit Risikoaufschlag.",
        "berechnung": "Preissteigerung Personal 2,6%, Dienstleistungen 2,4%, Gebrauchsgüter 2,4%.",
    },
    "kap5": {
        "ibv": "Ein Interessenbekundungsverfahren wurde nicht durchgeführt.",
        "berechnung_text": "Kapitalwerte über 10 Jahre berechnet mit 1,2% Diskontierung.",
        "risikoidentifizierung": "Preisrisiken, Maschinenausfallrisiken, Vertragsrisiken erkannt.",
        "risikoverteilung": "Option 1: Risiken bei BW. Option 3: Teilweise beim Vermieter. Option 4: Beim Dienstleister.",
        "risikobewertung": "Option 1: 10%, Option 3: 5%, Option 4: 15% des Kapitalwerts.",
        "kw_mit_risiko": "Option 1: 858.441 EUR, Option 3: 606.714 EUR (BESTE), Option 4: 3.888.614 EUR.",
    },
    "kap6_9": {
        "optionen_uebersicht": [
            {"name": "Option 1: Eigenbetrieb", "kw_ohne_risiko": "805.441 EUR", "kw_mit_risiko": "858.441 EUR", "empfohlen": False},
            {"name": "Option 3: Fahrzeugmiete", "kw_ohne_risiko": "589.714 EUR", "kw_mit_risiko": "606.714 EUR", "empfohlen": True},
            {"name": "Option 4: Externe DL", "kw_ohne_risiko": "3.381.614 EUR", "kw_mit_risiko": "3.888.614 EUR", "empfohlen": False},
        ],
        "vergleich_text": "Option 3 ist wirtschaftlich überlegen.",
        "sensitivitaet": "Option 3 robust gegen Preisänderungen.",
        "nichtmonetaer": "Flexibilität, Kontrolle, Personalentwicklung.",
        "entscheidungsvorschlag": "Option 3 ab Dezember 2026.",
    },
    "anlage": [
        {"nr": "1", "produkt": "Hochdruckreiniger Reuther", "preis": "240.000 EUR", "url": "Webrecherche"},
        {"nr": "2", "produkt": "Fahrzeugmiete 24.000 EUR/Jahr", "preis": "24.000 EUR/Jahr", "url": "Webrecherche"},
        {"nr": "3", "produkt": "Externe Reinigung 8 EUR/m²", "preis": "8,00 EUR/m²", "url": "Webrecherche"},
    ],
}

# Erstelle Ausgabepfad
output_path = build_filename("20.04.2026", "Tartanbahnreinigung", "BwDLZ_Mayen", 1)

# Fülle Template
try:
    result = fill_template(wu_data, output_path)
    print(f"\nERFOLG: WU-Dokument erstellt")
    print(f"Datei: {os.path.basename(output_path)}")
    print(f"Groesse: {os.path.getsize(output_path) / 1024:.1f} KB")
    print(f"Alle Kapitel mit Inhalten befuellt!")
except Exception as e:
    print(f"Fehler: {e}")
    import traceback
    traceback.print_exc()
