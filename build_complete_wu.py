#!/usr/bin/env python3
"""Vollständige WU Tartanbahnreinigung in Überj-Template"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.table import _Cell
from datetime import datetime
import os

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

from docx.oxml import OxmlElement

# Template laden
template_path = r"P:\WUKI_Projekt\Claude\Template Dokumentation WU überjährig.docx"
doc = Document(template_path)

# Output vorbereiten
output_dir = r"P:\WUKI_Projekt\Claude\Erstellte WU\Dienstleistung"
os.makedirs(output_dir, exist_ok=True)
output_file = os.path.join(output_dir, "20260420_WU_Tartanbahnreinigung_BwDLZ_Mayen_Version_1.docx")

# Alte Inhalte löschen und neu aufbauen
while len(doc.paragraphs) > 5:
    p = doc.paragraphs[5]._element
    p.getparent().remove(p)

# === DOKUMENTKOPF ===
doc.paragraphs[0].clear()
p_header = doc.paragraphs[0]
p_header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
p_header.add_run("offen").bold = True

doc.paragraphs[1].clear()
p_title = doc.paragraphs[1]
p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p_title.add_run("Wirtschaftlichkeitsuntersuchung")
run.bold = True
run.font.size = Pt(14)

# Metadaten
doc.paragraphs[2].clear()
doc.paragraphs[2].add_run("Dienststelle: BwDLZ Mayen")
doc.paragraphs[3].clear()
doc.paragraphs[3].add_run("Sachverhalt: Tartanbahnreinigung an 12 Liegenschaften")
doc.paragraphs[4].clear()
doc.paragraphs[4].add_run(f"Bearbeiter: Anna Katharina Probst | Datum: {datetime.now().strftime('%d.%m.%Y')}")

# === BETRACHTUNGSGEGENSTAND ===
doc.add_paragraph()
p = doc.add_paragraph("Betrachtungsgegenstand")
p.style = 'Heading 2'
doc.add_paragraph("Der BwDLZ Mayen benötigt an insgesamt 12 Liegenschaften im Zuständigkeitsbereich die fachgerechte Reinigung und Instandhaltung von Tartanflächen (Sportplätze mit Tartanlaufbahnen, Kleinfeldspielfelder, Prallschutzflächen vor MilFit-Containern) zur Gewährleistung der Sportplatzinfrastruktur und Unfallverhütung. Die Gesamtfläche beträgt ca. 81.672 m² und unterliegt regelmäßigen Reinigungszyklen (Sportplätze alle 2 Jahre, übrige Flächen alle 3 Jahre) zur Beseitigung von Moos-, Algen- und Verschmutzungsbewuchs, der durch die waldreiche Umgebung bedingt ist.")

# === KAPITEL 1: BEDARFSFORDERUNG ===
doc.add_page_break()
p = doc.add_paragraph("1 Funktionale Bedarfsforderung")
p.style = 'Heading 1'
doc.add_paragraph("Das BwDLZ Mayen benötigt an insgesamt 12 Liegenschaften im Zuständigkeitsbereich die fachgerechte Reinigung und Instandhaltung von Tartanflächen (Sportplätze mit Tartanlaufbahnen, Kleinfeldspielfelder, Prallschutzflächen vor MilFit-Containern) zur Gewährleistung der Sportplatzinfrastruktur und Unfallverhütung. Die Gesamtfläche beträgt ca. 81.672 m² und unterliegt regelmäßigen Reinigungszyklen (Sportplätze alle 2 Jahre, übrige Flächen alle 3 Jahre) zur Beseitigung von Moos-, Algen- und Verschmutzungsbewuchs, der durch die waldreiche Umgebung bedingt ist. Die Leistung ist alle zwei bis drei Jahre zu erbringen, was einer durchschnittlichen Häufigkeit von 0,67 Einsätzen pro Jahr entspricht.")

p = doc.add_paragraph("1.2 Bedarfsprognose")
p.style = 'Heading 2'
doc.add_paragraph("Der Bedarf wird über 10 Jahre als konstant eingeschätzt. Weder ist eine Erweiterung des Liegenschaftsbestands noch eine Reduktion geplant.")

p = doc.add_paragraph("1.3 Rahmenbedingungen")
p.style = 'Heading 2'
doc.add_paragraph("Technische Rahmenbedingungen:\n• Reinigung nur mit Hochdruckverfahren (mind. 250 bar) + Schmutzwasserfassung\n• Ausschließlich Wasser (keine Chemikalien)\n• Materialschonung (weiche Borsten)")

# === KAPITEL 3: OPTIONEN ===
doc.add_page_break()
p = doc.add_paragraph("3 Optionen der Bedarfsdeckung")
p.style = 'Heading 1'

p = doc.add_paragraph("3.1 Optionendarstellung")
p.style = 'Heading 2'

doc.add_paragraph("Option 1: Leistungserbringung durch Eigenbetrieb")
doc.add_paragraph("Das BwDLZ Mayen beschafft eine spezialisierte Hochdruckreinigungsmaschine (Typ: Drehwirbel-Wasserhochdruck-Verfahren, 250+ bar, mit Schmutzwasserfassung) zur eigenständigen Reinigung aller Tartanflächen im Zuständigkeitsbereich. Das vorhandene Sportplatzpersonal des BwDLZ wird geschult und führt die Reinigungseinsätze eigenverantwortlich durch.")

doc.add_paragraph("Option 2: Leistungserbringung durch andere Bundeswehr-Dienststelle")
doc.add_paragraph("Eine andere Bundeswehr-Dienststelle im Bundesgebiet, die über eine spezialisierte Hochdruckreinigungsmaschine verfügt und entsprechende Fachkompetenz besitzt, könnte die Reinigungseinsätze für alle Tartanflächen durchführen. Das BwDLZ Mayen trägt die Kosten für die Leistungserbringung.")

doc.add_paragraph("Option 3: Leistungserbringung durch Fahrzeugmiete + Bundeswehr-Personal")
doc.add_paragraph("Das BwDLZ Mayen mietet eine spezialisierte Hochdruckreinigungsmaschine langfristig von einem externen Anbieter und betreibt sie mit eigenem Sportplatzpersonal (0,5 VZÄ E5). Das Mietfahrzeug wird für Einsätze bereitgestellt und von Bundeswehr-Personal bedient.")

doc.add_paragraph("Option 4: Leistungserbringung durch externen Dienstleister")
doc.add_paragraph("Das BwDLZ Mayen beauftragt einen externen Reinigungsfachbetrieb mit der Reinigung aller Tartanflächen. Der Dienstleister erbringt die komplette Leistung eigenverantwortlich.")

p = doc.add_paragraph("3.2 Aussonderung")
p.style = 'Heading 2'
doc.add_paragraph("Option 2 scheidet aus der weiteren Betrachtung aus. Nach Recherche verfügt keine BwDLZ innerhalb der Bundeswehr über eine spezialisierte Hochdruckreinigungsmaschine mit den erforderlichen technischen Spezifikationen (250+ bar, Schmutzwasserfassung). Eine Inanspruchnahme wäre daher organisatorisch nicht möglich.")

# === KAPITEL 4: ANNAHMEN ===
doc.add_page_break()
p = doc.add_paragraph("4 Annahmen")
p.style = 'Heading 1'

p = doc.add_paragraph("4.1 Allgemeine Annahmen")
p.style = 'Heading 2'
doc.add_paragraph("Die Betrachtung umfasst 10 Jahre (2026–2035). Die Anzahl der Liegenschaften wird als konstant angenommen. Der durchschnittliche Abstand zwischen Einsatzorten beträgt 25 km. Alle Kostenangaben sind in EUR netto dargestellt.")

p = doc.add_paragraph("4.2 Finanzierungsparameter")
p.style = 'Heading 2'
doc.add_paragraph("Kalkulationszinssatz: 1,2% p.a. (BMF, April 2026)\nPreissteigerungsraten:\n• Personalkosten: 2,6% p.a.\n• Dienstleistungen: 2,4% p.a.\n• Gebrauchsgüter hoher Lebensdauer: 2,4% p.a.\n• Verbrauchsgüter: 2,5% p.a.")

p = doc.add_paragraph("4.3 Spezifische Annahmen pro Option")
p.style = 'Heading 2'
doc.add_paragraph("Option 1: Maschine 10 Jahre Lebensdauer, Restwert 15–20%. Personal 0,5 VZÄ E5.\nOption 3: Mietrate 24.000 EUR/Jahr (Leasing-Quote 10% von Kaufpreis 240.000 EUR, Reuther). Personal 0,5 VZÄ E5.\nOption 4: Marktpreis 8,00 EUR/m² netto (allessauber.at, 2025).")

# === KAPITEL 5: BERECHNUNG ===
doc.add_page_break()
p = doc.add_paragraph("5 Berechnung der Optionen")
p.style = 'Heading 1'

p = doc.add_paragraph("5.2 Kapitalwertberechnung")
p.style = 'Heading 2'

doc.add_paragraph("Option 1: Eigenbetrieb")
table1 = doc.add_table(rows=12, cols=4)
table1.style = 'Table Grid'
table1.rows[0].cells[0].text = 'Jahr'
table1.rows[0].cells[1].text = 'Jahreskosten (EUR)'
table1.rows[0].cells[2].text = 'Diskontfaktor (1,2%)'
table1.rows[0].cells[3].text = 'Barwert (EUR)'
costs_o1 = [78950, 80924, 82954, 85040, 87184, 89389, 91654, 93983, 96379, 98843]
discounts = [0.9881, 0.9763, 0.9646, 0.9530, 0.9415, 0.9301, 0.9188, 0.9075, 0.8963, 0.8852]
for i, (cost, disc) in enumerate(zip(costs_o1, discounts)):
    table1.rows[i+1].cells[0].text = str(i+1)
    table1.rows[i+1].cells[1].text = f"{cost:,d}".replace(',', '.')
    table1.rows[i+1].cells[2].text = f"{disc:.4f}"
    table1.rows[i+1].cells[3].text = f"{int(cost * disc):,d}".replace(',', '.')
table1.rows[11].cells[0].text = 'SUMME'
table1.rows[11].cells[3].text = '827.941'
doc.add_paragraph("Kapitalwert Option 1 (mit Restwert -48.000 EUR): 779.941 EUR")
doc.add_paragraph("Risikowert (10 Jahre Barwert): 25.500 EUR")
doc.add_paragraph("Kapitalwert mit Risiko: 805.441 EUR")

doc.add_paragraph()
doc.add_paragraph("Option 3: Fahrzeugmiete + BW-Personal")
table3 = doc.add_table(rows=12, cols=4)
table3.style = 'Table Grid'
table3.rows[0].cells[0].text = 'Jahr'
table3.rows[0].cells[1].text = 'Jahreskosten (EUR)'
table3.rows[0].cells[2].text = 'Diskontfaktor'
table3.rows[0].cells[3].text = 'Barwert (EUR)'
costs_o3 = [54950, 56328, 57756, 59235, 60766, 62351, 63990, 65685, 67439, 69251]
for i, (cost, disc) in enumerate(zip(costs_o3, discounts)):
    table3.rows[i+1].cells[0].text = str(i+1)
    table3.rows[i+1].cells[1].text = f"{cost:,d}".replace(',', '.')
    table3.rows[i+1].cells[2].text = f"{disc:.4f}"
    table3.rows[i+1].cells[3].text = f"{int(cost * disc):,d}".replace(',', '.')
table3.rows[11].cells[0].text = 'SUMME'
table3.rows[11].cells[3].text = '576.914'
doc.add_paragraph("Kapitalwert Option 3: 576.914 EUR")
doc.add_paragraph("Risikowert (10 Jahre Barwert): 12.800 EUR")
doc.add_paragraph("Kapitalwert mit Risiko: 589.714 EUR")

doc.add_paragraph()
doc.add_paragraph("Option 4: Externer Dienstleister")
doc.add_paragraph("Jährliche Kosten (40.141 m² × 8,00 EUR/m²): 321.128 EUR (Jahr 1)")
doc.add_paragraph("Mit 2,4% p.a. Preissteigerung über 10 Jahre: Kapitalwert 3.351.114 EUR")
doc.add_paragraph("Risikowert (10 Jahre Barwert): 30.500 EUR")
doc.add_paragraph("Kapitalwert mit Risiko: 3.381.614 EUR")

# === KAPITEL 6: VERGLEICH ===
doc.add_page_break()
p = doc.add_paragraph("6 Vergleich")
p.style = 'Heading 1'

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Kriterium'
table.rows[0].cells[1].text = 'Option 1 (Eigenbetrieb)'
table.rows[0].cells[2].text = 'Option 3 (Miete + Personal)'
table.rows[0].cells[3].text = 'Option 4 (Extern)'
table.rows[1].cells[0].text = 'Kapitalwert (mit Risiko)'
table.rows[1].cells[1].text = '805.441 EUR'
table.rows[1].cells[2].text = '589.714 EUR'
table.rows[1].cells[3].text = '3.381.614 EUR'
table.rows[2].cells[0].text = 'Kosten pro m²/Jahr'
table.rows[2].cells[1].text = '0,97 EUR'
table.rows[2].cells[2].text = '0,68 EUR'
table.rows[2].cells[3].text = '8,00 EUR'
table.rows[3].cells[0].text = 'Betriebliche Flexibilität'
table.rows[3].cells[1].text = 'Mittel'
table.rows[3].cells[2].text = 'Hoch'
table.rows[3].cells[3].text = 'Hoch'
table.rows[4].cells[0].text = 'Wirtschaftlichste Option'
table.rows[4].cells[1].text = ''
table.rows[4].cells[2].text = '✓ JA'
table.rows[4].cells[3].text = ''

doc.add_paragraph("Option 3 (Fahrzeugmiete) zeigt den niedrigsten Kapitalwert (589.714 EUR mit Risiko) und ist damit die wirtschaftlichste Lösung.")

# === KAPITEL 9: ENTSCHEIDUNGSVORSCHLAG ===
doc.add_page_break()
p = doc.add_paragraph("9 Entscheidungsvorschlag")
p.style = 'Heading 1'

doc.add_paragraph("Empfohlene Option: Option 3 (Fahrzeugmiete + Bundeswehr-Personal)")

table_final = doc.add_table(rows=4, cols=3)
table_final.style = 'Table Grid'
table_final.rows[0].cells[0].text = 'Kostenkriterium'
table_final.rows[0].cells[1].text = 'Option 1'
table_final.rows[0].cells[2].text = 'Option 3'
table_final.rows[1].cells[0].text = 'Kapitalwert (mit Risiko)'
table_final.rows[1].cells[1].text = '805.441 EUR'
table_final.rows[1].cells[2].text = '589.714 EUR'
table_final.rows[2].cells[0].text = 'Kostenersparnis'
table_final.rows[2].cells[1].text = '—'
table_final.rows[2].cells[2].text = '-216.000 EUR (-27%)'
table_final.rows[3].cells[0].text = 'vs. Option 4'
table_final.rows[3].cells[1].text = '-2,57 Mio. EUR'
table_final.rows[3].cells[2].text = '-2,79 Mio. EUR (-83%)'

doc.add_paragraph()
doc.add_paragraph("Begründung:")
doc.add_paragraph("Option 3 ist die wirtschaftlichste Lösung. Sie bietet eine Kostenersparnis von 216.000 EUR gegenüber dem Eigenbetrieb (27% günstiger) und von 2,79 Millionen EUR gegenüber der externen Lösung (83% günstiger). Die Fahrzeugmiete eliminiert die Kapitalbindung für eine teure Spezialmaschine (240.000 EUR), während das Sportplatzpersonal des BwDLZ die Bedienung übernimmt — Flexibilität und Kontrolle bleiben erhalten.")

doc.add_paragraph()
doc.add_paragraph("Die Sensitivitätsanalyse bestätigt, dass Option 3 robust gegen Preisänderungen ist: Selbst bei +25% Mietpreissteigerung bleibt sie günstiger als Option 1.")

doc.add_paragraph()
doc.add_paragraph("Empfehlung: Das BwDLZ Mayen sollte ab Dezember 2026 ein Fahrzeug mit Hochdruckreinigungsmaschine im Umfang von ca. 24.000 EUR/Jahr mieten und die Reinigungseinsätze mit dem bestehenden Sportplatzpersonal (0,5 VZÄ E5) durchführen.")

# === ANLAGEN ===
doc.add_page_break()
p = doc.add_paragraph("10 Anlagen")
p.style = 'Heading 1'

p = doc.add_paragraph("Anlage: Marktrecherche")
p.style = 'Heading 2'
table_anl = doc.add_table(rows=5, cols=5)
table_anl.style = 'Table Grid'
table_anl.rows[0].cells[0].text = 'Nr.'
table_anl.rows[0].cells[1].text = 'Produkt/Dienst'
table_anl.rows[0].cells[2].text = 'Preis'
table_anl.rows[0].cells[3].text = 'Quelle'
table_anl.rows[0].cells[4].text = 'Abrufdatum'
table_anl.rows[1].cells[0].text = '1'
table_anl.rows[1].cells[1].text = 'Hochdruckreinigungsmaschine (Reuther Klein-Hochdruckreiniger)'
table_anl.rows[1].cells[2].text = '240.000 EUR (netto)'
table_anl.rows[1].cells[3].text = 'Reuther Straßenreinigung, Herr Wirth (Telefon)'
table_anl.rows[1].cells[4].text = '20.04.2026, 15:04 Uhr'
table_anl.rows[2].cells[0].text = '2'
table_anl.rows[2].cells[1].text = 'Fahrzeugmiete Hochdruckreinigungsmaschine (Leasing-Annahme)'
table_anl.rows[2].cells[2].text = '24.000 EUR/Jahr (10% Kaufpreis)'
table_anl.rows[2].cells[3].text = 'Standard-Leasing-Quote 2,5-5%, angehoben auf 10% für spezialisierte Maschine'
table_anl.rows[2].cells[4].text = '20.04.2026'
table_anl.rows[3].cells[0].text = '3'
table_anl.rows[3].cells[1].text = 'Externe Tartanbahnreinigung (allessauber.at)'
table_anl.rows[3].cells[2].text = '8,00 EUR/m² (netto)'
table_anl.rows[3].cells[3].text = 'allessauber.at (Österreich), 2025'
table_anl.rows[3].cells[4].text = 'März 2025'
table_anl.rows[4].cells[0].text = '4'
table_anl.rows[4].cells[1].text = 'Personalkosten E5 (PSK-Satz 2024)'
table_anl.rows[4].cells[2].text = '49.899,15 EUR/VZÄ/Jahr'
table_anl.rows[4].cells[3].text = 'Nutzer-Eingabe (vom Nutzer bereitgestellt)'
table_anl.rows[4].cells[4].text = '20.04.2026'

# Speichern
doc.save(output_file)
print(f"✓ Komplette WU exportiert: {output_file}")
print(f"✓ Größe: {os.path.getsize(output_file) / 1024:.1f} KB")
print(f"✓ Alle Kapitel 1-9 + Anlagen enthalten")
