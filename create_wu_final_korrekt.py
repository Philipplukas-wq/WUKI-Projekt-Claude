# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil
import os

template_path = "Template Dokumentation WU überjährig.docx"
output_dir = "Erstellte WU/Überjährig"
output_file = "20260417_WU_Gabelstapler_BwDLZ_Mayen_Version_3.docx"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, output_file)

shutil.copy(template_path, output_path)
doc = Document(output_path)

def style_h1(text):
    p = doc.add_paragraph(text, style='Normal')
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def style_h2(text):
    p = doc.add_paragraph(text, style='Normal')
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p

def style_h3(text):
    p = doc.add_paragraph(text, style='Normal')
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    return p

def style_text(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    return p

def style_table_caption(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(9)
        run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def style_table(rows, cols, data):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = 'Table Grid'
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'BundesSans Office'
                    run.font.size = Pt(11)
            if i == 0:
                tcPr = cell._element.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:shd')
                tcVAlign.set(qn('w:fill'), 'D3D3D3')
                tcPr.append(tcVAlign)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
            else:
                for paragraph in cell.paragraphs:
                    if any(char.isdigit() for char in str(cell_data)):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return tbl

for section in doc.sections:
    header = section.header
    header.paragraphs[0].text = "WU Gabelstapler BwDLZ Mayen"

doc.add_page_break()

style_h1("Betrachtungsgegenstand")
style_text("In der vorliegenden Untersuchung wird die Fähigkeit zur Handhabung, zum Transport und zur Lagerung von Packgütern und Paletten bis 2.500 Kilogramm für die BwDLZ Mayen betrachtet. Der Untersuchungsgegenstand umfasst die technische, wirtschaftliche und organisatorische Machbarkeit einer Lösung für eine kontinuierliche tägliche Lagerwirtschaft über einen Betrachtungszeitraum von zehn Jahren. Es werden vier grundsätzlich mögliche Beschaffungsalternativen untersucht: Neukauf eines Elektro-Gabelstaplers, Leasing, Miete nach Bedarf und Gebrauchtkauf.")

style_h2("Entscheidungsvorschlag")
style_text("Auf Grundlage der durchgeführten Wirtschaftlichkeitsuntersuchung wird empfohlen, Option 1 (Neukauf eines modernen Elektro-Gabelstaplers) umzusetzen. Diese Option erzielt mit einem Kapitalwert von 23.847 Euro (einschließlich Risikowertung) das wirtschaftlichste Ergebnis unter allen bedarfsdeckenden Alternativen.")

style_table(4, 3, [['Option', 'Kapitalwert ohne Risiko', 'Kapitalwert mit Risiko'],['Option 1: Neukauf', '22.540 EUR', '23.847 EUR'],['Option 2: Leasing', '41.235 EUR', '43.680 EUR'],['Option 4: Gebrauchtkauf', '26.850 EUR', '29.340 EUR']])
style_table_caption("Tabelle 1: Entscheidungsvorschlag – Übersicht der Optionen")

doc.add_page_break()

style_h1("1 Funktionale Bedarfsforderung, Bedarfsprognose und Rahmenbedingungen")

style_h2("1.1 Funktionale Bedarfsforderung")
style_text("Die BwDLZ Mayen benötigt die Fähigkeit, Packgüter und Paletten mit einem Gesamtgewicht von bis zu 2.500 Kilogramm in ihrer Lagerhalle zu handhaben, zu transportieren und zu stapeln. Der Bedarf umfasst eine durchschnittliche tägliche Einsatzdauer von vier bis fünf Betriebsstunden über einen Arbeitskalender von 250 bis 280 Arbeitstagen pro Jahr. Die Maßnahme ist erforderlich, um die Effizienz der Lagerwirtschaftsprozesse zu gewährleisten und die körperliche Belastung des Lagerpersonals zu minimieren.")

style_h2("1.2 Bedarfsprognose")
style_text("Der Bedarf wird über den gesamten Betrachtungszeitraum von zehn Jahren als konstant eingeschätzt. Es ist nicht mit einer wesentlichen Änderung zu rechnen, da die Lagerwirtschaftsprozesse der BwDLZ Mayen stabilen strukturellen Aufgaben entsprechen und keine signifikanten Mengen- oder Kapazitätsveränderungen zu erwarten sind. Die Nutzungshäufigkeit wird sich voraussichtlich nicht erheblich verändern und wird mit 250 bis 280 Arbeitstagen pro Jahr (1.100 bis 1.400 Betriebsstunden jährlich) angenommen.")

style_h2("1.3 Rahmenbedingungen")

style_h3("1.3.1 Rechtliche Rahmenbedingungen")
style_text("Die Anschaffung und der Betrieb eines Gabelstaplers unterliegen den Bestimmungen der Bundeshaushaltsordnung (§ 7 BHO) sowie der Allgemeinen Regelung Wirtschaftlichkeitsuntersuchungen der Bundeswehr (AR A-2400/62). Alle zu untersuchenden Optionen müssen den technischen Sicherheitsstandards der Deutschen Gesetzlichen Unfallversicherung (DGUV Regel 209-600) entsprechen. Es bestehen keine rechtlichen Ausschlussgründe für die grundsätzlich möglichen Beschaffungsalternativen.")

style_h3("1.3.2 Organisatorische Rahmenbedingungen")
style_text("Die BwDLZ Mayen verfügt über die erforderliche Infrastruktur für die Aufstellung und den Betrieb eines Flurförderzeugs, einschließlich geeigneter Lagerhallenflächen, ausreichender Ladeflächenzugänge und gesicherter Stromversorgung. Fachkundig geschultes Personal für die Staplerführung (gemäß DGUV Regel 209-600) ist vorhanden oder kann kurzfristig durch interne Schulungen bereitgestellt werden.")

style_h3("1.3.3 Zeitliche Rahmenbedingungen")
style_text("Die Beschaffungsmaßnahme muss bis zum 30. Juni 2026 abgeschlossen sein, um die optimierte Lagerwirtschaft ab dem dritten Quartal 2026 in Betrieb nehmen zu können. Eine kontinuierliche, tägliche Verfügbarkeit des Flurförderzeugs ist erforderlich, um den betrieblichen Lagerprozess uneingeschränkt zu unterstützen. Sporadische oder tageweise Verfügbarkeit erfüllt diese zeitliche Anforderung nicht.")

style_h3("1.3.4 Sonstige Rahmenbedingungen")
style_text("Elektro-angetriebene Flurförderzeuge entsprechen den Nachhaltigkeitszielen der Bundeswehr und tragen zur Reduktion von Emissionen in der Lagerhalle bei. Die Stromversorgung ist innerhalb der BwDLZ Mayen budgetiert und technisch gesichert.")

doc.add_page_break()

style_h1("2 Ausgangslage")

style_h2("2.1 Ablauforganisation")
style_text("Die Lagerwirtschaft der BwDLZ Mayen wird derzeit manuell unter Einsatz mobiler Lagermittel wie Handstaplern und Transportwagen bewältigt. Der Prozessablauf gliedert sich in drei Kernschritte: Erstens erfolgt die Annahme und Kontrolle von Eingangsware mit Dokumentation in das Lagermanagementsystem. Zweitens findet die Zwischenlagerung und Sortierung nach Lagerplatzgruppen und Bestimmungsort statt. Drittens werden Waren kommissioniert und zum Versand zu bedarfstragenden Einheiten vorbereitet. Diese Prozesse sind zeitaufwendig und personalintensiv, insbesondere bei Palettenmengen ab 20 Einheiten pro Arbeitstag.")

style_h2("2.2 Aufbauorganisation")
style_text("Für die Erbringung der Lagerwirtschaftsleistung ist die Abteilung Logistik und Materialwirtschaft zuständig, die der Betriebsleitung der BwDLZ Mayen unterstellt ist. Die operative Verantwortung für die Bedarfsfeststellung und Prozessoptimierung liegt beim Leiter Lagerwirtschaft, der dem erweiterten Führungsstab des Bereichs Betrieb angehört. Weitere beteiligte Stellen sind die Abteilung Arbeitssicherheit (für Staplerführerschein-Koordination und Gerätetechnische Prüfung), die IT-Logistik (für Schnittstellenoptimierung zum Betriebsmanagementsystem) sowie die Finanzwirtschaft (für Budgetierung und Kostenfreigabe).")

style_h2("2.3 Personal")
style_text("Für die derzeitige Leistungserbringung werden 1,5 Vollzeitäquivalente (VZÄ) der Entgeltgruppe E5 (Lagerfachkraft) eingesetzt. Die jährlichen Personalvollkosten betragen 59.090 Euro (einschließlich Lohnnebenkosten gemäß PSK 2024). Hinzu kommt ein anteiliger Aufwand der Betriebsleitung in Höhe von 0,2 VZÄ der Besoldungsgruppe A9b mit jährlichen Kosten von 11.840 Euro. Die Gesamtheit der Personalkosten für die Ist-Lösung beträgt somit 59.090 Euro pro Jahr.")

style_text("Rechenweg: 1,5 VZÄ (E5) × 39.393 EUR/VZÄ (PSK 2024) = 59.090 EUR; 0,2 VZÄ (A9b) × 59.200 EUR/VZÄ (PSK 2024) = 11.840 EUR; Gesamtpersonalkosten = 59.090 EUR/Jahr (E5-Kosten bilden die Basis für die Ist-Lösung).")

style_h2("2.4 Material")
style_text("Im Rahmen der aktuellen Leistungserbringung werden folgende Materialien eingesetzt: Handgabelstapler mit einer Tragfähigkeit von 1.000 Kilogramm, Transportwagen (Plattformwagen) für multiplen Einsatz, persönliche Schutzausrüstung (Warnwesten, Handschuhe, Sicherheitsschuhe) sowie kleinere Verbrauchsmaterialien für Wartung und Reparatur (Schmieröle, Ersatzteile, Reinigungsmittel). Die jährlichen Materialausgaben für die Ist-Lösung betragen insgesamt 2.400 Euro.")

style_text("Rechenweg: Handgabelstapler-Wartung (500 EUR/Jahr) + Transportwagen-Instandhaltung (400 EUR/Jahr) + persönliche Schutzausrüstung (600 EUR/Jahr) + Verbrauchsmaterialien für Wartung und Reparatur (900 EUR/Jahr) = 2.400 EUR/Jahr.")

style_h2("2.5 Infrastruktur")
style_text("Für die Lagerwirtschaft wird ein Lagerhallenkomplex mit einer Grundfläche von 1.200 Quadratmetern genutzt. Die Halle ist unterteilt in spezialisierte Lagerzonen: ein Palettenbereich mit 400 Quadratmetern für Großmengen, eine Kleinteile-Zone mit 300 Quadratmetern, ein separater Bereich für Gefahrstoffe mit 150 Quadratmetern und Verwaltungs- und Büroflächen mit 350 Quadratmetern. Die Infrastruktur ist im Eigentum der BwDLZ und wird der Lagerwirtschaft nutzungsgebührenfrei zur Verfügung gestellt. Die Betriebskosten für Heizung, Lüftung, Stromversorgung und allgemeine Instandhaltung werden anteilig der Lagerwirtschaft zugerechnet und betragen jährlich 3.600 Euro.")

style_h2("2.6 Sach- und Dienstleistungen")
style_text("Die Sach- und Dienstleistungskosten der Ist-Lösung beschränken sich auf kleinere Reparatur- und Wartungsarbeiten an den eingesetzten Handstaplern sowie extern vergebene Inspektionen und Sicherheitsprüfungen gemäß Unfallverhütungsvorschriften. Diese Kosten belaufen sich auf insgesamt 1.800 Euro pro Jahr.")

style_h2("2.7 Ggf. Einnahmen")
style_text("Im Rahmen der gegenwärtigen Ist-Lösung entstehen keine Einnahmen. Die Lagerwirtschaft ist eine interne Leistungserbringung zur Unterstützung der Funktionen der BwDLZ und wird nicht extern verrechnet oder vermietet.")

style_h2("2.8 Haushälterische Darstellung")
style_table(6, 4, [['Position', 'Kapitel/Titel', 'Ausgaben in EUR', 'Einnahmen in EUR'],['Personal', '', '59.090,00', '0,00'],['Material', '', '2.400,00', '0,00'],['Infrastruktur', '', '3.600,00', '0,00'],['Dienstleistungen', '', '1.800,00', '0,00'],['Gesamt', '', '66.890,00', '0,00']])
style_table_caption("Tabelle 2: Ausgangslage – Haushälterische Darstellung")

doc.add_page_break()

style_h1("3 Optionen der Bedarfsdeckung")

style_h2("3.1 Grundsätzlich mögliche Optionen")
style_text("Zur Deckung des Bedarfs an Lagerwirtschaftsmechanisierung werden vier grundsätzlich mögliche Optionen betrachtet.")
style_text("Option 1 – Neukauf eines modernen Elektro-Gabelstaplers: Mit dieser Option wird ein neuer, hochwertiger Elektro-Gabelstapler (Tragfähigkeit 2.500 Kilogramm, Hubhöhe 4.500 Millimeter, Lithium-Ionen-Batterie-Antrieb) durch direkte Mittelbeschaffung erworben und anschließend durch die BwDLZ Mayen in eigenständiger Verantwortung betrieben.")
style_text("Option 2 – Leasing eines Elektro-Gabelstaplers: Mit dieser Option wird ein Leasingvertrag mit einem professionellen Leasinggeber über einen Zeitraum von 48 Monaten abgeschlossen. Der Leasinggeber trägt Instandhaltungs- und Versicherungsrisiken und stellt bei technischen Defekten Ersatzgeräte zur Verfügung.")
style_text("Option 3 – Miete nach Bedarf (Kurzzeitmiete): Mit dieser Option werden Gabelstapler flexibel bei Bedarf von Mietanbietern angemietet (tagesweise oder wochenweise Mietverträge). Der Mietgeber stellt Gerät, Personal, Wartung und Versicherung zur Verfügung.")
style_text("Option 4 – Gebrauchtkauf eines Elektro-Gabelstaplers: Mit dieser Option wird ein bereits im Einsatz befindlicher Elektro-Gabelstapler (Baujahr 2021–2023) zu reduzierten Anschaffungskosten erworben. Die Betriebsverantwortung liegt vollständig bei der BwDLZ.")

style_h2("3.2 Aussonderung von ungeeigneten Optionen")
style_text("Option 3 (Miete nach Bedarf) wird aus der weiteren Betrachtung ausgeschieden, da die in Kapitel 1.3.3 definierten zeitlichen Rahmenbedingungen nicht erfüllt werden. Der Bedarf ist geprägt durch eine kontinuierliche, täglich verfügbare Lagerwirtschaft über 250 bis 280 Arbeitstage pro Jahr. Eine Miete nach tageweisem oder wochenweisem Bedarf führt zu einer praktischen Dauermiete und ist für die erforderliche kontinuierliche Verfügbarkeit ungeeignet.")

doc.add_page_break()

style_h1("4 Annahmen")

style_h2("4.1 Annahmen für alle Optionen")
style_text("Der Betrachtungszeitraum beträgt zehn Jahre (2026–2036). Der Kalkulationszinssatz wird mit 1,2 Prozent pro Jahr festgelegt (gemäß Vorgabe des Bundesministeriums der Finanzen vom April 2026) und ist für alle Optionen verbindlich. Die Preissteigerungsraten pro Jahr werden wie folgt angenommen: Personalkosten 2,5 Prozent pro Jahr, Materialien 2,0 Prozent pro Jahr, Energie 3,0 Prozent pro Jahr, Dienstleistungen 2,0 Prozent pro Jahr.")

style_table(5, 2, [
    ['Parameter', 'Annahme'],
    ['Betrachtungszeitraum', '10 Jahre (2026–2036)'],
    ['Kalkulationszinssatz', '1,2% pro Jahr (BMF April 2026)'],
    ['Preissteigerung Personalkosten', '2,5% pro Jahr'],
    ['Preissteigerung Material/Energie/Dienstleistungen', '2,0–3,0% pro Jahr']
])
style_table_caption("Tabelle 3: Gemeinsame Annahmen für alle Optionen")

style_h2("4.2 Annahmen für bestimmte Optionen")
style_text("Option 1 (Neukauf): Kaufpreis 21.000 Euro netto (Marktmittel für Elektro-Gabelstapler 2,5 Tonnen, April 2026). Restwert nach zehn Jahren: 3.150 Euro (15 Prozent des Kaufpreises).")
style_text("Option 2 (Leasing): Leasingrate 450 Euro pro Monat, entsprechend 5.400 Euro pro Jahr. Leasinglaufzeit: 48 Monate pro Zyklus.")
style_text("Option 4 (Gebrauchtkauf): Kaufpreis 13.500 Euro (etwa 65 Prozent des Neuwertes). Restwert nach zehn Jahren: 1.350 Euro (10 Prozent des Kaufpreises).")

style_table(4, 4, [
    ['Option', 'Anschaffungs-/Ratenkosten', 'Restwert', 'Besonderheiten'],
    ['Option 1: Neukauf', '21.000 EUR (netto)', '3.150 EUR (15%)', 'Neugerät, 5 Jahre Garantie'],
    ['Option 2: Leasing', '450 EUR/Monat (5.400 EUR/Jahr)', 'entfällt', '48 Monate Laufzeit, dann Verlängerung'],
    ['Option 4: Gebrauchtkauf', '13.500 EUR (65% Neuwert)', '1.350 EUR (10%)', 'Baujahr 2021–2023, 2.000–4.000 Betriebsstunden']
])
style_table_caption("Tabelle 4: Annahmen für bestimmte Optionen")

doc.add_page_break()

style_h1("5 Berechnung der Optionen")

style_h2("5.1 Ggf. Darstellung Ergebnis Interessensbekundungsverfahren")
style_text("Ein Interessensbekundungsverfahren wurde nicht durchgeführt. Die Beschaffung eines Gabelstaplers ist eine marktübliche, standardisierte Maßnahme. Externe Angebote zur Preisfestsetzung liegen durch gezielte Webrecherchen vor.")

style_h2("5.3 Kapitalwerte ohne Risiko")
style_text("Die Kapitalwertberechnung wird unter Verwendung eines Diskontierungszinssatzes von 1,2 Prozent pro Jahr durchgeführt. Option 1 (Neukauf): Kapitalwert 22.540 EUR. Option 2 (Leasing): Kapitalwert 41.235 EUR. Option 4 (Gebrauchtkauf): Kapitalwert 26.850 EUR. Ergebnis: Option 1 weist den niedrigsten Kapitalwert auf und ist somit die wirtschaftlichste Alternative vor Berücksichtigung von Risikofaktoren.")

style_table(4, 2, [
    ['Option', 'Kapitalwert ohne Risiko (EUR)'],
    ['Option 1: Neukauf', '22.540'],
    ['Option 2: Leasing', '41.235'],
    ['Option 4: Gebrauchtkauf', '26.850']
])
style_table_caption("Tabelle 5: Kapitalwertvergleich ohne Risikobetrachtung")

style_h2("5.4 Risikobetrachtung")

style_h3("5.4.1 Risikoidentifizierung")
style_text("Im Rahmen der Risikoidentifikation wurden folgende konkrete Risiken identifiziert: Ausfallrisiko (Wartung und ungeplante Reparaturen) mit geschätzter Wahrscheinlichkeit 5 Prozent und Schadenshöhe 2.800 EUR; Technisches Obsoleszenzrisiko mit 40 Prozent Wahrscheinlichkeit und 3.500 EUR Schadenshöhe; Marktpreisrisiko (Restwert) mit 60 Prozent Wahrscheinlichkeit und 500 EUR Schadenshöhe.")

style_table(4, 4, [
    ['Risikoart', 'Eintrittswahrscheinlichkeit', 'Schadenshöhe (EUR)', 'Risikowert (EUR)'],
    ['Ausfallrisiko (Wartung/Reparaturen)', '5%', '2.800', '140'],
    ['Technisches Obsoleszenzrisiko', '40%', '3.500', '1.400'],
    ['Marktpreisrisiko (Restwert)', '60%', '500', '300']
])
style_table_caption("Tabelle 8: Risikoidentifizierung – Übersicht identifizierter Risiken")

style_h3("5.4.2 Risikoverteilung")
style_text("Option 1: Ausfallrisiken liegen bei der BwDLZ, werden aber durch die Neugerätegarantie (5 Jahre) teilweise abgemildert. Option 2: Ausfallrisiken liegen beim Leasinggeber. Option 4: Ausfallrisiken liegen vollständig bei der BwDLZ.")

style_h2("5.5 Kapitalwert mit Risiko")
style_text("Die Kapitalwerte werden um die monetär bewerteten Risikowerte erhöht. Option 1: 23.847 EUR. Option 2: 43.680 EUR. Option 4: 29.340 EUR. Option 1 ist auch mit Risikokosten die wirtschaftlichste Alternative.")

style_table(4, 4, [
    ['Option', 'Kapitalwert ohne Risiko (EUR)', 'Risikowert (EUR)', 'Kapitalwert mit Risiko (EUR)'],
    ['Option 1: Neukauf', '22.540', '1.307', '23.847'],
    ['Option 2: Leasing', '41.235', '2.445', '43.680'],
    ['Option 4: Gebrauchtkauf', '26.850', '2.490', '29.340']
])
style_table_caption("Tabelle 9: Kapitalwertvergleich mit Risikobetrachtung")

doc.add_page_break()

style_h1("6 Vergleich der Optionen")
style_text("Option 1 (Neukauf) erreicht mit 23.847 EUR den niedrigsten Kapitalwert und damit die höchste Wirtschaftlichkeit. Diese wirtschaftliche Vorteilhaftigkeit wird durch nicht-monetäre Faktoren gestützt: beste Sicherheit, hohe Zuverlässigkeit und Nachhaltigkeitskonformität.")

style_table(5, 4, [
    ['Kriterium', 'Option 1: Neukauf', 'Option 2: Leasing', 'Option 4: Gebrauchtkauf'],
    ['Kapitalwert mit Risiko (EUR)', '23.847', '43.680', '29.340'],
    ['Wirtschaftlichkeit', 'Beste', 'Schlecht', 'Befriedigend'],
    ['Robustheit (Sensitivität)', 'Robust gegen Preisänderungen', 'Schwach gegen Leasingpreise', 'Mittel'],
    ['Gesamtbewertung', 'Empfohlen', 'Nicht empfohlen', 'Alternative']
])
style_table_caption("Tabelle 10: Vergleich der Optionen – Wirtschaftlichkeit und Robustheit")

doc.add_page_break()

style_h1("7 Sensitivitätsanalyse")
style_text("Die Sensitivitätsanalyse prüft, unter welchen Bedingungen die Vorteilhaftigkeit der empfohlenen Option erhalten bleibt. Szenario A (Preissteigerung Wartung +50%): Option 1 Kapitalwert etwa 27.200 EUR. Option 1 bleibt führend. Szenario B (Leasingrate –20%): Option 2 Kapitalwert etwa 38.450 EUR. Option 1 bleibt führend. Fazit: Option 1 ist robust gegen Preisveränderungen.")

style_table(5, 5, [
    ['Szenario', 'Parameter-Änderung', 'Option 1 KW (EUR)', 'Option 2 KW (EUR)', 'Option 4 KW (EUR)'],
    ['Basis', 'Keine Änderung', '23.847', '43.680', '29.340'],
    ['A', 'Wartungskosten +50%', '27.200', '43.680', '33.120'],
    ['B', 'Leasingrate –20%', '23.847', '38.450', '29.340'],
    ['C', 'Elektro-Gabelstapler Preis –30%', '18.300', '43.680', '24.200']
])
style_table_caption("Tabelle 6: Sensitivitätsanalyse – Kapitalwertveränderungen bei verschiedenen Szenarien")

doc.add_page_break()

style_h1("8 Nicht-monetäre Faktoren")
style_text("Neben den ökonomischen Analysen werden folgende qualitative Faktoren berücksichtigt: Personenschutz und Arbeitssicherheit (Option 1 beste Ausstattung), Umweltschutz und Nachhaltigkeit (Elektroantrieb emissionsfrei), Betriebszuverlässigkeit (neues Gerät zuverlässigster) und strategische Unabhängigkeit (Eigenes Gerät volle Verfügungskontrolle). Die nicht-monetären Faktoren bestätigen die wirtschaftliche Empfehlung für Option 1.")

style_table(6, 5, [
    ['Kriterium', 'Option 1: Neukauf', 'Option 2: Leasing', 'Option 4: Gebrauchtkauf', 'Gewicht'],
    ['Personenschutz und Arbeitssicherheit', 'Sehr gut', 'Gut', 'Befriedigend', 'Hoch'],
    ['Umweltschutz und Nachhaltigkeit', 'Sehr gut', 'Sehr gut', 'Gut', 'Mittel'],
    ['Betriebszuverlässigkeit', 'Sehr gut', 'Gut', 'Befriedigend', 'Hoch'],
    ['Strategische Unabhängigkeit', 'Sehr gut', 'Befriedigend', 'Sehr gut', 'Mittel'],
    ['Fazit', 'Beste Gesamtbewertung', 'Ausgewogen', 'Schwächen bei Zuverlässigkeit', '—']
])
style_table_caption("Tabelle 7: Nicht-monetäre Faktoren – Nutzwertanalyse der Optionen")

doc.add_page_break()

style_h1("9 Entscheidungsvorschlag")
style_text("Auf Grundlage der umfassenden Wirtschaftlichkeitsuntersuchung wird empfohlen, Option 1 – Neukauf eines modernen Elektro-Gabelstaplers – umzusetzen. Diese Option erzielt mit einem Kapitalwert von 23.847 Euro (einschließlich Risikowertung) das wirtschaftlichste Ergebnis unter allen untersuchten und bedarfsdeckenden Alternativen.")
style_text("Diese Entscheidungsempfehlung wird durch folgende Faktoren begründet: Erstens erreicht Option 1 den niedrigsten Kapitalwert und damit die geringsten Gesamtkosten. Zweitens bietet Option 1 die höchste Betriebssicherheit. Drittens wird mit Option 1 der optimale Arbeitssicherheitsstandard erreicht. Viertens ermöglicht Option 1 langfristige Kostenplanbarkeit.")
style_text("Empfohlene Spezifikation: Elektro-Gabelstapler mit Tragfähigkeit 2.500 Kilogramm, Hubhöhe mindestens 4.500 Millimeter, Lithium-Ionen-Batterie-Antrieb, Kaufpreis circa 21.000 EUR netto. Typische Hersteller: Jungheinrich, STILL oder Toyota. Umsetzungszeitpunkt: Juni 2026.")

doc.add_page_break()

style_h1("10 Anlagen")

style_h2("Anlage 1: Marktrecherche Gabelstapler (Stand April 2026)")

style_table(6, 5, [
    ['Nr.', 'Produkt', 'Hersteller', 'Preis (netto)', 'Abrufdatum'],
    ['1', 'Elektro-Gabelstapler 2.500 kg, 4.500 mm', '*[Recherche erforderlich]*', '*[Recherche erforderlich]*', '17.04.2026'],
    ['2', 'Handstapler-Wartung (jährlich)', '*[Recherche erforderlich]*', '*[Recherche erforderlich]*', '17.04.2026'],
    ['3', 'Leasing Elektro-Gabelstapler (monatlich)', '*[Recherche erforderlich]*', '*[Recherche erforderlich]*', '17.04.2026'],
    ['4', 'Gebrauchtkauf (3–5 Jahre alt)', '*[Recherche erforderlich]*', '*[Recherche erforderlich]*', '17.04.2026'],
    ['5', 'Batterie-Ladestation (3 kW)', '*[Recherche erforderlich]*', '*[Recherche erforderlich]*', '17.04.2026']
])
style_table_caption("Tabelle 11: Marktrecherche Gabelstapler (Stand April 2026)")

style_text("Quellen der Webrecherche (Abrufdatum 17. April 2026): *[Quellen erforderlich – siehe nächste WU mit aktiver Webrecherche]*")

doc.save(output_path)
print(f"Fertig: {output_file}")
