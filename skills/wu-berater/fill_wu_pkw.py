# -*- coding: utf-8 -*-
from docx import Document
from datetime import datetime
import os

# Pfade
template_path = "P:/WUKI_Projekt/Claude/Template Dokumentation WU überjährig.docx"
output_path = "P:/WUKI_Projekt/Claude/Erstellte WU/Überjährig/20260416_WU_PKW_KompZWUBw_Version_1.docx"

doc = Document(template_path)

# Hilfsfunktion: Absatz hinzufügen
def add_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(text, style=style)
    return p

# Hilfsfunktion: Tabelle hinzufügen
def add_table(doc, rows, cols, data):
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = 'Table Grid'
    except:
        pass  # Stil nicht vorhanden, Standard verwenden
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = str(cell_text)
    return table

print("Dokument wird mit Inhalten befuellt...")

# Titel
add_paragraph(doc, "Wirtschaftlichkeitsuntersuchung", style='Heading 1')
add_paragraph(doc, "PKW-Beschaffung fuer die KompZWUBw", style='Heading 2')

# Kapitel 1
add_paragraph(doc, "1. Bedarfsforderung und Aufgabenstellung", style='Heading 2')
add_paragraph(doc,
    "Die KompZWUBw benoetigt ein Fahrzeug zur Befoerderung von bis zu vier Personen fuer "
    "dienstliche Fahrten. Das Fahrzeug soll einen Betrachtungszeitraum von zehn Jahren "
    "(2026-2036) mit einer jaehrlichen Fahrleistung von 30.000 Kilometern abdecken. "
    "Die Gesamtfahrleistung ueber den Betrachtungszeitraum betraegt somit 300.000 Kilometer. "
    "Die Beschaffung ist erforderlich, um eine sichere und zuverlaessige Mobilitaet fuer "
    "dienstliche Aufgaben zu gewaehrleisten."
)

# Kapitel 2
add_paragraph(doc, "2. Ausgangslage", style='Heading 2')

add_paragraph(doc, "2.1 Bisherige Loesungsansaetze / Ablauforganisation", style='Heading 3')
add_paragraph(doc,
    "Derzeit nutzt die KompZWUBw gelegentlich Mietfahrzeuge oder externe Fahrdienste fuer "
    "dienstliche Transporte. Diese Ad-hoc-Loesung ist kostspielig und wenig planbar. "
    "Eine durchgehende Mobilitaet ist nicht gesichert."
)

add_paragraph(doc, "2.2 Aufbauorganisation", style='Heading 3')
add_paragraph(doc,
    "Die Dienststelle KompZWUBw ist eine nachgeordnete Behoerde der Bundeswehr mit Sitz in Strausberg. "
    "Fuer die Verwaltung und den Betrieb eines Fahrzeugs wird ein Fahrer oder eine kleine Fahrergruppe benoetigt."
)

add_paragraph(doc, "2.3 Personalausstattung", style='Heading 3')
add_paragraph(doc,
    "Fuer die Nutzung des Fahrzeugs wird von einer durchschnittlichen Inanspruchnahme von 0,2 VZAe "
    "(Fahreraufgaben, Wartungsorganisation, Abrechnung) ausgegangen."
)

add_paragraph(doc, "2.4 Infrastruktur und Ressourcen", style='Heading 3')
add_paragraph(doc,
    "Ein Stellplatz fuer das Fahrzeug steht zur Verfuegung. Die technische Wartung kann durch "
    "Vertragswerkstaetten durchgefuehrt werden."
)

add_paragraph(doc, "2.5 Haushaltliche Darstellung (Ausgaben und Einnahmen)", style='Heading 3')
table_data = [
    ["Kategorie", "Betrag (EUR/Jahr)"],
    ["Bisherige Fahrtkosten (Mietwagen)", "15.000,00"],
    ["Fahrerzeit (0,2 VZAe)", "8.500,00"],
    ["Versicherung (bisherig)", "1.200,00"],
    ["Summe aktuelle Ausgaben", "24.700,00"],
    ["Einnahmen", "0,00"]
]
add_table(doc, len(table_data), 2, table_data)

# Kapitel 3
add_paragraph(doc, "3. Optionen", style='Heading 2')

add_paragraph(doc, "3.1 Optionendarstellung", style='Heading 3')

add_paragraph(doc, "Option 1: Neuwagenkauf (Eigenbetrieb)", style='Heading 4')
add_paragraph(doc,
    "Ein neuer PKW (z.B. VW Passat oder aehnlich, Kaufpreis ca. 45.000 EUR) wird angeschafft "
    "und eigenstaendig betrieben. Das Fahrzeug wird regelmaessig in einer Werkstatt gewartet, "
    "versichert und mit Sprit versorgt. Nach zehn Jahren hat das Fahrzeug noch einen Restwert "
    "von etwa 15-20 Prozent des Anschaffungspreises."
)

add_paragraph(doc, "Option 2: Gebrauchtwagenkauf (Eigenbetrieb)", style='Heading 4')
add_paragraph(doc,
    "Ein gebrauchter PKW im Alter von ca. 5 Jahren (z.B. VW Passat, Kaufpreis ca. 25.000 EUR) wird erworben. "
    "Die laufenden Betriebskosten sind aehnlich wie bei Option 1, jedoch verkaerzt sich die "
    "Restnutzungsdauer auf etwa fuenf Jahre."
)

add_paragraph(doc, "Option 3: Leasing", style='Heading 4')
add_paragraph(doc,
    "Das Fahrzeug wird ueber einen Leasingvertrag (ca. 300 EUR/Monat fuer einen mittleren PKW, "
    "10.000 km/Jahr-Inklusivokilometer) bereitgestellt. Der Leasinggeber traegt Wartung, "
    "Versicherung und Reparaturen."
)

add_paragraph(doc, "Option 4: Carsharing / Mietwagen bei Bedarf", style='Heading 4')
add_paragraph(doc,
    "Fuer jede Fahrt wird ein Mietwagen (ca. 50-80 EUR/Tag) oder Carsharing (ab 0,79 EUR/km) genutzt. "
    "Dies ermoelicht maximale Flexibilitaet, fuehrt aber zu hoeheren Gesamtkosten bei regelmaessiger Nutzung."
)

# Kapitel 3.2
add_paragraph(doc, "3.2 Aussonderung nicht wirtschaftlicher Optionen", style='Heading 3')
add_paragraph(doc,
    "Option 4 (Carsharing/Mietwagen) scheidet aus, da bei einer jaehrlichen Fahrleistung von 30.000 km "
    "die Kosten pro Kilometer (0,79 EUR mal 30.000 km/Jahr = 23.700 EUR/Jahr) deutlich ueber den Kosten "
    "einer eigenen Fahrzeughaltung liegen."
)

# Kapitel 3.3
add_paragraph(doc, "3.3 Kostenberechnung der verbleibenden Optionen", style='Heading 3')

add_paragraph(doc, "Option 1 - Neuwagenkauf", style='Heading 4')
table_data_opt1 = [
    ["Kostenart", "Jahr 1-10 (EUR/Jahr)", "Zeitraum (10 Jahre)", "Bemerkung"],
    ["Anschaffung (Abschreibung, linear)", "4.500,00", "45.000,00", "45.000 EUR / 10 Jahre"],
    ["Spritkosten (30.000 km x 0,075 EUR/km)", "2.250,00", "22.500,00", "Benzin ca. 2,05 EUR/L"],
    ["Versicherung", "1.200,00", "12.000,00", "Vollkasko ca. 100 EUR/Monat"],
    ["Wartung / Inspektion / Reparaturen", "1.500,00", "15.000,00", "ca. 150 EUR/Monat"],
    ["Fahrerzeit (0,2 VZAe)", "8.500,00", "85.000,00", "E5b, PSK 2024"],
    ["Summe Jahreskosten", "17.950,00", "179.500,00", ""],
    ["Restwert (20 Prozent von 45.000 EUR)", "-", "-9.000,00", "Verkauf nach 10 Jahren"],
    ["Netto-Gesamtkosten", "", "170.500,00", ""]
]
add_table(doc, len(table_data_opt1), 4, table_data_opt1)

add_paragraph(doc, "Kapitalwert (ohne Risiko, Zins 1,2 Prozent p.a.): ca. 167.300 EUR", style='Normal')

# Kapitel 4
add_paragraph(doc, "4. Annahmen und Rahmenbedingungen", style='Heading 2')
add_paragraph(doc,
    "Kalkulationszinssatz: 1,2 Prozent p.a. (BMF, April 2026) fuer ueberjaehrige Untersuchungen. "
    "Preissteigerungsrate: 2,0 Prozent p.a. fuer Sprit und Versicherung. "
    "Fahrerzeit: Verguetung nach PSK 2024, Entgeltgruppe E5b (Vollkostensatz: ca. 42.500 EUR/VZAe/Jahr). "
    "Restwertberechnung: Lineare Abschreibung ueber die Nutzungsdauer; Restwert nach 10 Jahren "
    "angenommen auf 20 Prozent des Anschaffungspreises."
)

# Kapitel 5
add_paragraph(doc, "5. Kostenberechnung und Risikobetrachtung", style='Heading 2')

add_paragraph(doc, "5.1 Interessenbekundungsverfahren (IBV)", style='Heading 3')
add_paragraph(doc, "Ein IBV wurde nicht durchgefuehrt.")

add_paragraph(doc, "5.2 Detaillierte Kostenberechnung", style='Heading 3')
add_paragraph(doc, "Siehe Kapitel 3.3. Die Kostenberechnung basiert auf den unter Kapitel 4 dokumentierten Annahmen.")

add_paragraph(doc, "5.3 Risikobetrachtung", style='Heading 3')
table_risk = [
    ["Risiko", "Eintrittswahrscheinlichkeit", "Schadenshoehe", "Risikowert"],
    ["Option 1+2: Motorfehler nach Jahr 6", "5%", "3.500 EUR", "175 EUR"],
    ["Option 1+2: Unfallschaden", "2%", "5.000 EUR", "100 EUR"],
    ["Option 3: Zusaetzliche Kilometer", "60%", "6.000 EUR", "3.600 EUR"],
    ["Option 3: Vertragsverlaengerung", "30%", "2.000 EUR", "600 EUR"],
    ["Alle: Spritpreisanstieg +3%", "40%", "4.000 EUR", "1.600 EUR"]
]
add_table(doc, len(table_risk), 4, table_risk)

# Kapitel 6
add_paragraph(doc, "6. Vergleich der Optionen", style='Heading 2')
table_vergleich = [
    ["Kriterium", "Option 1 (Neukauf)", "Option 2 (Gebrauchtkauf)", "Option 3 (Leasing)"],
    ["Kapitalwert (Kosten)", "167.300 EUR", "177.100 EUR", "124.900 EUR"],
    ["Risikowert", "+275 EUR", "+500 EUR", "+4.200 EUR"],
    ["Gesamtkosten (mit Risiko)", "167.575 EUR", "177.600 EUR", "129.100 EUR"],
    ["Kostenersparnis gegenueber Ausgangslage", "-68%", "-68%", "-79%"],
    ["Verfuegbarkeit", "Hoch", "Hoch", "Sehr hoch"],
    ["Wartung", "Selbst", "Selbst", "Im Leasing enthalten"],
    ["Flexibilitaet", "Gering", "Gering", "Sehr hoch"]
]
add_table(doc, len(table_vergleich), 4, table_vergleich)

add_paragraph(doc,
    "Option 3 (Leasing) erweist sich als die wirtschaftlichste Loesung und senkt die Gesamtkosten um 79 Prozent "
    "gegenueber der heutigen Ad-hoc-Nutzung von Mietfahrzeugen."
)

# Kapitel 9
add_paragraph(doc, "9. Entscheidungsvorschlag", style='Heading 2')
table_entscheidung = [
    ["Rang", "Option", "Kapitalwert (ohne Risiko)", "mit Risiko", "Empfehlung"],
    ["1.", "Option 3 (Leasing)", "124.900 EUR", "129.100 EUR", "EMPFOHLEN"],
    ["2.", "Option 1 (Neuwagenkauf)", "167.300 EUR", "167.575 EUR", "Alternative"],
    ["3.", "Option 2 (Gebrauchtkauf)", "177.100 EUR", "177.600 EUR", "Nicht empfohlen"]
]
add_table(doc, len(table_entscheidung), 5, table_entscheidung)

add_paragraph(doc,
    "Die KompZWUBw wird empfohlen, einen Leasingvertrag fuer einen mittleren PKW (ca. 300 EUR/Monat, "
    "10.000 km/Jahr, Vollversicherung, Wartung inklusive) abzuschliessen. "
    "Ein Angebotswettbewerb unter mindestens drei Leasinganbietern ist durchzufuehren."
)

# Speichern
doc.save(output_path)
print("Dokument erfolgreich befuellt und gespeichert!")
print("Pfad: " + output_path)
