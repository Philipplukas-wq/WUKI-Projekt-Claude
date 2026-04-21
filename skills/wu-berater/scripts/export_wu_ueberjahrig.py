"""
WU-Berater: Export-Skript für überjährige Wirtschaftlichkeitsuntersuchungen
Füllt das Word-Template 'Template Dokumentation WU überjährig.docx' mit WU-Inhalten.

Verwendung:
    from export_wu_ueberjahrig import fill_template
    fill_template(wu_data, output_path)

wu_data-Struktur: siehe Abschnitt WU_DATA_SCHEMA am Ende dieser Datei.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os
from wu_file_handler import save_document_safely

# ── Überschriften-Formatierungsvorgaben ────────────────────────────────────────
HEADING_FONT   = 'BundesSans Office'
HEADING_COLOR  = RGBColor(0x00, 0x00, 0x00)   # Schwarz

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), '..', '..', '..', 'Template Dokumentation WU überjährig.docx')
TEMPLATE_PATH = os.path.normpath(TEMPLATE_PATH)

OUTPUT_DIR = os.path.normpath(os.path.join(os.path.dirname(__file__), '..', '..', '..', 'Erstellte WU', 'Überjährige'))


# ── Hilfsfunktionen ────────────────────────────────────────────────────────────

def _format_heading_run(run):
    """Setzt BundesSans Office, schwarz, fett auf einem Heading-Run."""
    run.font.name     = HEADING_FONT
    run.font.color.rgb = HEADING_COLOR
    run.bold          = True


def format_alle_ueberschriften(doc: Document) -> None:
    """
    Durchläuft alle Heading-Paragraphen im Dokument und setzt:
      - Schriftart: BundesSans Office
      - Farbe: schwarz
      - Fett: ja
      - Abstand nach Absatz: 0 pt (kein Leerzeile zwischen Überschrift und Text)
    """
    from docx.shared import Pt as _Pt
    for para in doc.paragraphs:
        if not para.style.name.startswith('Heading'):
            continue
        # Paragraph-Abstand nach = 0 (kein Extra-Zeilenumbruch)
        para.paragraph_format.space_after = _Pt(0)
        # Jeden Run formatieren
        for run in para.runs:
            _format_heading_run(run)
        # Wenn der Paragraph Runs hat, auch den XML-font-Eintrag setzen
        # (manche Styles überschreiben den Run-Font via rPr-Ableitung)
        if para.runs:
            for run in para.runs:
                rPr = run._r.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'),    HEADING_FONT)
                rFonts.set(qn('w:hAnsi'),    HEADING_FONT)
                rFonts.set(qn('w:eastAsia'), HEADING_FONT)
                rFonts.set(qn('w:cs'),       HEADING_FONT)
                # Vorhandene rFonts ersetzen
                existing = rPr.find(qn('w:rFonts'))
                if existing is not None:
                    rPr.remove(existing)
                rPr.insert(0, rFonts)


def set_text(para, text, size=10, bold=False):
    para.clear()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold

def find_heading(doc, text, style):
    for i, p in enumerate(doc.paragraphs):
        if text in p.text and p.style.name == style:
            return i
    return -1

def insert_para_after(para, doc) -> object:
    """Fügt einen neuen leeren Normal-Paragraphen direkt nach 'para' ein."""
    new_p = OxmlElement('w:p')
    para._element.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.style = doc.styles['Normal']
            return p
    return None


def insert_table_after(para, tbl_element) -> None:
    """Verschiebt ein bestehendes Tabellen-XML-Element direkt nach 'para'."""
    para._element.addnext(tbl_element)


def set_after_heading(doc, heading_text, style, content, size=10):
    """
    Schreibt 'content' in den ersten Normal-Paragraphen nach der Überschrift.
    Falls der nächste Paragraph selbst eine Überschrift ist, wird vorher
    ein neuer Normal-Paragraph eingefügt (Überschriften werden nie überschrieben).
    """
    i = find_heading(doc, heading_text, style)
    if i < 0:
        return
    next_para = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
    if next_para is None:
        return
    # Prüfen ob der nächste Paragraph eine Überschrift ist
    if next_para.style.name.startswith('Heading'):
        # Neuen Paragraphen einfügen, ohne die Überschrift zu überschreiben
        target = insert_para_after(doc.paragraphs[i], doc)
    else:
        target = next_para
    if target:
        set_text(target, content, size)

def add_heading_after_para(doc, para, text, level):
    styles = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
    new_p = OxmlElement('w:p')
    para._element.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.style = doc.styles[styles[level]]
            p.paragraph_format.space_after = Pt(0)
            p.clear()
            run = p.add_run(text)
            run.font.size = Pt(12 if level <= 2 else 11)
            _format_heading_run(run)
            return p
    return None

def add_text_para_after(doc, para, text, size=10):
    new_p = OxmlElement('w:p')
    para._element.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.style = doc.styles['Normal']
            run = p.add_run(text)
            run.font.size = Pt(size)
            return p
    return None


def insert_toc_and_lof(doc: Document, insert_after_para_index: int) -> None:
    """
    Fügt Inhaltsverzeichnis (TOC) und Tabellenverzeichnis (LOF) nach einem
    bestimmten Paragraph ein, mit Seitenwechsel vorher und nachher.

    :param doc: Das Document-Objekt
    :param insert_after_para_index: Index des Paragraphs, nach dem eingefügt werden soll
                                    (typisch: Index nach Dokumentkopf)
    """
    # Stelle sicher, dass der Index gültig ist
    if insert_after_para_index < 0 or insert_after_para_index >= len(doc.paragraphs):
        return

    anchor_para = doc.paragraphs[insert_after_para_index]

    # ── Seitenwechsel vor dem Verzeichnis ────────────────────────────────────────
    page_break_p = OxmlElement('w:p')
    anchor_para._element.addnext(page_break_p)

    # Finde den neu eingefügten Paragraph
    for p in doc.paragraphs:
        if p._element is page_break_p:
            # Seitenwechsel = <w:br w:type="page"/>
            pPr = p._element.get_or_add_pPr()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            pPr.append(br)
            break

    # ── Überschrift: Inhaltsverzeichnis ──────────────────────────────────────────
    toc_heading_p = OxmlElement('w:p')
    page_break_p.addnext(toc_heading_p)

    for p in doc.paragraphs:
        if p._element is toc_heading_p:
            p.style = doc.styles['Heading 1']
            run = p.add_run('Inhaltsverzeichnis')
            _format_heading_run(run)
            p.paragraph_format.space_after = Pt(0)
            break

    # ── TOC-Feld einfügen ───────────────────────────────────────────────────────
    toc_p = OxmlElement('w:p')
    toc_heading_p.addnext(toc_p)

    for p in doc.paragraphs:
        if p._element is toc_p:
            # Word-Feld für Inhaltsverzeichnis: \o "1-3" erzeugt 1–3 Ebenen
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            p_elem = p._element
            p_elem.append(fldChar1)
            p_elem.append(instrText)
            p_elem.append(fldChar2)

            # Standardtext (wird beim Aktualisieren überschrieben)
            run = p.add_run('[Inhaltsverzeichnis wird hier eingefügt]')
            run.font.italic = True
            run.font.size = Pt(9)
            break

    # ── Seitenwechsel nach TOC ──────────────────────────────────────────────────
    page_break_p2 = OxmlElement('w:p')
    toc_p.addnext(page_break_p2)

    for p in doc.paragraphs:
        if p._element is page_break_p2:
            pPr = p._element.get_or_add_pPr()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            pPr.append(br)
            break

    # ── Überschrift: Tabellenverzeichnis ────────────────────────────────────────
    lof_heading_p = OxmlElement('w:p')
    page_break_p2.addnext(lof_heading_p)

    for p in doc.paragraphs:
        if p._element is lof_heading_p:
            p.style = doc.styles['Heading 1']
            run = p.add_run('Tabellenverzeichnis')
            _format_heading_run(run)
            p.paragraph_format.space_after = Pt(0)
            break

    # ── LOF-Feld einfügen (nutzt Caption-Style) ────────────────────────────────
    lof_p = OxmlElement('w:p')
    lof_heading_p.addnext(lof_p)

    for p in doc.paragraphs:
        if p._element is lof_p:
            # Word-Feld für Tabellenverzeichnis (Abbildungen/Tabellen)
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'TOC \\h \\z \\c "Tabellentitel"'  # Nutzt Caption-Style

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            p_elem = p._element
            p_elem.append(fldChar1)
            p_elem.append(instrText)
            p_elem.append(fldChar2)

            # Standardtext (wird beim Aktualisieren überschrieben)
            run = p.add_run('[Tabellenverzeichnis wird hier eingefügt]')
            run.font.italic = True
            run.font.size = Pt(9)
            break

    # ── Seitenwechsel nach LOF ──────────────────────────────────────────────────
    page_break_p3 = OxmlElement('w:p')
    lof_p.addnext(page_break_p3)

    for p in doc.paragraphs:
        if p._element is page_break_p3:
            pPr = p._element.get_or_add_pPr()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            pPr.append(br)
            break

def add_tbl_header(tbl, headers, fill='1F3864'):
    row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

def add_tbl_row(tbl, values, bold=False, fill=None):
    row = tbl.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(9)
        run.bold = bold
        if fill:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)

def set_cell(cell, text, bold=False, size=10, color=None):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


# ── Haupt-Exportfunktion ───────────────────────────────────────────────────────

def fill_template(wu_data: dict, output_path: str):
    """
    Befüllt das überjährige WU-Template mit den übergebenen Daten und speichert
    das Ergebnis unter output_path.

    :param wu_data: Dictionary mit WU-Inhalten (Struktur siehe WU_DATA_SCHEMA)
    :param output_path: Pfad der Ausgabedatei (.docx)
    """
    doc = Document(TEMPLATE_PATH)
    meta    = wu_data.get('meta', {})
    kap1    = wu_data.get('kap1', {})
    kap2    = wu_data.get('kap2', {})
    kap3    = wu_data.get('kap3', {})
    kap4    = wu_data.get('kap4', {})
    kap5    = wu_data.get('kap5', {})
    kap6_9  = wu_data.get('kap6_9', {})
    anlage  = wu_data.get('anlage', [])

    # ── Titel ────────────────────────────────────────────────────────────────
    set_text(doc.paragraphs[4],  meta.get('titel', ''), size=16, bold=True)
    set_text(doc.paragraphs[6],  meta.get('kurztitel', ''))
    set_text(doc.paragraphs[36], meta.get('titel', ''), size=16, bold=True)
    set_text(doc.paragraphs[38], meta.get('kurztitel', ''))
    for p in doc.paragraphs:
        if 'Dokumentinformationen' in p.text:
            set_text(p, (
                f"Dienststelle: {meta.get('dienststelle', '')}  |  "
                f"Bearbeiter: {meta.get('bearbeiter', '')}  |  "
                f"Datum: {meta.get('datum', '')}  |  "
                f"Schutz: {meta.get('schutz', 'offen')}  |  "
                f"Version: {meta.get('version', '1')}"
            ))
            break

    # ── Tabelle 2: Dokumentkopf ──────────────────────────────────────────────
    # Zeile 0: Az (Aktenzeichen) – nicht überschreiben, manuell zu ergänzen
    # Zeile 1: Stand (Datum)
    # Zeile 2: Version
    # Zeile 5: Einstufung (Schutz)
    tbl2 = doc.tables[2]
    tbl2.rows[1].cells[1].paragraphs[0].text = meta.get('datum', '')
    tbl2.rows[2].cells[1].paragraphs[0].text = meta.get('version', '1')
    tbl2.rows[5].cells[1].paragraphs[0].text = meta.get('schutz', 'offen')

    # ── Tabelle 3: Projektinfo ───────────────────────────────────────────────
    tbl3 = doc.tables[3]
    tbl3.rows[0].cells[1].paragraphs[0].text = meta.get('titel', '')
    tbl3.rows[1].cells[1].paragraphs[0].text = f"Wirtschaftlichkeitsuntersuchung\n{meta.get('titel', '')}"
    tbl3.rows[3].cells[1].paragraphs[0].text = meta.get('dienststelle', '')
    tbl3.rows[4].cells[1].paragraphs[0].text = f"Bearbeiter: {meta.get('bearbeiter', '')}"

    # ── Inhalts- und Tabellenverzeichnis einfügen ───────────────────────────────
    # Finde den letzten Paragraph des Dokumentkopfs (üblicherweise vor "Überblick")
    # und füge die Verzeichnisse danach ein
    dokumentkopf_end_idx = find_heading(doc, 'Überblick', 'Heading 1')
    if dokumentkopf_end_idx > 0:
        insert_toc_and_lof(doc, dokumentkopf_end_idx - 1)

    # ── Überblick ────────────────────────────────────────────────────────────
    set_after_heading(doc, 'Betrachtungsgegenstand', 'Heading 2',
        wu_data.get('ueberblick', {}).get('betrachtungsgegenstand', ''))
    set_after_heading(doc, 'Entscheidungsvorschlag', 'Heading 2',
        wu_data.get('ueberblick', {}).get('entscheidungsvorschlag', ''))

    # ── Tabelle 4: Entscheidungsvorschlag-Übersicht ──────────────────────────
    # Zeilenanzahl dynamisch: genau 1 Headerzeile + n Optionen (keine Leerzeilen)
    optionen_ue = kap6_9.get('optionen_uebersicht', [])
    n_opt = len(optionen_ue)

    # Erstelle Table 4 falls nicht vorhanden
    if len(doc.tables) < 5:
        tbl4 = doc.add_table(rows=1, cols=3)
        tbl4.style = 'Table Grid'
        add_tbl_header(tbl4, ['Option', 'Kapitalwert (ohne Risiko)', 'Kapitalwert (mit Risiko)'])
    else:
        tbl4 = doc.tables[4]

    # Fehlende Zeilen hinzufügen
    while len(tbl4.rows) - 1 < n_opt:
        tbl4.add_row()

    # Überzählige Leerzeilen entfernen (von unten)
    while len(tbl4.rows) - 1 > n_opt:
        row_to_del = tbl4.rows[-1]
        row_to_del._element.getparent().remove(row_to_del._element)

    # Optionen eintragen
    for i, opt in enumerate(optionen_ue):
        row = tbl4.rows[i + 1]
        is_best = opt.get('empfohlen', False)
        color = RGBColor(0x1F, 0x38, 0x64) if is_best else None
        set_cell(row.cells[0], opt.get('name', ''),            bold=is_best, color=color)
        set_cell(row.cells[1], opt.get('kw_ohne_risiko', ''),  bold=is_best, color=color)
        set_cell(row.cells[2], opt.get('kw_mit_risiko', ''),   bold=is_best, color=color)
        if is_best:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D9E1F2')
                tcPr.append(shd)

    # ── Kapitel 1 ────────────────────────────────────────────────────────────
    set_after_heading(doc, 'Funktionale Bedarfsforderung', 'Heading 2', kap1.get('bedarfsforderung', ''))
    set_after_heading(doc, 'Bedarfsprognose', 'Heading 2',              kap1.get('bedarfsprognose', ''))
    set_after_heading(doc, 'Rechtliche Rahmenbedingungen', 'Heading 3', kap1.get('rb_rechtlich', 'Entfällt.'))
    set_after_heading(doc, 'Organisatorische Rahmenbedingungen', 'Heading 3', kap1.get('rb_organisatorisch', 'Entfällt.'))
    set_after_heading(doc, 'Zeitliche Rahmenbedingungen', 'Heading 3',  kap1.get('rb_zeitlich', 'Entfällt.'))
    set_after_heading(doc, 'Sonstige Rahmenbedingungen', 'Heading 3',   kap1.get('rb_sonstige', 'Entfällt.'))

    # ── Kapitel 2 ────────────────────────────────────────────────────────────
    import sys as _sys
    _sys.path.insert(0, os.path.dirname(__file__))
    from bwdlz_grafiken import (ist_bwdlz, erzeuge_organigramm, erzeuge_workflow,
                                 erzeuge_workflow_generisch, erzeuge_organigramm_generisch)

    _sachverhalt = (
        kap1.get('bedarfsforderung', '') + ' ' +
        wu_data.get('ueberblick', {}).get('betrachtungsgegenstand', '')
    )
    _ist_bwdlz = ist_bwdlz(meta.get('dienststelle', ''))

    # 2.1 Ablauforganisation – immer als Workflow-Grafik
    i_ablauf = find_heading(doc, 'Ablauforganisation', 'Heading 2')
    if i_ablauf >= 0:
        ref_ablauf = doc.paragraphs[i_ablauf + 1]
        if not ref_ablauf.style.name.startswith('Heading'):
            set_text(ref_ablauf, '')
        else:
            ref_ablauf = insert_para_after(doc.paragraphs[i_ablauf], doc)
        if _ist_bwdlz:
            erzeuge_workflow(doc, _sachverhalt)
        else:
            erzeuge_workflow_generisch(doc, kap2.get('ablauforganisation', ''))
        insert_table_after(ref_ablauf, doc.tables[-1]._tbl)

    # 2.2 Aufbauorganisation – immer als Organigramm-Grafik
    i_aufbau = find_heading(doc, 'Aufbauorganisation', 'Heading 2')
    if i_aufbau >= 0:
        ref_aufbau = doc.paragraphs[i_aufbau + 1]
        if not ref_aufbau.style.name.startswith('Heading'):
            set_text(ref_aufbau, '')
        else:
            ref_aufbau = insert_para_after(doc.paragraphs[i_aufbau], doc)
        if _ist_bwdlz:
            erzeuge_organigramm(doc, _sachverhalt, meta.get('dienststelle', ''))
        else:
            erzeuge_organigramm_generisch(doc, kap2.get('aufbauorganisation', ''),
                                          meta.get('dienststelle', ''))
        insert_table_after(ref_aufbau, doc.tables[-1]._tbl)

    set_after_heading(doc, 'Personal', 'Heading 2',                 kap2.get('personal', ''))
    set_after_heading(doc, 'Material', 'Heading 2',                 kap2.get('material', ''))
    set_after_heading(doc, 'Infrastruktur', 'Heading 2',            kap2.get('infrastruktur', ''))
    set_after_heading(doc, 'Sach- und Dienstleistungen', 'Heading 2', kap2.get('sach_dienstleistungen', ''))
    set_after_heading(doc, 'Ggf. Einnahmen', 'Heading 2',           kap2.get('einnahmen', 'Keine Einnahmen.'))
    set_after_heading(doc, 'Haushalterische Darstellung', 'Heading 2', kap2.get('haushalterische_darstellung', ''))

    # Tabelle Haushalterische Darstellung: per Inhalt suchen (Index kann sich nach
    # Grafik-Einfügungen verschieben). Erkennungsmerkmal: erste Zelle = "Position"
    # oder erste Zelle enthält "Jahr der Ausgangslage".
    tbl_haush = None
    for t in doc.tables:
        if len(t.rows) >= 2 and len(t.columns) >= 4:
            z0 = t.rows[0].cells[0].text.strip()
            z1 = t.rows[1].cells[0].text.strip() if len(t.rows) > 1 else ''
            if 'Jahr der Ausgangslage' in z0 or z0 == 'Position' or z1 == 'Personal':
                tbl_haush = t
                break

    if tbl_haush:
        # "Jahr der Ausgangslage"-Zeile entfernen falls noch vorhanden
        if 'Jahr der Ausgangslage' in tbl_haush.rows[0].cells[0].text:
            r = tbl_haush.rows[0]
            r._element.getparent().remove(r._element)
        # Jetzt: Zeile 0 = Spaltenheader, Zeile 1+ = Datenzeilen
        for i, row_data in enumerate(kap2.get('haushalterische_tabelle', [])):
            if i + 1 < len(tbl_haush.rows):
                row = tbl_haush.rows[i + 1]
                for j, val in enumerate(row_data):
                    if j < len(row.cells):
                        row.cells[j].paragraphs[0].text = str(val)

    # ── Kapitel 3.1 ──────────────────────────────────────────────────────────
    set_after_heading(doc, 'Grundsätzlich', 'Heading 2', kap3.get('optionen_grundsaetzlich', ''))
    set_after_heading(doc, 'Aussonderung von ungeeigneten', 'Heading 2', kap3.get('aussonderung', ''))

    # ── Kapitel 3.3 – Optionen ausführlich ───────────────────────────────────
    optionen = kap3.get('optionen_detail', [])
    felder = ['ablauforganisation', 'aufbauorganisation', 'personal',
              'material', 'infrastruktur', 'sach_dienstleistungen', 'einnahmen']
    feld_labels = ['Ablauforganisation', 'Aufbauorganisation', 'Personal',
                   'Material', 'Infrastruktur', 'Sach- und Dienstleistungen', 'Ggf. Einnahmen']

    if optionen:
        # Option 1: in das vorhandene Template einfügen
        opt1 = optionen[0]
        for feld, label in zip(felder, feld_labels):
            i_h = find_heading(doc, label, 'Heading 4')
            if i_h >= 0:
                set_text(doc.paragraphs[i_h + 1], opt1.get(feld, ''))

        # Optionen 2+ als neue Abschnitte einfügen
        i_ein1 = find_heading(doc, 'Ggf. Einnahmen', 'Heading 4')
        ref = doc.paragraphs[i_ein1 + 1]

        for opt_nr, opt in enumerate(optionen[1:], start=2):
            h_opt = add_heading_after_para(doc, ref, f'3.3.{opt_nr} {opt.get("titel", f"Option {opt_nr}")}', 3)
            last = h_opt
            for sub_nr, (feld, label) in enumerate(zip(felder, feld_labels), start=1):
                h_sub = add_heading_after_para(doc, last, f'3.3.{opt_nr}.{sub_nr} {label}', 4)
                last  = add_text_para_after(doc, h_sub, opt.get(feld, ''))
            ref = last

    # ── Kapitel 4 ────────────────────────────────────────────────────────────
    set_after_heading(doc, 'Annahmen für alle Optionen',      'Heading 2', kap4.get('alle_optionen', ''))
    set_after_heading(doc, 'Annahmen für bestimmte Optionen', 'Heading 2', kap4.get('bestimmte_optionen', 'Entfällt.'))
    set_after_heading(doc, 'Annahmen für die Berechnung',     'Heading 2', kap4.get('berechnung', ''))

    # ── Kapitel 5 ────────────────────────────────────────────────────────────
    set_after_heading(doc, 'Interessenbekundungsverfahren', 'Heading 2',
        kap5.get('ibv', 'Entfällt (wird zu einem späteren Zeitpunkt ergänzt).'))

    # Tabelle Kap 5.2
    i_52 = find_heading(doc, 'Berechnung der Option', 'Heading 2')
    if i_52 >= 0 and kap5.get('berechnung_text'):
        ref_para = doc.paragraphs[i_52 + 1]
        set_text(ref_para, kap5['berechnung_text'])
        if kap5.get('berechnung_tabelle'):
            tbl52 = doc.add_table(rows=1, cols=len(kap5['berechnung_tabelle']['headers']))
            tbl52.style = 'Table Grid'
            add_tbl_header(tbl52, kap5['berechnung_tabelle']['headers'])
            for ri, row_data in enumerate(kap5['berechnung_tabelle']['rows']):
                is_sep = row_data.get('separator', False)
                is_sum = row_data.get('summe', False)
                fill = 'D9E1F2' if (is_sep or is_sum) else None
                add_tbl_row(tbl52, row_data['values'], bold=is_sum, fill=fill)
            ref_para._element.addnext(tbl52._tbl)

    set_after_heading(doc, 'Kapitalwerte ohne Risiko', 'Heading 2', kap5.get('kw_ohne_risiko', ''))

    # 5.4.1 Risikoidentifizierung – Tabelle wenn vorhanden, sonst Text
    # Überschrift immer erhalten – Text/Tabelle NUR darunter einfügen
    risiko_tbl = kap5.get('risiko_tabelle')
    if risiko_tbl:
        i_rid = find_heading(doc, 'Risikoidentifizierung', 'Heading 3')
        if i_rid >= 0:
            next_rid = doc.paragraphs[i_rid + 1]
            if next_rid.style.name.startswith('Heading'):
                ref_rid = insert_para_after(doc.paragraphs[i_rid], doc)
            else:
                set_text(next_rid, '')
                ref_rid = next_rid
            tbl_rid = doc.add_table(rows=1, cols=len(risiko_tbl['headers']))
            tbl_rid.style = 'Table Grid'
            add_tbl_header(tbl_rid, risiko_tbl['headers'])
            for row_data in risiko_tbl['rows']:
                is_sum = row_data.get('summe', False)
                fill = 'D9E1F2' if is_sum else None
                add_tbl_row(tbl_rid, row_data['values'], bold=is_sum, fill=fill)
            ref_rid._element.addnext(tbl_rid._tbl)
    else:
        set_after_heading(doc, 'Risikoidentifizierung', 'Heading 3', kap5.get('risikoidentifizierung', ''))

    set_after_heading(doc, 'Risikoverteilung', 'Heading 3', kap5.get('risikoverteilung', ''))

    # 5.4.3 Monetäre Risikobewertung – Tabelle wenn vorhanden, sonst Text
    # Überschrift immer erhalten – Text/Tabelle NUR darunter einfügen
    risiko_bew_tbl = kap5.get('risikobewertung_tabelle')
    if risiko_bew_tbl:
        i_rbw = find_heading(doc, 'Risikobewertung', 'Heading 3')
        if i_rbw >= 0:
            next_rbw = doc.paragraphs[i_rbw + 1]
            if next_rbw.style.name.startswith('Heading'):
                ref_rbw = insert_para_after(doc.paragraphs[i_rbw], doc)
            else:
                set_text(next_rbw, '')
                ref_rbw = next_rbw
            tbl_rbw = doc.add_table(rows=1, cols=len(risiko_bew_tbl['headers']))
            tbl_rbw.style = 'Table Grid'
            add_tbl_header(tbl_rbw, risiko_bew_tbl['headers'])
            for row_data in risiko_bew_tbl['rows']:
                is_sum = row_data.get('summe', False)
                fill = 'D9E1F2' if is_sum else None
                add_tbl_row(tbl_rbw, row_data['values'], bold=is_sum, fill=fill)
            ref_rbw._element.addnext(tbl_rbw._tbl)
    else:
        set_after_heading(doc, 'Risikobewertung', 'Heading 3', kap5.get('risikobewertung', ''))

    set_after_heading(doc, 'Kapitalwert mit Risiko', 'Heading 2', kap5.get('kw_mit_risiko', ''))

    # ── Kapitel 6: Vergleich als Tabelle ─────────────────────────────────────
    vergleich_tbl = kap6_9.get('vergleich_tabelle')
    if vergleich_tbl:
        i_v = find_heading(doc, 'Vergleich der Optionen', 'Heading 1')
        if i_v >= 0:
            ref_v = doc.paragraphs[i_v + 1]
            set_text(ref_v, kap6_9.get('vergleich_text', ''))
            tbl_v = doc.add_table(rows=1, cols=len(vergleich_tbl['headers']))
            tbl_v.style = 'Table Grid'
            add_tbl_header(tbl_v, vergleich_tbl['headers'])
            for rd in vergleich_tbl['rows']:
                is_best = rd.get('empfohlen', False)
                fill = 'D9E1F2' if is_best else None
                add_tbl_row(tbl_v, rd['values'], bold=is_best, fill=fill)
            ref_v._element.addnext(tbl_v._tbl)
    else:
        set_after_heading(doc, 'Vergleich der Optionen', 'Heading 1', kap6_9.get('vergleich', ''))

    set_after_heading(doc, 'Sensitivit',             'Heading 1', kap6_9.get('sensitivitaet', ''))
    set_after_heading(doc, 'Nicht monet',            'Heading 1', kap6_9.get('nichtmonetaer', ''))
    set_after_heading(doc, 'Entscheidungsvorschlag', 'Heading 1', kap6_9.get('entscheidungsvorschlag', ''))

    # ── Alle Überschriften einheitlich formatieren ────────────────────────────
    # BundesSans Office, schwarz, fett, kein Extra-Abstand nach Überschrift
    format_alle_ueberschriften(doc)

    # ── Anlage Marktrecherche ─────────────────────────────────────────────────
    if anlage:
        doc.add_page_break()
        h_anlage = doc.add_heading('Anlage: Marktrecherche', level=1)
        if h_anlage.runs:
            _format_heading_run(h_anlage.runs[0])
            h_anlage.paragraph_format.space_after = Pt(0)
        p_datum = doc.add_paragraph(f"Abrufdatum: {meta.get('datum', '')}")
        p_datum.runs[0].font.size = Pt(9)
        p_datum.runs[0].italic = True

        tbl_a = doc.add_table(rows=1, cols=4)
        tbl_a.style = 'Table Grid'
        add_tbl_header(tbl_a, ['Nr.', 'Produkt / Beschreibung', 'Preis', 'Quelle (URL)'])
        for eintrag in anlage:
            add_tbl_row(tbl_a, [
                eintrag.get('nr', ''),
                eintrag.get('produkt', ''),
                eintrag.get('preis', ''),
                eintrag.get('url', ''),
            ])
        note = doc.add_paragraph(
            'Hinweis: Screenshots der oben genannten Webseiten sind diesem Dokument '
            'als Anlage beizufügen (Einfügen → Objekt/Bild).')
        note.runs[0].font.size = Pt(9)
        note.runs[0].italic = True

    # Speichere mit robustem Lock-Handling
    actual_file = save_document_safely(doc, output_path)
    return actual_file


# ── Hilfsfunktionen für den Skill-Dialog ──────────────────────────────────────

def baue_haushalterische_tabelle(
    personal_eur: float = 0,
    personal_kap: str = '2.3',
    material_eur: float = 0,
    infrastruktur_eur: float = 0,
    dienstleistungen_eur: float = 0,
    einnahmen_eur: float = 0,
    jahr: str = ''
) -> list:
    """
    Baut die haushalterische Darstellung (Tabelle 5) automatisch aus
    den in Kapitel 2.3–2.7 erfassten Werten.

    :return: Liste von Zeilen für kap2['haushalterische_tabelle']
    """
    def fmt(betrag):
        return f'{betrag:,.0f} €'.replace(',', '.') if betrag else '–'

    gesamt_ausg = personal_eur + material_eur + infrastruktur_eur + dienstleistungen_eur
    gesamt_einnahmen = einnahmen_eur

    return [
        ['Personal',         personal_kap,              fmt(personal_eur),       '–'],
        ['Material',         '2.4',                     fmt(material_eur) if material_eur else '–', '–'],
        ['Infrastruktur',    '2.5',                     fmt(infrastruktur_eur) if infrastruktur_eur else '–', '–'],
        ['Dienstleistungen', '2.6 Sach- u. Dienstleist.',fmt(dienstleistungen_eur) if dienstleistungen_eur else '–', '–'],
        ['Gesamt',           '',                        fmt(gesamt_ausg),        fmt(gesamt_einnahmen) if gesamt_einnahmen else '–'],
    ]


def baue_vergleich_tabelle(optionen_uebersicht: list) -> dict:
    """
    Erzeugt die Vergleichstabelle für Kapitel 6 aus der Optionenübersicht.

    :param optionen_uebersicht: Liste aus erstelle_kw_uebersicht()
    :return: Dict mit headers und rows für kap6_9['vergleich_tabelle']
    """
    headers = ['Option', 'Kapitalwert ohne Risiko', 'Kapitalwert mit Risiko', 'Rang', 'Vorteil ggü. nächstbester']
    rows = []
    beste_kw = None

    import re as _re
    def _parse_kw(s):
        try:
            return float(_re.sub(r'[^\d]', '', str(s)))
        except Exception:
            return 0
    for rang, opt in enumerate(sorted(optionen_uebersicht, key=lambda x: _parse_kw(x['kw_mit_risiko'])), 1):
        try:
            kw_str = opt['kw_mit_risiko']
            # Entferne alle nicht-numerischen Zeichen außer Dezimaltrenner
            import re
            kw_num = float(re.sub(r'[^\d]', '', kw_str))
        except Exception:
            kw_num = 0

        if beste_kw is None:
            beste_kw = kw_num
            vorteil = '–'
        else:
            diff = kw_num - beste_kw
            vorteil = f'+{diff:,.0f} €'.replace(',', '.')

        rows.append({
            'values':    [opt['name'], opt['kw_ohne_risiko'], opt['kw_mit_risiko'], str(rang), vorteil],
            'empfohlen': opt.get('empfohlen', False),
        })

    return {'headers': headers, 'rows': rows}


def erstelle_abschlusscheckliste(wu_data: dict, output_path: str) -> str:
    """
    Gibt eine Abschlusscheckliste zurück: was wurde automatisch befüllt,
    was muss noch manuell ergänzt werden.

    :param wu_data:     Das wu_data-Dictionary
    :param output_path: Pfad der gespeicherten Datei
    :return:            Checkliste als formatierten String
    """
    meta = wu_data.get('meta', {})
    kap3 = wu_data.get('kap3', {})
    anlage = wu_data.get('anlage', [])

    auto_befuellt = [
        f'Dokumentkopf (Dienststelle, Bearbeiter, Datum, Schutz: {meta.get("schutz", "offen")})',
        'Inhaltsverzeichnis (automatisch aus Headings generiert)',
        'Tabellenverzeichnis (automatisch aus Tabellen generiert)',
        'Kapitel 1: Bedarfsforderung, Bedarfsprognose, Rahmenbedingungen',
        'Kapitel 2: Ausgangslage inkl. haushalterischer Darstellung',
        f'Kapitel 3: {len(kap3.get("optionen_detail", []))} Optionen vollständig beschrieben',
        'Kapitel 4: Annahmen (betrieblich + kalkulatorisch)',
        'Kapitel 5: Kapitalwertberechnung (automatisch), Risikobetrachtung',
        'Kapitel 6: Vergleich der Optionen (Tabelle)',
        'Kapitel 7: Sensitivitätsanalyse (Break-even)',
        'Kapitel 8: Nichtmonetäre Faktoren',
        'Kapitel 9: Entscheidungsvorschlag',
        f'Anlage Marktrecherche: {len(anlage)} Quellen',
        'Tabelle Entscheidungsvorschlag (Seite 6)',
        'Tabelle Haushalterische Darstellung',
    ]

    manuell_erforderlich = [
        'Organigramme in Kap. 2.2 und 3.3.x.2 (für jede Option) einfügen',
        'Screenshots der Marktrecherche-Webseiten in Anlage einfügen (Einfügen → Objekt/Bild)',
        'Dokumentkopf (Az., Projekttitel) in Tabelle 3 prüfen/ergänzen',
        'In Word: Inhalts-/Tabellenverzeichnis aktualisieren (Rechtsklick → Felder aktualisieren) nach Änderungen',
        'Finalisierungs-Hinweisseite (Seite 1) vor Abgabe entfernen',
    ]

    checkliste = (
        f'\n{"="*60}\n'
        f'ABSCHLUSS-CHECKLISTE WU-DOKUMENT\n'
        f'Datei: {output_path}\n'
        f'{"="*60}\n\n'
        f'Automatisch befüllt:\n'
    )
    for item in auto_befuellt:
        checkliste += f'  [x] {item}\n'

    checkliste += '\nNoch manuell zu ergänzen:\n'
    for item in manuell_erforderlich:
        checkliste += f'  [ ] {item}\n'

    checkliste += f'\n{"="*60}\n'
    return checkliste


# ── Dateiname nach Konvention erzeugen ────────────────────────────────────────

def build_filename(datum: str, sachverhalt: str, dienststelle: str, version: int = 1) -> str:
    """
    Erzeugt vollständigen Ausgabepfad nach Konvention:
    <OUTPUT_DIR>/JJJJMMTT_WU_[Sachverhalt]_[Dienststelle]_Version_N.docx
    datum: 'TT.MM.JJJJ'
    """
    teile = datum.split('.')
    jjmmtt = teile[2] + teile[1] + teile[0] if len(teile) == 3 else datum.replace('.', '')
    sachverhalt_clean  = sachverhalt.replace(' ', '_').replace('/', '_')
    dienststelle_clean = dienststelle.replace(' ', '_').replace('/', '_')
    filename = f"{jjmmtt}_WU_{sachverhalt_clean}_{dienststelle_clean}_Version_{version}.docx"
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    return os.path.join(OUTPUT_DIR, filename)


# ── WU_DATA_SCHEMA (Dokumentation der Datenstruktur) ─────────────────────────
WU_DATA_SCHEMA = {
    "meta": {
        "titel":        "Vollständiger Titel der WU",
        "kurztitel":    "WU Kurzbezeichnung",
        "dienststelle": "BAIUDBw ...",
        "bearbeiter":   "Vorname Nachname",
        "datum":        "TT.MM.JJJJ",
        "schutz":       "offen",
        "version":      "1",
    },
    "ueberblick": {
        "betrachtungsgegenstand": "...",
        "entscheidungsvorschlag": "...",
    },
    "kap1": {
        "bedarfsforderung":    "...",
        "bedarfsprognose":     "...",
        "rb_rechtlich":        "Entfällt / ...",
        "rb_organisatorisch":  "Entfällt / ...",
        "rb_zeitlich":         "Entfällt / ...",
        "rb_sonstige":         "...",
    },
    "kap2": {
        "ablauforganisation":       "...",
        "aufbauorganisation":       "...",
        "personal":                 "...",
        "material":                 "...",
        "infrastruktur":            "...",
        "sach_dienstleistungen":    "...",
        "einnahmen":                "Keine Einnahmen.",
        "haushalterische_darstellung": "...",
        "haushalterische_tabelle": [
            ["Personal", "2.3 Personal", "X.XXX €", "–"],
            ["Material", "2.4 Material", "–", "–"],
            ["Infrastruktur", "2.5 Infrastruktur", "–", "–"],
            ["Dienstleistungen", "2.6 ...", "X.XXX €", "–"],
            ["Gesamt", "", "XX.XXX €", "–"],
        ],
    },
    "kap3": {
        "optionen_grundsaetzlich": "Option 1: ... Option 2: ... Option 3: ...",
        "aussonderung":            "...",
        "optionen_detail": [
            {
                "titel":               "Option 1: Kauf",
                "ablauforganisation":  "...",
                "aufbauorganisation":  "...",
                "personal":            "...",
                "material":            "...",
                "infrastruktur":       "...",
                "sach_dienstleistungen": "...",
                "einnahmen":           "Keine Einnahmen.",
            },
        ],
    },
    "kap4": {
        "alle_optionen":      "Betriebliche Annahmen: ...\nKalkulatorische Annahmen: ...",
        "bestimmte_optionen": "Option 1: ... Option 2: ... Option 3: ...",
        "berechnung":         "Kalkulationszinssatz: 1,2 % ...",
    },
    "kap5": {
        "ibv":              "Entfällt.",
        "berechnung_text":  "Alle Barwerte bei 1,2 % ...",
        "berechnung_tabelle": {
            "headers": ["Kostenart", "Option 1", "Option 2", "Option 3"],
            "rows": [
                {"values": ["Investition", "25.000 €", "–", "–"]},
                {"values": ["Personal/Jahr", "10.800 €", "10.800 €", "3.600 €"]},
                {"values": ["", "", "", ""], "separator": True},
                {"values": ["Kapitalwert", "154.900 €", "175.700 €", "518.200 €"], "summe": True},
            ],
        },
        "kw_ohne_risiko":      "Option 1: X € | Option 2: X € | Option 3: X €",
        "risikoidentifizierung": "...",
        "risikoverteilung":    "...",
        "risikobewertung":     "...",
        "kw_mit_risiko":       "Option 1: X € | Option 2: X € | Option 3: X €",
    },
    "kap6_9": {
        "optionen_uebersicht": [
            {"name": "Option 1: Kauf", "kw_ohne_risiko": "154.900 €", "kw_mit_risiko": "156.700 €", "empfohlen": True},
            {"name": "Option 2: Leasing", "kw_ohne_risiko": "175.700 €", "kw_mit_risiko": "176.700 €", "empfohlen": False},
            {"name": "Option 3: Dienstleistung", "kw_ohne_risiko": "518.200 €", "kw_mit_risiko": "538.450 €", "empfohlen": False},
        ],
        "vergleich":             "...",
        "sensitivitaet":         "...",
        "nichtmonetaer":         "...",
        "entscheidungsvorschlag": "...",
    },
    "anlage": [
        {"nr": "1", "produkt": "...", "preis": "...", "url": "https://..."},
    ],
}
