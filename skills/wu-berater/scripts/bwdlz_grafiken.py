"""
WU-Berater: BwDLZ-Grafiken für Aufbau- und Ablauforganisation

Erzeugt auf Basis des Sachverhalts automatisch:
  - ein auf die relevanten Bereiche reduziertes Organigramm (Aufbauorganisation)
  - eine reduzierte Workflow-Grafik (Ablauforganisation)

Beide Grafiken werden als formatierte Tabellen direkt in ein python-docx-Dokument
eingefügt und nutzen die Farbgestaltung des WU-Templates.

Verwendung:
    from bwdlz_grafiken import ist_bwdlz, erzeuge_organigramm, erzeuge_workflow
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ── Farben (angelehnt an WU-Template) ─────────────────────────────────────────
C_DUNKEL   = '1F3864'   # Dunkelblau  – BwDLZ-Leitung
C_MITTEL   = '2E5496'   # Mittelblau  – direkt beteiligte Bereiche
C_HELL     = 'BDD7EE'   # Hellblau    – indirekt beteiligte Bereiche
C_GRAU     = 'D9D9D9'   # Grau        – nicht beteiligte Bereiche (ausgeblendet)
C_WEISS    = 'FFFFFF'
C_PFEIL    = 'F4F4F4'   # Hellgrau    – Pfeil-Zwischenzeilen
C_STEP_BG  = 'DEEAF1'   # Hellblau    – Workflow-Schritte


# ── Vollständige BwDLZ-Struktur (aus Organigramm BwDLZ.txt) ───────────────────
BWDLZ_STRUKTUR = {
    'BwDLZ-Leitung': {
        'untereinheiten': ['Stabsstellen'],
        'stichwoerter':   [],
        'pflicht':        True,   # immer im Organigramm
    },
    'Facility Management': {
        'untereinheiten': ['Liegenschaften', 'Gebäude- und Anlagenbetrieb', 'Instandhaltung / Störungsmanagement'],
        'stichwoerter':   ['reinigung', 'gebäude', 'liegenschaft', 'sportplatz', 'tartanbahn', 'anlage',
                           'instandhaltung', 'wartung', 'bau', 'fläche', 'grünanlage', 'sanitär',
                           'heizung', 'strom', 'wasser', 'entsorgung', 'abfall'],
        'pflicht':        False,
    },
    'Beschaffung / Materialwirtschaft': {
        'untereinheiten': ['Bedarfsermittlung', 'Beschaffung', 'Materialverwaltung'],
        'stichwoerter':   [],
        'pflicht':        True,   # immer im Organigramm (Kernprozess Vergabe)
    },
    'Finanzen / Interne Dienste': {
        'untereinheiten': ['Haushalt', 'Kasse / Rechnungswesen', 'interne Verwaltungsprozesse'],
        'stichwoerter':   [],
        'pflicht':        True,   # immer im Organigramm (Budget/Zahlung)
    },
    'Personalmanagement': {
        'untereinheiten': ['Personalbetreuung', 'Aus- und Fortbildung', 'Wohnungsfürsorge / soziale Angelegenheiten'],
        'stichwoerter':   ['personal', 'mitarbeiter', 'stellen', 'ausbildung', 'fortbildung', 'schulung'],
        'pflicht':        False,
    },
    'Verpflegung': {
        'untereinheiten': ['Truppenküchen', 'Verpflegungsplanung'],
        'stichwoerter':   ['verpflegung', 'küche', 'catering', 'essen', 'kantinen', 'truppenverpflegung'],
        'pflicht':        False,
    },
    'Standortservice / sonstige Services': {
        'untereinheiten': ['Empfang / Auskunft', 'Schlüssel- und Servicefragen', 'Standortunterstützung'],
        'stichwoerter':   ['service', 'empfang', 'standort', 'schlüssel', 'auskunft', 'transport',
                           'fuhrpark', 'poststelle'],
        'pflicht':        False,
    },
}

# ── Vollständiger BwDLZ-Workflow (aus BwDLZ Workflow.docx) ────────────────────
BWDLZ_WORKFLOW = [
    {'nr': '1',  'schritt': 'Bedarf entsteht',              'rolle': 'Bedarfsträger',                    'entscheidung': 'Ja: Bedarf beschreiben'},
    {'nr': '2',  'schritt': 'Bedarf beschreiben',           'rolle': 'Bedarfsträger',                    'entscheidung': 'Vollständig: Fachprüfung'},
    {'nr': '3',  'schritt': 'Fachliche Prüfung',            'rolle': 'Fachbereich',                      'entscheidung': 'Begründet: Budgetprüfung'},
    {'nr': '4',  'schritt': 'Budgetprüfung',                'rolle': 'Haushalts-/Finanzstelle',          'entscheidung': 'Budget OK: Beschaffungsweg'},
    {'nr': '5',  'schritt': 'Beschaffungsweg festlegen',    'rolle': 'Beschaffung',                      'entscheidung': 'Klar: Freigabe einholen'},
    {'nr': '6',  'schritt': 'Freigabe einholen',            'rolle': 'Freigabeberecht. Führungskraft',   'entscheidung': 'Genehmigt: Auftrag auslösen'},
    {'nr': '7',  'schritt': 'Bestellung/Auftrag auslösen',  'rolle': 'Beschaffung',                      'entscheidung': 'Korrekt: Lieferung abwarten'},
    {'nr': '8',  'schritt': 'Lieferung/Leistung erbracht',  'rolle': 'Lieferant/Auftragnehmer',          'entscheidung': '→ Abnahme'},
    {'nr': '9',  'schritt': 'Abnahme',                      'rolle': 'Abnehmende Stelle',                'entscheidung': 'OK: Rechnung prüfen'},
    {'nr': '10', 'schritt': 'Rechnung prüfen',              'rolle': 'Rechnungsstelle',                  'entscheidung': 'Korrekt: Zahlung'},
    {'nr': '11', 'schritt': 'Zahlung veranlassen',          'rolle': 'Finanzstelle',                     'entscheidung': 'OK: Dokumentation'},
    {'nr': '12', 'schritt': 'Dokumentation/Abschluss',      'rolle': 'Beschaffung/Rechnungsstelle',      'entscheidung': 'Vollständig: Archivieren'},
]

# Rollen → BwDLZ-Bereich-Mapping (für Workflow-Filterung)
ROLLE_BEREICH = {
    'Bedarfsträger':                    None,           # je nach Sachverhalt
    'Fachbereich':                      None,
    'Haushalts-/Finanzstelle':          'Finanzen / Interne Dienste',
    'Beschaffung':                      'Beschaffung / Materialwirtschaft',
    'Freigabeberecht. Führungskraft':   'BwDLZ-Leitung',
    'Lieferant/Auftragnehmer':          None,           # extern
    'Abnehmende Stelle':                None,
    'Rechnungsstelle':                  'Finanzen / Interne Dienste',
    'Finanzstelle':                     'Finanzen / Interne Dienste',
    'Beschaffung/Rechnungsstelle':      'Beschaffung / Materialwirtschaft',
}


# ── Hilfsfunktionen ────────────────────────────────────────────────────────────

def ist_bwdlz(dienststelle: str) -> bool:
    """Prüft ob die Dienststelle ein BwDLZ ist."""
    d = dienststelle.lower()
    return 'bwdlz' in d or 'bundeswehrdienstleistungszentrum' in d


def erkenne_relevante_bereiche(sachverhalt: str) -> list[str]:
    """
    Bestimmt anhand von Stichwörtern im Sachverhalt welche BwDLZ-Bereiche
    im Organigramm hervorgehoben werden.

    :param sachverhalt: Sachverhaltstext (Bezeichnung + Beschreibung)
    :return: Liste der relevanten Bereichsnamen
    """
    s = sachverhalt.lower()
    relevante = []
    for bereich, info in BWDLZ_STRUKTUR.items():
        if info['pflicht']:
            relevante.append(bereich)
        elif any(sw in s for sw in info['stichwoerter']):
            relevante.append(bereich)
    return relevante


def _set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold=False, size=9, color_hex=None,
                   align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )


def _merge_row(table, row_idx: int):
    """Verbindet alle Zellen einer Zeile."""
    row = table.rows[row_idx]
    cells = row.cells
    for i in range(1, len(cells)):
        cells[0].merge(cells[i])
    return cells[0]


# ── Organigramm ────────────────────────────────────────────────────────────────

def erzeuge_organigramm(doc: Document, sachverhalt: str, dienststelle: str = '') -> None:
    """
    Fügt ein auf relevante Bereiche reduziertes BwDLZ-Organigramm als
    formatierte Tabelle in das Dokument ein.

    :param doc:          python-docx Document-Objekt
    :param sachverhalt:  Sachverhaltstext für Relevanz-Erkennung
    :param dienststelle: Name der Dienststelle (für Überschrift)
    """
    relevante = erkenne_relevante_bereiche(sachverhalt)

    # Nur Fachbereiche (ohne Pflichtfelder die sowieso immer dabei sind)
    fachbereiche = [b for b in BWDLZ_STRUKTUR.keys()
                    if b != 'BwDLZ-Leitung']

    # Tabelle: 1 Kopfzeile + 1 Bereichszeile + 1 Unterzeile
    # Spaltenanzahl = Anzahl Fachbereiche
    n_cols = len(fachbereiche)
    tbl = doc.add_table(rows=4, cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Spaltenbreiten gleichmäßig verteilen (max. 16 cm)
    col_w = Cm(16.0 / n_cols)
    for col in tbl.columns:
        for cell in col.cells:
            cell.width = col_w

    # Zeile 0: BwDLZ-Leitung (gemergt, dunkelblau)
    leitung_cell = _merge_row(tbl, 0)
    _set_cell_shading(leitung_cell, C_DUNKEL)
    _set_cell_text(leitung_cell,
                   f'BwDLZ-Leitung — {dienststelle}' if dienststelle else 'BwDLZ-Leitung',
                   bold=True, size=10, color_hex=C_WEISS)

    # Zeile 1: Verbindungslinie (grau, klein)
    conn_cell = _merge_row(tbl, 1)
    _set_cell_shading(conn_cell, C_PFEIL)
    _set_cell_text(conn_cell, '│', size=7, color_hex='AAAAAA')

    # Zeile 2: Fachbereiche
    for col_idx, bereich in enumerate(fachbereiche):
        cell = tbl.rows[2].cells[col_idx]
        ist_relevant = bereich in relevante
        farbe = C_MITTEL if ist_relevant else C_GRAU
        text_farbe = C_WEISS if ist_relevant else '666666'
        _set_cell_shading(cell, farbe)
        _set_cell_text(cell, bereich, bold=ist_relevant, size=8, color_hex=text_farbe)

    # Zeile 3: Untereinheiten (nur bei relevanten Bereichen)
    for col_idx, bereich in enumerate(fachbereiche):
        cell = tbl.rows[3].cells[col_idx]
        ist_relevant = bereich in relevante
        if ist_relevant:
            untereinheiten = BWDLZ_STRUKTUR[bereich]['untereinheiten']
            text = '\n'.join(f'· {u}' for u in untereinheiten)
            _set_cell_shading(cell, C_HELL)
            _set_cell_text(cell, text, size=7, align=WD_ALIGN_PARAGRAPH.LEFT)
        else:
            _set_cell_shading(cell, 'F2F2F2')
            _set_cell_text(cell, '', size=7)

    # Hinweis unter Tabelle
    hinweis = doc.add_paragraph()
    hinweis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hinweis.add_run(
        f'Organigramm BwDLZ (auf relevante Bereiche reduziert) — '
        f'Beteiligte Bereiche: {", ".join(relevante)}'
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ── Workflow ───────────────────────────────────────────────────────────────────

def erzeuge_workflow(doc: Document, sachverhalt: str) -> None:
    """
    Fügt eine auf den Sachverhalt angepasste BwDLZ-Workflow-Tabelle in das
    Dokument ein.

    :param doc:         python-docx Document-Objekt
    :param sachverhalt: Sachverhaltstext für Schritt-Anpassung
    """
    s = sachverhalt.lower()

    # Schritt-Bezeichnungen sachverhaltsbezogen anpassen
    workflow = []
    for step in BWDLZ_WORKFLOW:
        angepasst = step.copy()
        # Schritt 1 konkretisieren
        if step['nr'] == '1':
            angepasst['schritt'] = f"Bedarf entsteht\n({sachverhalt[:60]}{'...' if len(sachverhalt)>60 else ''})"
        # Bedarfsträger konkretisieren
        if 'reinigung' in s or 'tartanbahn' in s or 'sportplatz' in s or 'gebäude' in s or 'liegenschaft' in s:
            if angepasst['rolle'] == 'Bedarfsträger':
                angepasst['rolle'] = 'Bedarfsträger\n(Facility Management)'
            if angepasst['rolle'] == 'Abnehmende Stelle':
                angepasst['rolle'] = 'Abnehmende Stelle\n(Facility Management)'
        workflow.append(angepasst)

    # Tabelle: Header + 1 Zeile je Schritt
    tbl = doc.add_table(rows=1 + len(workflow), cols=3)
    tbl.style = 'Table Grid'

    # Spaltenbreiten: Nr (1cm), Schritt/Rolle (9cm), Entscheidung (6cm)
    col_widths = [Cm(1.2), Cm(9.5), Cm(5.5)]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    # Header
    headers = ['Nr.', 'Prozessschritt / Zuständige Rolle', 'Entscheidung / Folgeschritt']
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _set_cell_shading(c, C_DUNKEL)
        _set_cell_text(c, h, bold=True, size=9, color_hex=C_WEISS)

    # Workflow-Schritte
    for row_idx, step in enumerate(workflow, start=1):
        row = tbl.rows[row_idx]
        fill = C_STEP_BG if row_idx % 2 == 0 else C_WEISS

        _set_cell_shading(row.cells[0], fill)
        _set_cell_text(row.cells[0], step['nr'], size=9, bold=True)

        _set_cell_shading(row.cells[1], fill)
        _set_cell_text(row.cells[1],
                       f"{step['schritt']}\n→ {step['rolle']}",
                       size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

        _set_cell_shading(row.cells[2], fill)
        _set_cell_text(row.cells[2], step['entscheidung'], size=9,
                       align=WD_ALIGN_PARAGRAPH.LEFT)

    # Hinweis
    hinweis = doc.add_paragraph()
    run = hinweis.add_run(
        'Workflow BwDLZ (Beschaffungs-/Vergabeprozess, auf Sachverhalt angepasst)'
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ── Generische Fallbacks (nicht-BwDLZ) ────────────────────────────────────────

def erzeuge_workflow_generisch(doc: Document, text: str) -> None:
    """
    Generischer Workflow für nicht-BwDLZ-Einheiten.
    Zerlegt den Ablauftext in Schritte und stellt sie als Workflow-Tabelle dar.
    """
    # Schritte aus Text extrahieren (Zeilenumbrüche oder Sätze)
    import re
    zeilen = [z.strip() for z in re.split(r'\n|(?<=\.)\s+', text) if z.strip()]
    if not zeilen:
        zeilen = [text] if text.strip() else ['Ablauforganisation (wird manuell ergänzt)']

    tbl = doc.add_table(rows=1 + len(zeilen), cols=2)
    tbl.style = 'Table Grid'

    col_widths = [Cm(1.5), Cm(14.0)]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    # Header
    _set_cell_shading(tbl.rows[0].cells[0], C_DUNKEL)
    _set_cell_text(tbl.rows[0].cells[0], 'Nr.', bold=True, size=9, color_hex=C_WEISS)
    _set_cell_shading(tbl.rows[0].cells[1], C_DUNKEL)
    _set_cell_text(tbl.rows[0].cells[1], 'Prozessschritt', bold=True, size=9, color_hex=C_WEISS)

    for idx, zeile in enumerate(zeilen, start=1):
        fill = C_STEP_BG if idx % 2 == 0 else C_WEISS
        row = tbl.rows[idx]
        _set_cell_shading(row.cells[0], fill)
        _set_cell_text(row.cells[0], str(idx), size=9, bold=True)
        _set_cell_shading(row.cells[1], fill)
        _set_cell_text(row.cells[1], zeile, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

    hinweis = doc.add_paragraph()
    run = hinweis.add_run('Ablauforganisation (Prozessschritte)')
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def erzeuge_organigramm_generisch(doc: Document, text: str, dienststelle: str = '') -> None:
    """
    Generisches Organigramm für nicht-BwDLZ-Einheiten.
    Zeigt Dienststelle als Kopf und die Aufbauorganisationseinheiten darunter.
    """
    import re
    zeilen = [z.strip() for z in re.split(r'\n|;', text) if z.strip()]
    if not zeilen:
        zeilen = [text] if text.strip() else ['(Aufbauorganisation wird manuell ergänzt)']

    n_cols = min(len(zeilen), 4)
    tbl = doc.add_table(rows=3, cols=n_cols)
    tbl.style = 'Table Grid'
    from docx.enum.table import WD_TABLE_ALIGNMENT
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_w = Cm(16.0 / n_cols)
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = col_w

    # Zeile 0: Dienststelle (gemergt)
    top_cell = _merge_row(tbl, 0)
    _set_cell_shading(top_cell, C_DUNKEL)
    _set_cell_text(top_cell, dienststelle or 'Dienststelle', bold=True, size=10, color_hex=C_WEISS)

    # Zeile 1: Verbindungslinie
    conn_cell = _merge_row(tbl, 1)
    _set_cell_shading(conn_cell, C_PFEIL)
    _set_cell_text(conn_cell, '│', size=7, color_hex='AAAAAA')

    # Zeile 2: Einheiten
    for col_idx in range(n_cols):
        cell = tbl.rows[2].cells[col_idx]
        text_val = zeilen[col_idx] if col_idx < len(zeilen) else ''
        _set_cell_shading(cell, C_MITTEL)
        _set_cell_text(cell, text_val, bold=bool(text_val), size=8, color_hex=C_WEISS)

    hinweis = doc.add_paragraph()
    run = hinweis.add_run('Aufbauorganisation (auf beteiligte Einheiten reduziert)')
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
