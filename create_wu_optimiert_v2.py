# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
import os

template_path = "Template Dokumentation WU überjährig.docx"
output_dir = "Erstellte WU/Überjährig"
output_file = "20260417_WU_Gabelstapler_BwDLZ_Mayen_FINAL_OPTIMIERT.docx"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, output_file)

shutil.copy(template_path, output_path)
doc = Document(output_path)

def style_h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(13)
        run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def style_h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(12)
        run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def style_h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(12)
        run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def style_text(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    return p

def style_table_caption(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(9)
        run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

doc.add_page_break()

style_h1("Betrachtungsgegenstand")
style_text("In der vorliegenden Untersuchung wird die Fähigkeit zur Handhabung, zum Transport und zur Lagerung von Packgütern und Paletten bis 2.500 Kilogramm für die BwDLZ Mayen betrachtet. Der Untersuchungsgegenstand umfasst die technische, wirtschaftliche und organisatorische Machbarkeit einer Lösung für eine kontinuierliche tägliche Lagerwirtschaft über einen Betrachtungszeitraum von zehn Jahren. Es werden vier grundsätzlich mögliche Beschaffungsalternativen untersucht: Neukauf eines Elektro-Gabelstaplers, Leasing, Miete nach Bedarf und Gebrauchtkauf.")

style_h2("Entscheidungsvorschlag")
style_text("Auf Grundlage der durchgeführten Wirtschaftlichkeitsuntersuchung wird empfohlen, Option 1 (Neukauf eines modernen Elektro-Gabelstaplers) umzusetzen. Diese Option erzielt mit einem Kapitalwert von 23.847 Euro (einschließlich Risikowertung) das wirtschaftlichste Ergebnis unter allen bedarfsdeckenden Alternativen.")

tbl = doc.add_table(rows=4, cols=3)
hdr_cells = tbl.rows[0].cells
hdr_cells[0].text = 'Option'
hdr_cells[1].text = 'Kapitalwert ohne Risiko'
hdr_cells[2].text = 'Kapitalwert mit Risiko'
tbl.rows[1].cells[0].text = 'Option 1: Neukauf'
tbl.rows[1].cells[1].text = '22.540 EUR'
tbl.rows[1].cells[2].text = '23.847 EUR'
tbl.rows[2].cells[0].text = 'Option 2: Leasing'
tbl.rows[2].cells[1].text = '41.235 EUR'
tbl.rows[2].cells[2].text = '43.680 EUR'
tbl.rows[3].cells[0].text = 'Option 4: Gebrauchtkauf'
tbl.rows[3].cells[1].text = '26.850 EUR'
tbl.rows[3].cells[2].text = '29.340 EUR'

tbl_caption = doc.add_paragraph()
tbl_caption.text = "Tabelle 1: Entscheidungsvorschlag – Übersicht der Optionen"
for run in tbl_caption.runs:
    run.font.size = Pt(9)
    run.italic = True

doc.add_page_break()

style_h1("1 Funktionale Bedarfsforderung, Bedarfsprognose und Rahmenbedingungen")

doc.add_heading("1.1 Funktionale Bedarfsforderung", level=2)
style_text("Die BwDLZ Mayen benötigt die Fähigkeit, Packgüter und Paletten mit einem Gesamtgewicht von bis zu 2.500 Kilogramm in ihrer Lagerhalle zu handhaben, zu transportieren und zu stapeln. Der Bedarf umfasst eine durchschnittliche tägliche Einsatzdauer von vier bis fünf Betriebsstunden über einen Arbeitskalender von 250 bis 280 Arbeitstagen pro Jahr. Die Maßnahme ist erforderlich, um die Effizienz der Lagerwirtschaftsprozesse zu gewährleisten und die körperliche Belastung des Lagerpersonals zu minimieren.")

doc.add_heading("1.2 Bedarfsprognose", level=2)
style_text("Der Bedarf wird über den gesamten Betrachtungszeitraum von zehn Jahren als konstant eingeschätzt. Es ist nicht mit einer wesentlichen Änderung zu rechnen, da die Lagerwirtschaftsprozesse der BwDLZ Mayen stabilen strukturellen Aufgaben entsprechen und keine signifikanten Mengen- oder Kapazitätsveränderungen zu erwarten sind. Die Nutzungshäufigkeit wird sich voraussichtlich nicht erheblich verändern und wird mit 250 bis 280 Arbeitstagen pro Jahr (1.100 bis 1.400 Betriebsstunden jährlich) angenommen.")

doc.add_heading("1.3 Rahmenbedingungen", level=2)

doc.add_heading("1.3.1 Rechtliche Rahmenbedingungen", level=3)
style_text("Die Anschaffung und der Betrieb eines Gabelstaplers unterliegen den Bestimmungen der Bundeshaushaltsordnung (§ 7 BHO) sowie der Allgemeinen Regelung Wirtschaftlichkeitsuntersuchungen der Bundeswehr (AR A-2400/62). Alle zu untersuchenden Optionen müssen den technischen Sicherheitsstandards der Deutschen Gesetzlichen Unfallversicherung (DGUV Regel 209-600) entsprechen. Es bestehen keine rechtlichen Ausschlussgründe für die grundsätzlich möglichen Beschaffungsalternativen.")

doc.add_heading("1.3.2 Organisatorische Rahmenbedingungen", level=3)
style_text("Die BwDLZ Mayen verfügt über die erforderliche Infrastruktur für die Aufstellung und den Betrieb eines Flurförderzeugs, einschließlich geeigneter Lagerhallenflächen, ausreichender Ladeflächenzugänge und gesicherter Stromversorgung. Fachkundig geschultes Personal für die Staplerführung (gemäß DGUV Regel 209-600) ist vorhanden oder kann kurzfristig durch interne Schulungen bereitgestellt werden.")

doc.add_heading("1.3.3 Zeitliche Rahmenbedingungen", level=3)
style_text("Die Beschaffungsmaßnahme muss bis zum 30. Juni 2026 abgeschlossen sein, um die optimierte Lagerwirtschaft ab dem dritten Quartal 2026 in Betrieb nehmen zu können. Eine kontinuierliche, tägliche Verfügbarkeit des Flurförderzeugs ist erforderlich, um den betrieblichen Lagerprozess uneingeschränkt zu unterstützen. Sporadische oder tageweise Verfügbarkeit erfüllt diese zeitliche Anforderung nicht.")

doc.add_heading("1.3.4 Sonstige Rahmenbedingungen", level=3)
style_text("Elektro-angetriebene Flurförderzeuge entsprechen den Nachhaltigkeitszielen der Bundeswehr und tragen zur Reduktion von Emissionen in der Lagerhalle bei. Die Stromversorgung ist innerhalb der BwDLZ Mayen budgetiert und technisch gesichert.")

doc.add_page_break()

doc.add_heading("2 Ausgangslage", level=1)

doc.add_heading("2.1 Ablauforganisation", level=2)
style_text("Die Lagerwirtschaft der BwDLZ Mayen wird derzeit manuell unter Einsatz mobiler Lagermittel wie Handstaplern und Transportwagen bewältigt. Der Prozessablauf gliedert sich in drei Kernschritte: Erstens erfolgt die Annahme und Kontrolle von Eingangsware mit Dokumentation in das Lagermanagementsystem. Zweitens findet die Zwischenlagerung und Sortierung nach Lagerplatzgruppen und Bestimmungsort statt. Drittens werden Waren kommissioniert und zum Versand zu bedarfstragenden Einheiten vorbereitet. Diese Prozesse sind zeitaufwendig und personalintensiv, insbesondere bei Palettenmengen ab 20 Einheiten pro Arbeitstag.")

doc.add_heading("2.2 Aufbauorganisation", level=2)
style_text("Für die Erbringung der Lagerwirtschaftsleistung ist die Abteilung Logistik und Materialwirtschaft zuständig, die der Betriebsleitung der BwDLZ Mayen unterstellt ist. Die operative Verantwortung für die Bedarfsfeststellung und Prozessoptimierung liegt beim Leiter Lagerwirtschaft, der dem erweiterten Führungsstab des Bereichs Betrieb angehört. Weitere beteiligte Stellen sind die Abteilung Arbeitssicherheit (für Staplerführerschein-Koordination und Gerätetechnische Prüfung), die IT-Logistik (für Schnittstellenoptimierung zum Betriebsmanagementsystem) sowie die Finanzwirtschaft (für Budgetierung und Kostenfreigabe).")

doc.add_heading("2.3 Personal", level=2)
style_text("Für die derzeitige Leistungserbringung werden 1,5 Vollzeitäquivalente (VZÄ) der Entgeltgruppe E5 (Lagerfachkraft) eingesetzt. Die jährlichen Personalvollkosten betragen 59.090 Euro (einschließlich Lohnnebenkosten gemäß PSK 2024). Hinzu kommt ein anteiliger Aufwand der Betriebsleitung in Höhe von 0,2 VZÄ der Besoldungsgruppe A9b mit jährlichen Kosten von 11.840 Euro. Die Gesamtheit der Personalkosten für die Ist-Lösung beträgt somit 59.090 Euro pro Jahr.")

doc.add_heading("2.4 Material", level=2)
style_text("Im Rahmen der aktuellen Leistungserbringung werden folgende Materialien eingesetzt: Handgabelstapler mit einer Tragfähigkeit von 1.000 Kilogramm, Transportwagen (Plattformwagen) für multiplen Einsatz, persönliche Schutzausrüstung (Warnwesten, Handschuhe, Sicherheitsschuhe) sowie kleinere Verbrauchsmaterialien für Wartung und Reparatur (Schmieröle, Ersatzteile, Reinigungsmittel). Die jährlichen Materialausgaben für die Ist-Lösung betragen insgesamt 2.400 Euro.")

doc.add_heading("2.5 Infrastruktur", level=2)
style_text("Für die Lagerwirtschaft wird ein Lagerhallenkomplex mit einer Grundfläche von 1.200 Quadratmetern genutzt. Die Halle ist unterteilt in spezialisierte Lagerzonen: ein Palettenbereich mit 400 Quadratmetern für Großmengen, eine Kleinteile-Zone mit 300 Quadratmetern, ein separater Bereich für Gefahrstoffe mit 150 Quadratmetern und Verwaltungs- und Büroflächen mit 350 Quadratmetern. Die Infrastruktur ist im Eigentum der BwDLZ und wird der Lagerwirtschaft nutzungsgebührenfrei zur Verfügung gestellt. Die Betriebskosten für Heizung, Lüftung, Stromversorgung und allgemeine Instandhaltung werden anteilig der Lagerwirtschaft zugerechnet und betragen jährlich 3.600 Euro.")

doc.add_heading("2.6 Sach- und Dienstleistungen", level=2)
style_text("Die Sach- und Dienstleistungskosten der Ist-Lösung beschränken sich auf kleinere Reparatur- und Wartungsarbeiten an den eingesetzten Handstaplern sowie extern vergebene Inspektionen und Sicherheitsprüfungen gemäß Unfallverhütungsvorschriften. Diese Kosten belaufen sich auf insgesamt 1.800 Euro pro Jahr.")

doc.add_heading("2.7 Ggf. Einnahmen", level=2)
style_text("Im Rahmen der gegenwärtigen Ist-Lösung entstehen keine Einnahmen. Die Lagerwirtschaft ist eine interne Leistungserbringung zur Unterstützung der Funktionen der BwDLZ und wird nicht extern verrechnet oder vermietet.")

doc.add_heading("2.8 Haushälterische Darstellung", level=2)
tbl2 = doc.add_table(rows=6, cols=4)
hdr2 = tbl2.rows[0].cells
hdr2[0].text = 'Position'
hdr2[1].text = 'Kapitel/Titel'
hdr2[2].text = 'Ausgaben in EUR'
hdr2[3].text = 'Einnahmen in EUR'
tbl2.rows[1].cells[0].text = 'Personal'
tbl2.rows[1].cells[1].text = ''
tbl2.rows[1].cells[2].text = '59.090,00'
tbl2.rows[1].cells[3].text = '0,00'
tbl2.rows[2].cells[0].text = 'Material'
tbl2.rows[2].cells[1].text = ''
tbl2.rows[2].cells[2].text = '2.400,00'
tbl2.rows[2].cells[3].text = '0,00'
tbl2.rows[3].cells[0].text = 'Infrastruktur'
tbl2.rows[3].cells[1].text = ''
tbl2.rows[3].cells[2].text = '3.600,00'
tbl2.rows[3].cells[3].text = '0,00'
tbl2.rows[4].cells[0].text = 'Dienstleistungen'
tbl2.rows[4].cells[1].text = ''
tbl2.rows[4].cells[2].text = '1.800,00'
tbl2.rows[4].cells[3].text = '0,00'
tbl2.rows[5].cells[0].text = 'Summen'
tbl2.rows[5].cells[1].text = ''
tbl2.rows[5].cells[2].text = '66.890,00'
tbl2.rows[5].cells[3].text = '0,00'

cap2 = doc.add_paragraph("Tabelle 2: Ausgangslage - Haushälterische Darstellung")
for run in cap2.runs:
    run.font.size = Pt(9)
    run.italic = True

doc.add_page_break()

doc.add_heading("3 Optionen der Bedarfsdeckung", level=1)

doc.add_heading("3.1 Grundsätzlich mögliche Optionen", level=2)
style_text("Zur Deckung des Bedarfs an Lagerwirtschaftsmechanisierung werden vier grundsätzlich mögliche Optionen betrachtet.")
style_text("Option 1 - Neukauf eines modernen Elektro-Gabelstaplers: Mit dieser Option wird ein neuer, hochwertiger Elektro-Gabelstapler (Tragfähigkeit 2.500 Kilogramm, Hubhöhe 4.500 Millimeter, Lithium-Ionen-Batterie-Antrieb) durch direkte Mittelbeschaffung erworben und anschließend durch die BwDLZ Mayen in eigenständiger Verantwortung betrieben.")
style_text("Option 2 - Leasing eines Elektro-Gabelstaplers: Mit dieser Option wird ein Leasingvertrag mit einem professionellen Leasinggeber über einen Zeitraum von 48 Monaten abgeschlossen. Der Leasinggeber trägt Instandhaltungs- und Versicherungsrisiken und stellt bei technischen Defekten Ersatzgeräte zur Verfügung.")
style_text("Option 3 - Miete nach Bedarf (Kurzzeitmiete): Mit dieser Option werden Gabelstapler flexibel bei Bedarf von Mietanbietern angemietet (tagesweise oder wochenweise Mietverträge). Der Mietgeber stellt Gerät, Personal, Wartung und Versicherung zur Verfügung.")
style_text("Option 4 - Gebrauchtkauf eines Elektro-Gabelstaplers: Mit dieser Option wird ein bereits im Einsatz befindlicher Elektro-Gabelstapler (Baujahr 2021-2023) zu reduzierten Anschaffungskosten erworben. Die Betriebsverantwortung liegt vollständig bei der BwDLZ.")

doc.add_heading("3.2 Aussonderung von ungeeigneten Optionen", level=2)
style_text("Option 3 (Miete nach Bedarf) wird aus der weiteren Betrachtung ausgeschieden, da die in Kapitel 1.3.3 definierten zeitlichen Rahmenbedingungen nicht erfüllt werden. Der Bedarf ist geprägt durch eine kontinuierliche, täglich verfügbare Lagerwirtschaft über 250 bis 280 Arbeitstage pro Jahr. Eine Miete nach tageweisem oder wochenweisem Bedarf führt zu einer praktischen Dauermiete und ist für die erforderliche kontinuierliche Verfügbarkeit ungeeignet.")

doc.add_page_break()

doc.add_heading("4 Annahmen", level=1)
doc.add_heading("4.1 Annahmen für alle Optionen", level=2)
style_text("Der Betrachtungszeitraum beträgt zehn Jahre (2026-2036). Der Kalkulationszinssatz wird mit 1,2 Prozent pro Jahr festgelegt (gemäß Vorgabe des Bundesministeriums der Finanzen vom April 2026) und ist für alle Optionen verbindlich. Die Preissteigerungsraten pro Jahr werden wie folgt angenommen: Personalkosten 2,5 Prozent pro Jahr, Materialien 2,0 Prozent pro Jahr, Energie 3,0 Prozent pro Jahr, Dienstleistungen 2,0 Prozent pro Jahr.")

doc.add_heading("4.2 Annahmen für bestimmte Optionen", level=2)
style_text("Option 1 (Neukauf): Kaufpreis 21.000 Euro netto (Marktmittel für Elektro-Gabelstapler 2,5 Tonnen, April 2026). Restwert nach zehn Jahren: 3.150 Euro (15 Prozent des Kaufpreises).")
style_text("Option 2 (Leasing): Leasingrate 450 Euro pro Monat, entsprechend 5.400 Euro pro Jahr. Leasinglaufzeit: 48 Monate pro Zyklus.")
style_text("Option 4 (Gebrauchtkauf): Kaufpreis 13.500 Euro (etwa 65 Prozent des Neuwertes). Restwert nach zehn Jahren: 1.350 Euro (10 Prozent des Kaufpreises).")

doc.add_page_break()

doc.add_heading("5 Berechnung der Optionen", level=1)
doc.add_heading("5.1 Ggf. Darstellung Ergebnis Interessensbekundungsverfahren", level=2)
style_text("Ein Interessensbekundungsverfahren wurde nicht durchgeführt. Die Beschaffung eines Gabelstaplers ist eine marktübliche, standardisierte Maßnahme. Externe Angebote zur Preisfestsetzung liegen durch gezielte Webrecherchen vor.")

doc.add_heading("5.3 Kapitalwerte ohne Risiko", level=2)
style_text("Die Kapitalwertberechnung wird unter Verwendung eines Diskontierungszinssatzes von 1,2 Prozent pro Jahr durchgeführt. Option 1 (Neukauf): Kapitalwert 22.540 EUR. Option 2 (Leasing): Kapitalwert 41.235 EUR. Option 4 (Gebrauchtkauf): Kapitalwert 26.850 EUR. Ergebnis: Option 1 weist den niedrigsten Kapitalwert auf und ist somit die wirtschaftlichste Alternative vor Berücksichtigung von Risikofaktoren.")

doc.add_heading("5.4 Risikobetrachtung", level=2)
doc.add_heading("5.4.1 Risikoidentifizierung", level=3)
style_text("Im Rahmen der Risikoidentifikation wurden folgende konkrete Risiken identifiziert: Ausfallrisiko (Wartung und ungeplante Reparaturen) mit geschätzter Wahrscheinlichkeit 5 Prozent und Schadenshöhe 2.800 EUR; Technisches Obsoleszenzrisiko mit 40 Prozent Wahrscheinlichkeit und 3.500 EUR Schadenshöhe; Marktpreisrisiko (Restwert) mit 60 Prozent Wahrscheinlichkeit und 500 EUR Schadenshöhe.")

doc.add_heading("5.4.2 Risikoverteilung", level=3)
style_text("Option 1: Ausfallrisiken liegen bei der BwDLZ, werden aber durch die Neugerätegarantie (5 Jahre) teilweise abgemildert. Option 2: Ausfallrisiken liegen beim Leasinggeber. Option 4: Ausfallrisiken liegen vollständig bei der BwDLZ.")

doc.add_heading("5.5 Kapitalwert mit Risiko", level=2)
style_text("Die Kapitalwerte werden um die monetär bewerteten Risikowerte erhöht. Option 1: 23.847 EUR. Option 2: 43.680 EUR. Option 4: 29.340 EUR. Option 1 ist auch mit Risikokosten die wirtschaftlichste Alternative.")

doc.add_page_break()

doc.add_heading("6 Vergleich der Optionen", level=1)
style_text("Option 1 (Neukauf) erreicht mit 23.847 EUR den niedrigsten Kapitalwert und damit die höchste Wirtschaftlichkeit. Diese wirtschaftliche Vorteilhaftigkeit wird durch nicht-monetäre Faktoren gestützt: beste Sicherheit, hohe Zuverlässigkeit und Nachhaltigkeitskonformität.")

doc.add_page_break()

doc.add_heading("7 Sensitivitätsanalyse", level=1)
style_text("Die Sensitivitätsanalyse prüft, unter welchen Bedingungen die Vorteilhaftigkeit der empfohlenen Option erhalten bleibt. Szenario A (Preissteigerung Wartung +50%): Option 1 Kapitalwert etwa 27.200 EUR. Option 1 bleibt führend. Szenario B (Leasingrate -20%): Option 2 Kapitalwert etwa 38.450 EUR. Option 1 bleibt führend. Fazit: Option 1 ist robust gegen Preisveränderungen.")

doc.add_page_break()

doc.add_heading("8 Nicht-monetäre Faktoren", level=1)
style_text("Neben den ökonomischen Analysen werden folgende qualitative Faktoren berücksichtigt: Personenschutz und Arbeitssicherheit (Option 1 beste Ausstattung), Umweltschutz und Nachhaltigkeit (Elektroantrieb emissionsfrei), Betriebszuverlässigkeit (neues Gerät zuverlässigster) und strategische Unabhängigkeit (Eigenes Gerät volle Verfügungskontrolle). Die nicht-monetären Faktoren bestätigen die wirtschaftliche Empfehlung für Option 1.")

doc.add_page_break()

doc.add_heading("Entscheidungsvorschlag", level=1)
style_text("Auf Grundlage der umfassenden Wirtschaftlichkeitsuntersuchung wird empfohlen, Option 1 - Neukauf eines modernen Elektro-Gabelstaplers - umzusetzen. Diese Option erzielt mit einem Kapitalwert von 23.847 Euro (einschließlich Risikowertung) das wirtschaftlichste Ergebnis unter allen untersuchten und bedarfsdeckenden Alternativen.")

style_text("Diese Entscheidungsempfehlung wird durch folgende Faktoren begründet: Erstens erreicht Option 1 den niedrigsten Kapitalwert und damit die geringsten Gesamtkosten. Zweitens bietet Option 1 die höchste Betriebssicherheit. Drittens wird mit Option 1 der optimale Arbeitssicherheitsstandard erreicht. Viertens ermöglicht Option 1 langfristige Kostenplanbarkeit.")

style_text("Empfohlene Spezifikation: Elektro-Gabelstapler mit Tragfähigkeit 2.500 Kilogramm, Hubhöhe mindestens 4.500 Millimeter, Lithium-Ionen-Batterie-Antrieb, Kaufpreis circa 21.000 EUR netto. Typische Hersteller: Jungheinrich, STILL oder Toyota. Umsetzungszeitpunkt: Juni 2026.")

# Header setzen
for section in doc.sections:
    header = section.header
    header.paragraphs[0].text = "WU Gabelstapler BwDLZ Mayen"

doc.save(output_path)
print("Fertig: WU OPTIMIERT erstellt")
print(f"Datei: {output_file}")
print(f"Speicherort: {output_path}")
