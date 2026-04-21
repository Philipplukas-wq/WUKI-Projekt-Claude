#!/usr/bin/env python3
"""Generate WU Word Document for Tartanbahnreinigung"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

doc = Document()

# Header
header = doc.add_paragraph()
header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
run = header.add_run("offen")
run.font.bold = True

# Title
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run("Wirtschaftlichkeitsuntersuchung")
run.font.bold = True
run.font.size = Pt(14)

# Metadata
doc.add_paragraph(f"Dienststelle: BwDLZ Mayen")
doc.add_paragraph(f"Sachverhalt: Tartanbahnreinigung")
doc.add_paragraph(f"Bearbeiter: Anna Katharina Probst")
doc.add_paragraph(f"Datum: {datetime.now().strftime('%d.%m.%Y')}")
doc.add_paragraph()

# Betrachtungsgegenstand
h = doc.add_heading("Betrachtungsgegenstand", level=2)
doc.add_paragraph("Der BwDLZ Mayen benoetigt an insgesamt 12 Liegenschaften im Zustaendigkeitsbereich die fachgerechte Reinigung und Instandhaltung von Tartanflaechen. Die Gesamtflaeche betraegt ca. 81.672 m2 und unterliegt Reinigungszyklen alle 2-3 Jahre.")

doc.add_page_break()

# Kapitel 1
doc.add_heading("1 Funktionale Bedarfsforderung", level=1)
doc.add_paragraph("Bedarfsforderung und Details.")

doc.add_page_break()

# Kapitel 6
doc.add_heading("6 Vergleich", level=1)
table = doc.add_table(rows=2, cols=4)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Kriterium'
table.rows[0].cells[1].text = 'Option 1'
table.rows[0].cells[2].text = 'Option 3'
table.rows[0].cells[3].text = 'Option 4'
table.rows[1].cells[0].text = 'Kapitalwert'
table.rows[1].cells[1].text = '805.441 EUR'
table.rows[1].cells[2].text = '589.714 EUR'
table.rows[1].cells[3].text = '3.381.614 EUR'

doc.add_page_break()

# Kapitel 9
doc.add_heading("9 Entscheidungsvorschlag", level=1)
doc.add_paragraph("Empfohlene Option: Option 3 (Fahrzeugmiete + Bundeswehr-Personal)")
doc.add_paragraph("Option 3 ist die wirtschaftlichste Loesung mit 589.714 EUR Kapitalwert.")

# Save
output_dir = r"P:\WUKI_Projekt\Claude\Erstellte WU\Dienstleistung"
os.makedirs(output_dir, exist_ok=True)
output_file = os.path.join(output_dir, "20260420_WU_Tartanbahnreinigung_BwDLZ_Mayen_Version_1.docx")

doc.save(output_file)
print(f"OK: {output_file}")
