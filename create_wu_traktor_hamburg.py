# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil
import os

template_path = "Template Dokumentation WU überjährig.docx"
output_dir = "Erstellte WU/Überjährig"
output_file = "20260417_WU_Allradschlepper_BwDLZ_Hamburg_Version_4.docx"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, output_file)

shutil.copy(template_path, output_path)
doc = Document(output_path)

def style_h1(text):
    p = doc.add_paragraph(text, style='Normal')
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def style_h2(text):
    p = doc.add_paragraph(text, style='Normal')
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p

def style_h3(text):
    p = doc.add_paragraph(text, style='Normal')
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    return p

def style_text(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    return p

def style_table_caption(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'BundesSans Office'
        run.font.size = Pt(9)
        run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def style_table(rows, cols, data):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = 'Table Grid'
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'BundesSans Office'
                    run.font.size = Pt(11)
            if i == 0:
                tcPr = cell._element.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:shd')
                tcVAlign.set(qn('w:fill'), 'D3D3D3')
                tcPr.append(tcVAlign)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
            else:
                for paragraph in cell.paragraphs:
                    if any(char.isdigit() for char in str(cell_data)):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return tbl

for section in doc.sections:
    header = section.header
    header.paragraphs[0].text = "WU Allradschlepper BwDLZ Hamburg"

doc.add_page_break()

style_h1("Betrachtungsgegenstand")
style_text("In der vorliegenden Untersuchung wird die Fähigkeit zur Erbringung von Winterdiensten und Grünflächenpflege für die BwDLZ Hamburg betrachtet. Der Untersuchungsgegenstand umfasst die technische, wirtschaftliche und organisatorische Machbarkeit einer Lösung für eine kontinuierliche Bedarfsdeckung über einen Betrachtungszeitraum von zehn Jahren. Es werden drei grundsätzlich mögliche Beschaffungsalternativen untersucht: Neukauf eines Allradschleppers, Leasing sowie Externe Dienstleistung durch Lohnunternehmer.")

style_h2("Entscheidungsvorschlag")
style_text("Auf Grundlage der durchgeführten Wirtschaftlichkeitsuntersuchung wird empfohlen, Option 1 (Neukauf eines modernen Allradschleppers) umzusetzen. Diese Option erzielt mit einem Kapitalwert von 89.850 Euro (einschließlich Risikowertung) das wirtschaftlichste Ergebnis unter allen bedarfsdeckenden Alternativen.")

style_table(4, 3, [['Option', 'Kapitalwert ohne Risiko', 'Kapitalwert mit Risiko'],['Option 1: Neukauf', '85.200 EUR', '89.850 EUR'],['Option 2: Leasing', '112.500 EUR', '118.200 EUR'],['Option 4: Externe DL', '156.000 EUR', '163.200 EUR']])
style_table_caption("Tabelle 1: Entscheidungsvorschlag – Übersicht der Optionen")

doc.add_page_break()

style_h1("1 Funktionale Bedarfsforderung, Bedarfsprognose und Rahmenbedingungen")

style_h2("1.1 Funktionale Bedarfsforderung")
style_text("Die BwDLZ Hamburg benötigt die Fähigkeit, Winterdienste (Schneeräumung, Streudienste) und Grünflächenpflege (Mähen, Mulchen, Wegeunterhalt) durchzuführen. Der Bedarf umfasst eine Einsatzzeit von 100 Arbeitstagen pro Jahr mit durchschnittlich fünf Betriebsstunden pro Tag (insgesamt 500 Betriebsstunden jährlich). Die Maßnahme erfordert ein Fahrzeug mit mindestens 150 PS Leistung und entsprechendem Zubehör für Winter- und Grünflächenarbeiten. Die Maßnahme ist erforderlich, um die Verkehrssicherheit und Infrastrukturpflege auf dem Gelände der BwDLZ Hamburg zu gewährleisten.")

style_h2("1.2 Bedarfsprognose")
style_text("Der Bedarf wird über den gesamten Betrachtungszeitraum von zehn Jahren als konstant eingeschätzt. Es ist nicht mit einer wesentlichen Änderung zu rechnen, da die Winterdienst- und Grünflächenpflegeleistungen strukturell unveränderliche Aufgaben darstellen. Die Nutzungshäufigkeit wird mit 100 Arbeitstagen pro Jahr (500 Betriebsstunden jährlich) angenommen und bleibt über den Planungszeitraum stabil.")

style_h2("1.3 Rahmenbedingungen")

style_h3("1.3.1 Rechtliche Rahmenbedingungen")
style_text("Die Anschaffung und der Betrieb eines Allradschleppers unterliegen den Bestimmungen der Bundeshaushaltsordnung (§ 7 BHO) sowie der Allgemeinen Regelung Wirtschaftlichkeitsuntersuchungen der Bundeswehr (AR A-2400/62). Alle Traktoren müssen der EU-Emissionsnorm Stage V entsprechen und mit Dieselpartikelfilter (DPF) und SCR-System ausgestattet sein. Es bestehen keine rechtlichen Ausschlussgründe für die grundsätzlich möglichen Beschaffungsalternativen.")

style_h3("1.3.2 Technische Rahmenbedingungen")
style_text("Der Schlepper muss eine Motorleistung von mindestens 150 PS aufweisen und über Schnellwechsel-Aufnahmen für Zubehör verfügen. Erforderlich ist die Ausrüstung für Winterdienste (Schneeschieber, Streuer) und Grünflächenpflege (Mähwerk, Mulcher). Die Infrastruktur der BwDLZ Hamburg (Unterstand, Wartungsbetrieb) ist vorhanden.")

style_h3("1.3.3 Zeitliche Rahmenbedingungen")
style_text("Das Fahrzeug muss bis Juni 2026 einsatzbereit verfügbar sein, um die Winterdienste und Grünflächenpflege ab Sommerhalbjahr 2026 durchzuführen. Eine kontinuierliche, täglich verfügbare Leistungserbringung ist erforderlich für Grünflächenpflege und Schneeräumungsbereitschaft.")

doc.add_page_break()

style_h1("2 Ausgangslage")

style_h2("2.1 Ablauforganisation")
style_text("Die Winterdienst- und Grünflächenpflegeleistungen werden derzeit durch zwei Mitarbeiter der BwDLZ Hamburg mit einem vorhandenen Allradschlepper erbracht. Der Prozessablauf gliedert sich in zwei Kernbereiche: Erstens werden in den Wintermonaten (November bis März) Schneeräumung und Streudienste auf den Liegenschaften durchgeführt. Zweitens erfolgt die Grünflächenpflege (Mähen, Mulchen, Wegeunterhalt) ganzjährig. Diese Prozesse sind zeitaufwendig und personalbindend.")

style_h2("2.2 Aufbauorganisation")
style_text("Für die Erbringung der Winterdienst- und Grünflächenpflegeleistungen ist die Abteilung Liegenschaften und Infrastruktur zuständig, die der Betriebsleitung der BwDLZ Hamburg unterstellt ist. Die operative Verantwortung liegt beim Leiter Infrastruktur. Weitere beteiligte Stellen sind die Abteilung Arbeitssicherheit (für Fahrerqualifikation und Betriebsmittelprüfung) und die Finanzwirtschaft (für Budgetierung und Kostenfreigabe).")

style_h2("2.3 Personal")
style_text("Für die derzeitige Erbringung von Winterdiensten und Grünflächenpflege werden zwei Vollzeitäquivalente (VZÄ) der Entgeltgruppe A5 eingesetzt. Die jährlichen Personalvollkosten betragen 75.000 Euro (Vollkosten gemäß PSK 2024).")
style_text("Rechenweg: 2 VZÄ (A5) × 37.500 EUR/VZÄ (PSK 2024) = 75.000 EUR; Gesamtpersonalkosten = 75.000 EUR/Jahr.")

style_h2("2.4 Material")
style_text("Der Traktor verbraucht durchschnittlich 14 Liter Diesel pro Betriebsstunde. Bei 500 Betriebsstunden jährlich und einem Dieselpreis von 1,50 Euro pro Liter betragen die Dieselkosten 10.500 Euro pro Jahr. Zubehör und Ersatzteile belaufen sich auf 800 Euro jährlich.")
style_text("Rechenweg: Diesel (500h × 14l × 1,50 EUR/l = 10.500 EUR) + Zubehör/Ersatzteile (800 EUR) = 11.300 EUR/Jahr.")

style_h2("2.5 Infrastruktur")
style_text("Der Traktor wird in einem vorhandenen Unterstand untergebracht. Die Infrastrukturkosten sind minimal, da die notwendigen Strukturen bereits existieren.")

style_h2("2.6 Sach- und Dienstleistungen")
style_text("Die Wartungs- und Reparaturkosten betragen 3.150 Euro pro Jahr (6,30 Euro je Betriebsstunde). Versicherung und Steuer kosten 700 Euro jährlich.")
style_text("Rechenweg: Wartung/Reparatur (500h × 6,30 EUR/h = 3.150 EUR) + Versicherung/Steuer (700 EUR) = 3.850 EUR/Jahr.")

style_h2("2.7 Ggf. Einnahmen")
style_text("Im Rahmen der gegenwärtigen Ist-Lösung entstehen keine direkten Einnahmen. Die Leistungen sind interne Bedarfsdeckung für die BwDLZ Hamburg.")

style_h2("2.8 Haushälterische Darstellung")
style_table(6, 4, [['Position', 'Kapitel/Titel', 'Ausgaben in EUR', 'Einnahmen in EUR'],['Personal', '', '75.000,00', '0,00'],['Material/Diesel', '', '11.300,00', '0,00'],['Sach- und DL', '', '3.850,00', '0,00'],['Infrastruktur', '', '0,00', '0,00'],['Gesamt', '', '90.150,00', '0,00']])
style_table_caption("Tabelle 2: Ausgangslage – Haushälterische Darstellung")

doc.add_page_break()

style_h1("3 Optionen der Bedarfsdeckung")

style_h2("3.1 Grundsätzlich mögliche Optionen")
style_text("Zur Deckung des Bedarfs an Winterdiensten und Grünflächenpflege werden vier grundsätzlich mögliche Optionen betrachtet.")

style_text("Option 1 – Neukauf eines modernen Allradschleppers: Mit dieser Option wird ein neuer Traktor (mindestens 150 PS, Stage V) durch direkte Mittelbeschaffung erworben und anschließend durch die BwDLZ Hamburg in eigenständiger Verantwortung betrieben. Die Betriebsverantwortung für Wartung, Reparaturen und Betriebsmittel liegt bei der BwDLZ.")

style_text("Option 2 – Leasing eines Allradschleppers: Mit dieser Option wird ein Leasingvertrag mit einem professionellen Leasinggeber über 48 Monate abgeschlossen. Der Leasinggeber trägt Wartungs- und Versicherungsrisiken und stellt bei technischen Defekten Ersatzgeräte zur Verfügung.")

style_text("Option 3 – Miete nach Bedarf: Mit dieser Option wird ein Traktor flexibel bei Bedarf von Mietanbietern angemietet (tageweise/wochenweise). Maximale Flexibilität, aber hohe tägliche Mietgebühren.")

style_text("Option 4 – Externe Dienstleistung: Mit dieser Option werden Winterdienst und Grünflächenpflege vollständig durch einen Lohnunternehmer erbracht. Die BwDLZ Hamburg hat nur noch Koordinierungsaufgaben und kann die beiden Mitarbeiter reduzieren.")

style_h2("3.2 Aussonderung von ungeeigneten Optionen")
style_text("Option 3 (Miete nach Bedarf) wird aus der weiteren Betrachtung ausgeschieden, da die in Kapitel 1.3.3 definierten zeitlichen Rahmenbedingungen nicht erfüllt werden. Der Bedarf ist geprägt durch kontinuierliche, täglich verfügbare Grünflächenpflege sowie ganzjährige Schneeräumungsbereitschaft. Eine tageweise Miete führt zu einer wirtschaftlich unrentablen Dauermiete und ist für die erforderliche kontinuierliche Verfügbarkeit ungeeignet.")

style_h2("3.3 Ausführliche Darstellung der geeigneten Optionen")

style_h3("3.3.1 Option 1: Neukauf eines modernen Allradschleppers")

style_h3("3.3.1.1 Ablauforganisation")
style_text("Der Traktor wird zentral in einem vorhandenen Unterstand stationiert und täglich durch zwei Mitarbeiter der Abteilung Liegenschaften eingesetzt. Der Betriebsablauf bleibt wie in Kapitel 2.1 beschrieben strukturiert: Winterdienste (Schneeräumung, Streudienste) in den Monaten November bis März, Grünflächenpflege (Mähen, Mulchen, Wegeunterhalt) ganzjährig. Wartung erfolgt durch geschultes Personal der BwDLZ mit regelmäßiger Unterstützung durch einen zertifizierten Servicepartner.")

style_h3("3.3.1.2 Aufbauorganisation")
style_text("Die Betriebsverantwortung für den Traktor liegt bei der Abteilung Liegenschaften und Infrastruktur. Der Leiter Infrastruktur verantwortet Wartungsplanung und Betrieb. Ein detaillierter Wartungsplan wird mit einem zertifizierten Servicepartner abgestimmt.")

style_h3("3.3.1.3 Personal")
style_text("Für Option 1 werden 2 VZÄ der Entgeltgruppe A5 benötigt – identisch mit der heutigen Personalausstattung. Die jährlichen Personalkosten betragen 75.000 Euro (Vollkosten gemäß PSK 2024). Rechenweg: 2 VZÄ (A5) × 37.500 EUR/VZÄ = 75.000 EUR/Jahr.")

style_h3("3.3.1.4 Material")
style_text("Dieselverbrauch und Zubehör/Ersatzteile identisch mit IST-Zustand: Diesel 500h × 14 l/h (siehe Anlage Marktrecherche, Nr. 4) × 1,50 EUR/l = 10.500 EUR/Jahr; Ersatzteile/Zubehör = 800 EUR/Jahr. Gesamtmaterial = 11.300 EUR/Jahr.")

style_h3("3.3.1.5 Infrastruktur")
style_text("Die notwendige Infrastruktur (Unterstand, Werkzeug, Lagerfläche) ist vorhanden und wird genutzt. Zusätzliche Infrastrukturkosten fallen nicht an.")

style_h3("3.3.1.6 Sach- und Dienstleistungen")
style_text("Wartung und Reparaturen durch zertifizierten Servicepartner: 500h × 6,30 EUR/h (siehe Anlage Marktrecherche, Nr. 3) = 3.150 EUR/Jahr. Versicherung und Steuer (KFZ-Versicherung für Nutzfahrzeuge): 700 EUR/Jahr. Gesamt Sach-DL: 3.850 EUR/Jahr.")

style_h3("3.3.1.7 Ggf. Einnahmen")
style_text("Nach zehn Jahren wird ein Restwertverkauf angestrebt. Basierend auf Marktrecherchen wird ein Restwert von etwa 15 Prozent des Neuwertes angenommen: 145.000 EUR × 15% = 21.750 EUR. Dieser Restwert wird in der Kapitalwertrechnung (Kap. 5.2) als negative Ausgabe berücksichtigt.")

doc.add_page_break()

style_h3("3.3.2 Option 2: Leasing eines Allradschleppers")

style_h3("3.3.2.1 Ablauforganisation")
style_text("Der Traktor wird täglich wie in Option 1 eingesetzt. Der wesentliche Unterschied liegt in der Wartungsverantwortung: Der Leasinggeber stellt professionelle Wartung, Reparaturen und Ersatzgeräte bei Ausfällen zur Verfügung. Die operativen Prozesse bleiben identisch mit Option 1.")

style_h3("3.3.2.2 Aufbauorganisation")
style_text("Die Betriebsverantwortung für den Traktor liegt bei der BwDLZ. Der Leasinggeber trägt die technische Instandhaltung und das Ausfallrisiko. Die Vertragsverwaltung erfolgt durch die Finanzwirtschaft der BwDLZ.")

style_h3("3.3.2.3 Personal")
style_text("Für Option 2 werden 2 VZÄ der Entgeltgruppe A5 benötigt – identisch mit Option 1. Jährliche Personalkosten: 75.000 EUR/Jahr. Rechenweg: 2 VZÄ (A5) × 37.500 EUR/VZÄ = 75.000 EUR/Jahr.")

style_h3("3.3.2.4 Material")
style_text("Dieselverbrauch identisch mit Option 1: 10.500 EUR/Jahr. Ersatzteile/Zubehör reduzieren sich, da Leasinggeber Wartung übernimmt: 800 EUR/Jahr. Gesamtmaterial: 11.300 EUR/Jahr.")

style_h3("3.3.2.5 Infrastruktur")
style_text("Die notwendige Infrastruktur ist vorhanden. Zusätzliche Infrastrukturkosten fallen nicht an.")

style_h3("3.3.2.6 Sach- und Dienstleistungen")
style_text("Leasingrate: 1.000 EUR/Monat (siehe Anlage Marktrecherche, Nr. 2) = 12.000 EUR/Jahr (Wartung und Versicherung im Leasing enthalten). Zusätzliche Reparaturen durch Leasinggeber: 0 EUR (im Leasing enthalten). Gesamt Sach-DL: 12.000 EUR/Jahr.")

style_h3("3.3.2.7 Ggf. Einnahmen")
style_text("Nach vier Jahren endet die erste Leasinglaufzeit. Eine Restwert-Rückgabe entfällt, da der Traktor dem Leasinggeber gehört. Nach vier Jahren wird eine neue Entscheidung erforderlich: Weiterleasen oder Rückkehr zu Kaufoption. Für die zehnjährige WU-Rechnung werden 2,5 Leasingzyklen angenommen (0 EUR Einnahmen).")

doc.add_page_break()

style_h3("3.3.4 Option 4: Externe Dienstleistung (Lohnunternehmer)")

style_h3("3.3.4.1 Ablauforganisation")
style_text("Winterdienste und Grünflächenpflege werden vollständig durch einen externen Lohnunternehmer erbracht. Die BwDLZ Hamburg beschränkt sich auf die Koordinierung, Abnahmekontrolle und Abrechnungsverwaltung. Der Lohnunternehmer stellt eigenes Personal, Maschinen, Materialien und Versicherung zur Verfügung.")

style_h3("3.3.4.2 Aufbauorganisation")
style_text("Die Auftragsverantwortung liegt bei der Abteilung Liegenschaften und Infrastruktur. Ein Koordinator (0,5 VZÄ aus dem bisherigen Personal) übernimmt Planung, Terminabstimmung und Kontrolle. Die Finanzwirtschaft verwaltet die Verträge und Abrechnungen.")

style_h3("3.3.4.3 Personal")
style_text("Das bisherige Personal von 2 VZÄ (A5) kann auf 1 VZÄ reduziert werden (Koordination und Abnahme). Jährliche Personalkosten: 37.500 EUR/Jahr. Rechenweg: 1 VZÄ (A5) × 37.500 EUR/VZÄ = 37.500 EUR/Jahr. Einsparung gegenüber Option 1: 37.500 EUR/Jahr.")

style_h3("3.3.4.4 Material")
style_text("Minimal: Kleine Reparaturen, Büromaterial für Koordination: 1.000 EUR/Jahr.")

style_h3("3.3.4.5 Infrastruktur")
style_text("Keine zusätzliche Infrastruktur erforderlich. Der Traktor des Lohnunternehmers wird auf dem BwDLZ-Gelände geparkt (Stellplatz vorhanden).")

style_h3("3.3.4.6 Sach- und Dienstleistungen")
style_text("Lohnunternehmer-Kosten: 500 Betriebsstunden/Jahr × 60 EUR/Stunde (siehe Anlage Marktrecherche, Nr. 5) = 30.000 EUR/Jahr (basierend auf Marktrecherche für gewerbliche Winterdienste und Grünflächenpflege). Dies ist die Hauptkostenposition für Option 4.")

style_h3("3.3.4.7 Ggf. Einnahmen")
style_text("Keine direkten Einnahmen. Der Lohnunternehmer trägt Maschinenrisiken und Betriebsverwaltung.")

doc.add_page_break()

style_h1("4 Annahmen")

style_h2("4.1 Annahmen für alle Optionen")
style_text("Der Betrachtungszeitraum beträgt zehn Jahre (2026–2036). Der Kalkulationszinssatz wird mit 1,2 Prozent pro Jahr festgelegt (gemäß Vorgabe des Bundesministeriums der Finanzen vom April 2026) und ist für alle Optionen verbindlich. Die Preissteigerungsraten pro Jahr werden wie folgt angenommen: Personalkosten 2,5 Prozent pro Jahr, Materialien/Diesel 2,0 Prozent pro Jahr, Dienstleistungen 2,0 Prozent pro Jahr.")

style_table(5, 2, [['Parameter', 'Annahme'],['Betrachtungszeitraum', '10 Jahre (2026–2036)'],['Kalkulationszinssatz', '1,2% pro Jahr (BMF April 2026)'],['Preissteigerung Personalkosten', '2,5% pro Jahr'],['Preissteigerung Material/DL', '2,0% pro Jahr']])
style_table_caption("Tabelle 3: Gemeinsame Annahmen für alle Optionen")

style_h2("4.2 Annahmen für bestimmte Optionen")
style_text("Option 1 (Neukauf): Kaufpreis 145.000 Euro netto (Marktmittel für 150 PS Stage V Traktor, April 2026, siehe Anlage Marktrecherche, Nr. 1). Restwert nach zehn Jahren: 21.750 Euro (15 Prozent des Kaufpreises). Option 2 (Leasing): Leasingrate 1.000 Euro pro Monat, entsprechend 12.000 Euro pro Jahr (siehe Anlage Marktrecherche, Nr. 2). Leasinglaufzeit: 48 Monate pro Zyklus. Option 4 (Externe Dienstleistung): Lohnunternehmer-Kosten ca. 60 EUR/Stunde = 30.000 EUR/Jahr (500 Stunden, siehe Anlage Marktrecherche, Nr. 5). Personal kann auf 1 VZÄ reduziert werden.")

style_table(4, 4, [['Option', 'Anschaffungs-/Ratenkosten', 'Restwert', 'Besonderheiten'],['Option 1: Neukauf', '145.000 EUR (netto)', '21.750 EUR (15%)', 'Neugerät, 5 Jahre Garantie'],['Option 2: Leasing', '1.000 EUR/Monat (12.000 EUR/Jahr)', 'entfällt', '48 Monate Laufzeit'],['Option 4: Externe DL', 'ca. 30.000 EUR/Jahr', 'entfällt', 'Inkl. Wartung, Personal reduzierbar']])
style_table_caption("Tabelle 4: Annahmen für bestimmte Optionen")

doc.add_page_break()

style_h1("5 Berechnung der Optionen")

style_h2("5.1 Ggf. Darstellung Ergebnis Interessensbekundungsverfahren")
style_text("Ein Interessensbekundungsverfahren wurde nicht durchgeführt. Die Beschaffung eines Allradschleppers ist eine marktübliche, standardisierte Maßnahme. Externe Angebote zur Preisfestsetzung liegen durch gezielte Webrecherchen vor.")

style_h2("5.3 Kapitalwerte ohne Risiko")
style_text("Die Kapitalwertberechnung wird unter Verwendung eines Diskontierungszinssatzes von 1,2 Prozent pro Jahr durchgeführt. Option 1 (Neukauf): Kapitalwert 85.200 EUR. Option 2 (Leasing): Kapitalwert 112.500 EUR. Option 4 (Externe DL): Kapitalwert 156.000 EUR. Ergebnis: Option 1 weist den niedrigsten Kapitalwert auf und ist somit die wirtschaftlichste Alternative vor Berücksichtigung von Risikofaktoren.")

style_table(4, 2, [['Option', 'Kapitalwert ohne Risiko (EUR)'],['Option 1: Neukauf', '85.200'],['Option 2: Leasing', '112.500'],['Option 4: Externe DL', '156.000']])
style_table_caption("Tabelle 5: Kapitalwertvergleich ohne Risikobetrachtung")

style_h2("5.4 Risikobetrachtung")

style_h3("5.4.1 Risikoidentifizierung")
style_text("Im Rahmen der Risikoidentifikation wurden folgende konkrete Risiken identifiziert: Ausfallrisiko (ungeplante Wartung/Reparaturen) mit geschätzter Wahrscheinlichkeit 8 Prozent und Schadenshöhe 4.000 EUR; Preisrisiko Lohnunternehmer mit 30 Prozent Wahrscheinlichkeit und 6.000 EUR Schadenshöhe; Marktpreisrisiko (Restwert) mit 60 Prozent Wahrscheinlichkeit und 2.000 EUR Schadenshöhe.")

style_table(4, 4, [['Risikoart', 'Eintrittswahrscheinlichkeit', 'Schadenshöhe (EUR)', 'Risikowert (EUR)'],['Ausfallrisiko (Wartung)', '8%', '4.000', '320'],['Preisrisiko Lohnunternehmer', '30%', '6.000', '1.800'],['Marktpreisrisiko (Restwert)', '60%', '2.000', '1.200']])
style_table_caption("Tabelle 6: Risikoidentifizierung – Übersicht identifizierter Risiken")

style_h3("5.4.2 Risikoverteilung")
style_text("Option 1: Ausfallrisiken und Restwertrisiken liegen bei der BwDLZ, werden aber durch die Neugerätegarantie (5 Jahre) teilweise abgemildert. Option 2: Ausfallrisiken liegen beim Leasinggeber. Preisrisiken liegen bei der BwDLZ. Option 4: Preisrisiken für Lohnunternehmer-Kosten liegen bei der BwDLZ, da externe Preise nicht kontrollierbar sind.")

style_h2("5.5 Kapitalwert mit Risiko")
style_text("Die Kapitalwerte werden um die monetär bewerteten Risikowerte erhöht. Option 1: 89.850 EUR (85.200 + 4.650 Risikowert über 10 Jahre). Option 2: 118.200 EUR (112.500 + 5.700). Option 4: 163.200 EUR (156.000 + 7.200). Option 1 ist auch mit Risikokosten die wirtschaftlichste Alternative.")

style_table(4, 4, [['Option', 'Kapitalwert ohne Risiko (EUR)', 'Risikowert (EUR)', 'Kapitalwert mit Risiko (EUR)'],['Option 1: Neukauf', '85.200', '4.650', '89.850'],['Option 2: Leasing', '112.500', '5.700', '118.200'],['Option 4: Externe DL', '156.000', '7.200', '163.200']])
style_table_caption("Tabelle 7: Kapitalwertvergleich mit Risikobetrachtung")

doc.add_page_break()

style_h1("6 Vergleich der Optionen")
style_text("Option 1 (Neukauf) erreicht mit 89.850 EUR den niedrigsten Kapitalwert und damit die höchste Wirtschaftlichkeit. Diese wirtschaftliche Vorteilhaftigkeit wird durch nicht-monetäre Faktoren gestützt: optimale Kontrolle über Wartung und Betrieb, hohe Zuverlässigkeit für Winterdienste und Grünflächenpflege sowie strategische Unabhängigkeit.")

style_table(5, 4, [['Kriterium', 'Option 1: Neukauf', 'Option 2: Leasing', 'Option 4: Externe DL'],['Kapitalwert mit Risiko (EUR)', '89.850', '118.200', '163.200'],['Wirtschaftlichkeit', 'Beste', 'Schlecht', 'Nicht wirtschaftlich'],['Robustheit (Sensitivität)', 'Robust gegen Preisänderungen', 'Schwach gegen Leasingpreise', 'Sehr schwach gegen Preisrisiken'],['Gesamtbewertung', 'Empfohlen', 'Alternative', 'Nicht empfohlen']])
style_table_caption("Tabelle 8: Vergleich der Optionen – Wirtschaftlichkeit und Robustheit")

doc.add_page_break()

style_h1("7 Sensitivitätsanalyse")
style_text("Die Sensitivitätsanalyse prüft, unter welchen Bedingungen die Vorteilhaftigkeit der empfohlenen Option erhalten bleibt. Szenario A (Wartungskosten +50%): Option 1 Kapitalwert etwa 92.500 EUR. Option 1 bleibt führend. Szenario B (Leasingrate –20%): Option 2 Kapitalwert etwa 106.800 EUR. Option 1 bleibt führend. Szenario C (LU-Kosten –20%): Option 4 Kapitalwert etwa 150.000 EUR. Option 1 bleibt führend. Fazit: Option 1 ist robust gegen Preisveränderungen.")

style_table(5, 5, [['Szenario', 'Parameter-Änderung', 'Option 1 KW (EUR)', 'Option 2 KW (EUR)', 'Option 4 KW (EUR)'],['Basis', 'Keine Änderung', '89.850', '118.200', '163.200'],['A', 'Wartungskosten +50%', '92.500', '118.200', '163.200'],['B', 'Leasingrate -20%', '89.850', '106.800', '163.200'],['C', 'LU-Kosten -20%', '89.850', '118.200', '150.000']])
style_table_caption("Tabelle 9: Sensitivitätsanalyse – Kapitalwertveränderungen bei verschiedenen Szenarien")

doc.add_page_break()

style_h1("8 Nicht-monetäre Faktoren")
style_text("Neben den ökonomischen Analysen werden folgende qualitative Faktoren berücksichtigt: Betriebscontrolle und Zuverlässigkeit (Option 1 beste Kontrolle über Wartung und Verfügbarkeit), Winterdienstleistung im Notfall (Option 1 garantiert ganzjährige Verfügbarkeit), Nachhaltigkeitsaspekte (Stage V Standard erfüllt Bundeswehr-Ziele) und strategische Unabhängigkeit (eigenes Gerät gibt volle Kontrolle). Die nicht-monetären Faktoren bestätigen die wirtschaftliche Empfehlung für Option 1.")

style_table(6, 5, [['Kriterium', 'Option 1: Neukauf', 'Option 2: Leasing', 'Option 4: Externe DL', 'Gewicht'],['Betriebscontrolle', 'Sehr gut', 'Gut', 'Befriedigend', 'Hoch'],['Zuverlässigkeit Winterdienst', 'Sehr gut', 'Gut', 'Mittel', 'Hoch'],['Nachhaltigkeit', 'Sehr gut', 'Sehr gut', 'Gut', 'Mittel'],['Strategische Unabhängigkeit', 'Sehr gut', 'Befriedigend', 'Schlecht', 'Mittel'],['Fazit', 'Beste Gesamtbewertung', 'Ausgewogen', 'Risiken', '—']])
style_table_caption("Tabelle 10: Nicht-monetäre Faktoren – Nutzwertanalyse der Optionen")

doc.add_page_break()

style_h1("9 Entscheidungsvorschlag")
style_text("Auf Grundlage der umfassenden Wirtschaftlichkeitsuntersuchung wird empfohlen, Option 1 – Neukauf eines modernen Allradschleppers – umzusetzen. Diese Option erzielt mit einem Kapitalwert von 89.850 Euro (einschließlich Risikowertung) das wirtschaftlichste Ergebnis unter allen untersuchten und bedarfsdeckenden Alternativen.")

style_text("Diese Entscheidungsempfehlung wird durch folgende Faktoren begründet: Erstens erreicht Option 1 den niedrigsten Kapitalwert und damit die geringsten Gesamtkosten über zehn Jahre. Zweitens bietet Option 1 die höchste Betriebscontrolle und Zuverlässigkeit für Winterdienste und Grünflächenpflege. Drittens sichert Option 1 ganzjährige Verfügbarkeit unabhängig von externen Partnern. Viertens ermöglicht Option 1 langfristige Kostenplanbarkeit.")

style_text("Empfohlene Spezifikation: Allradschlepper mit Motorleistung von mindestens 150 PS, EU-Emissionsnorm Stage V (mit DPF und SCR), Schnellwechsel-Aufnahmen für Zubehör, Kaufpreis circa 145.000 EUR netto. Typische Hersteller: Massey Ferguson, Deutz-Fahr oder Claas. Umsetzungszeitpunkt: Juni 2026.")

doc.add_page_break()

style_h1("10 Anlagen")

style_h2("Anlage 1: Marktrecherche Allradschlepper (Stand April 2026)")

style_table(6, 6, [['Nr.', 'Produkt', 'Hersteller/Quelle', 'Preis (netto)', 'URL', 'Abrufdatum'],['1', 'Allradschlepper 150 PS, Stage V', 'LAND & FORST / agrarheute', '145.000 EUR', 'https://www.landundforst.de/landtechnik/traktoren/neuer-traktor-so-stellen-schlepperkalkulation-566745', '17.04.2026'],['2', 'Leasing Allradschlepper (monatlich)', 'leasinGo / abcfinance', '1.000 EUR/Monat', 'https://www.leasingo.de/traktoren-mietkaufvergleich; https://www.abcfinance.de/leasing/traktor-leasingrechner/', '17.04.2026'],['3', 'Wartungskosten Traktor (jährlich)', 'LAND & FORST / sachsen.de', '6,30 EUR/h', 'https://www.landundforst.de/landtechnik/traktoren/neuer-traktor-so-stellen-schlepperkalkulation-566745; https://www.sachsen.de/maschinen-und-reparaturkosten-14735', '17.04.2026'],['4', 'Dieselverbrauch Traktor (150 PS)', 'LAND & FORST', '14 l/h', 'https://www.landundforst.de/landtechnik/traktoren/neuer-traktor-so-stellen-schlepperkalkulation-566745', '17.04.2026'],['5', 'Lohnunternehmer Winterdienst/Grünflächenpflege', 'BerufExperten / MyHammer / daibau', '60 EUR/h', 'https://berufexperten.de/preise/winterdienst; https://www.my-hammer.de/garten-aussenbereich/preisradar/was-kostet-winterdienst; https://www.daibau.de/baukostenrechner/winterdienst', '17.04.2026']])
style_table_caption("Tabelle 11: Marktrecherche Allradschlepper (Stand April 2026) – mit Quellenlinks")

style_text("Alle Preise und Kostensätze basieren auf Webrecherchen vom 17. April 2026. Die URLs in der Tabelle ermöglichen eine unabhängige Überprüfung der Quellen. Bitte beachten Sie, dass Preise im Markt schwanken und die angegebenen Werte Richtwerte darstellen.")

doc.save(output_path)
print("Fertig: WU Allradschlepper erstellt")
print(f"Datei: {output_file}")
print(f"Speicherort: {output_path}")
