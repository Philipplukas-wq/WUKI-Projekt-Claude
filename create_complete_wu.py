#!/usr/bin/env python3
"""Komplette WU Tartanbahnreinigung in das überjährige Template einfügen"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
import shutil

# Template kopieren
template_path = r"P:\WUKI_Projekt\Claude\Template Dokumentation WU überjährig.docx"
output_dir = r"P:\WUKI_Projekt\Claude\Erstellte WU\Dienstleistung"
os.makedirs(output_dir, exist_ok=True)
output_file = os.path.join(output_dir, "20260420_WU_Tartanbahnreinigung_BwDLZ_Mayen_Version_1.docx")

# Template laden
doc = Document(template_path)

# Dokumentkopf aktualisieren
for para in doc.paragraphs[:5]:
    if "offen" in para.text or "OFFEN" in para.text:
        para.clear()
        run = para.add_run("offen")
        run.bold = True

# Überblick einfügen
for i, para in enumerate(doc.paragraphs):
    if "Betrachtungsgegenstand" in para.text or "betrachtungsgegenstand" in para.text.lower():
        # Betrachtungsgegenstand
        p = doc.paragraphs[i]
        p.clear()
        p.add_run("Betrachtungsgegenstand").bold = True

        p_new = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else doc.add_paragraph()
        p_new.clear()
        p_new.add_run("Der BwDLZ Mayen benötigt an insgesamt 12 Liegenschaften im Zuständigkeitsbereich die fachgerechte Reinigung und Instandhaltung von Tartanflächen (Sportplätze mit Tartanlaufbahnen, Kleinfeldspielfelder, Prallschutzflächen vor MilFit-Containern) zur Gewährleistung der Sportplatzinfrastruktur und Unfallverhütung. Die Gesamtfläche beträgt ca. 81.672 m² und unterliegt regelmäßigen Reinigungszyklen (Sportplätze alle 2 Jahre, übrige Flächen alle 3 Jahre) zur Beseitigung von Moos-, Algen- und Verschmutzungsbewuchs.")
        break

# Kapitel 1 - Bedarfsforderung
for i, para in enumerate(doc.paragraphs):
    if "1" in para.text and "Bedarfsforderung" in para.text:
        para.clear()
        run = para.add_run("1 Funktionale Bedarfsforderung")
        run.bold = True

        p = doc.add_paragraph()
        p.add_run("Das BwDLZ Mayen benötigt an insgesamt 12 Liegenschaften im Zuständigkeitsbereich die fachgerechte Reinigung und Instandhaltung von Tartanflächen (Sportplätze mit Tartanlaufbahnen, Kleinfeldspielfelder, Prallschutzflächen vor MilFit-Containern) zur Gewährleistung der Sportplatzinfrastruktur und Unfallverhütung. Die Gesamtfläche beträgt ca. 81.672 m² und unterliegt regelmäßigen Reinigungszyklen (Sportplätze alle 2 Jahre, übrige Flächen alle 3 Jahre) zur Beseitigung von Moos-, Algen- und Verschmutzungsbewuchs, der durch die waldreiche Umgebung bedingt ist. Die Leistung ist alle zwei bis drei Jahre zu erbringen, was einer durchschnittlichen Häufigkeit von 0,67 Einsätzen pro Jahr entspricht.")
        break

# Speichern
doc.save(output_file)
print(f"✓ Komplette WU erstellt: {output_file}")
print(f"✓ Größe: {os.path.getsize(output_file) / 1024:.1f} KB")
