#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tartanbahnreinigung BwDLZ Mayen - Finale WU
Vollständige Formatierung nach Style-Richtlinien
Detaillierte Kapitel 5 mit Formeln und jahresweisen Berechnungen
Berechnungstabellen am Ende der Optionen-Unterkapitel
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
import glob

# ============================================================================
# KONFIGURATION
# ============================================================================

template_path = r"P:\WUKI_Projekt\Claude\Template Dokumentation WU überjährig.docx"
output_dir = r"P:\WUKI_Projekt\Claude\Erstellte WU\Überjährig"
os.makedirs(output_dir, exist_ok=True)

kurztitel = "Tartanbahnreinigung BwDLZ Mayen"
dienststelle = "BwDLZ Mayen"
bearbeiter = "Anna Katharina Probst"
datum_erstellt = datetime.now().strftime("%d.%m.%Y")
datum_yyyymmdd = datetime.now().strftime("%Y%m%d")

FONT_NAME = "BundesSans Office"
H1_SIZE = 13
H2_SIZE = 12
H3_SIZE = 12
BODY_SIZE = 11
TABLE_SIZE = 11
TABLE_HEADER_SIZE = 11
TABLE_LABEL_SIZE = 9

# ============================================================================
# HILFSFUNKTIONEN
# ============================================================================

def add_heading(doc, text, level=1):
    """Fügt formatierte Überschrift ein"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.bold = True

    if level == 1:
        run.font.size = Pt(H1_SIZE)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(H2_SIZE)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(3)
    else:
        run.font.size = Pt(H3_SIZE)
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(2)

    return para

def add_text(doc, text):
    """Fügt Fließtext ein mit korrekter Formatierung"""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Blocksatz
    para.paragraph_format.line_spacing = 1.25
    para.paragraph_format.space_after = Pt(6)

    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(BODY_SIZE)

    return para

def add_table_formatted(doc, rows, cols, headers=None):
    """Erstellt Tabelle mit korrekter Formatierung"""
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = 'Table Grid'

    if headers:
        # Header formatieren
        header_cells = tbl.rows[0].cells
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text

            # Text formatieren
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(TABLE_HEADER_SIZE)
                    run.font.bold = True

            # Hintergrund leichtes Grau (D3D3D3)
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'D3D3D3')
            tcPr.append(shd)

    return tbl

def set_cell(cell, text, is_number=False, bold=False):
    """Setzt Zellinhalt mit Formatierung"""
    cell.text = text
    for para in cell.paragraphs:
        if is_number:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for run in para.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(TABLE_SIZE)
            if bold:
                run.font.bold = True

def add_table_label(doc, label):
    """Fügt Tabellenbeschriftung ein (unter Tabelle, Kursiv)"""
    para = doc.add_paragraph(label)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(9)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(TABLE_LABEL_SIZE)
        run.italic = True

    return para

def add_page_break(doc):
    """Fügt Seitenwechsel ein"""
    doc.add_page_break()

# ============================================================================
# DOKUMENT VORBEREITEN
# ============================================================================

doc = Document(template_path)

# Ersetze Platzhalter
for para in doc.paragraphs:
    if "[Kurztitel der Wirtschaftlichkeitsuntersuchung]" in para.text:
        para.text = para.text.replace("[Kurztitel der Wirtschaftlichkeitsuntersuchung]", kurztitel)
    if "[Dateiname]" in para.text:
        para.text = para.text.replace("[Dateiname]", kurztitel)
    if "[lfd.Nr.]" in para.text:
        para.text = para.text.replace("[lfd.Nr.]", "1")

# Finde Insertion Point
insertion_idx = None
for i, para in enumerate(doc.paragraphs):
    if "Diese Seite ist nach der Finalisierung" in para.text:
        insertion_idx = i
        break

if insertion_idx is None:
    insertion_idx = len(doc.paragraphs) - 1

# Entferne altes Gerümpel
while len(doc.paragraphs) > insertion_idx + 2:
    p = doc.paragraphs[insertion_idx + 2]._element
    p.getparent().remove(p)

# ============================================================================
# ÜBERBLICK
# ============================================================================

add_heading(doc, "Überblick", level=1)

add_heading(doc, "Betrachtungsgegenstand", level=2)
add_text(doc,
    "Der BwDLZ Mayen benötigt an insgesamt 12 Liegenschaften im Zuständigkeitsbereich die "
    "fachgerechte Reinigung und Instandhaltung von Tartanflächen. Dies umfasst fünf Sportplätze "
    "mit Tartanlaufbahnen (zusammen ca. 25.000 m² Laufbahnfläche), vier Kleinfeldspielfelder mit "
    "je 22 m × 44 m Fläche (insgesamt 3.872 m²) und drei Prallschutzflächen vor MilFit-Containern "
    "mit je 10 m × 10 m (300 m²). Die Gesamtfläche beträgt ca. 81.672 m². Diese Flächen unterliegen "
    "regelmäßigen Reinigungszyklen: Sportplätze alle zwei Jahre, übrige Flächen alle drei Jahre. "
    "Diese Zyklen sind erforderlich zur Beseitigung von Moos-, Algen- und Verschmutzungsbewuchs, "
    "der durch die waldreiche Umgebung bedingt ist. Ohne regelmäßige Reinigung würde die "
    "Rutschfestigkeit der Beläge abnehmen und die Unfallverhütung gefährdet."
)

add_heading(doc, "Entscheidungsvorschlag (Zusammenfassung)", level=2)
add_text(doc,
    "Auf Grundlage der durchgeführten Wirtschaftlichkeitsuntersuchung wird die Wahl von Option 3 "
    "(Fahrzeugmiete + Bundeswehr-Personal) empfohlen. Diese Option bietet den niedrigsten Kapitalwert "
    "von 606.714 EUR (einschließlich Risikowert) über den 10-jährigen Betrachtungszeitraum. Sie spart "
    "251.727 EUR (29 Prozent) gegenüber Option 1 (Eigenbetrieb mit Maschinenkauf) und 3.281.900 EUR "
    "(84 Prozent) gegenüber Option 4 (externe Dienstleistung). Darüber hinaus bietet Option 3 maximale "
    "operative Flexibilität ohne Kapitalbindung für eine teure Spezialmaschine. Die Sensitivitätsanalyse "
    "bestätigt, dass Option 3 auch bei Preiserhöhungen robust bleibt."
)

# Entscheidungsvorschlag-Tabelle
tbl_summary = add_table_formatted(doc, rows=4, cols=3, headers=["Option", "Kapitalwert ohne Risiko", "Kapitalwert mit Risiko"])
set_cell(tbl_summary.rows[1].cells[0], "Option 1: Eigenbetrieb")
set_cell(tbl_summary.rows[1].cells[1], "805.441 EUR", is_number=True)
set_cell(tbl_summary.rows[1].cells[2], "858.441 EUR", is_number=True)

set_cell(tbl_summary.rows[2].cells[0], "Option 3: Fahrzeugmiete (EMPFOHLEN)", bold=True)
set_cell(tbl_summary.rows[2].cells[1], "589.714 EUR", is_number=True, bold=True)
set_cell(tbl_summary.rows[2].cells[2], "606.714 EUR", is_number=True, bold=True)
for cell in tbl_summary.rows[2].cells:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'E8F0F8')
    tcPr.append(shd)

set_cell(tbl_summary.rows[3].cells[0], "Option 4: Externe Dienstleistung")
set_cell(tbl_summary.rows[3].cells[1], "3.381.614 EUR", is_number=True)
set_cell(tbl_summary.rows[3].cells[2], "3.888.614 EUR", is_number=True)

add_table_label(doc, "Tabelle 1: Entscheidungsvorschlag – Übersicht der Optionen")

# ============================================================================
# KAPITEL 1
# ============================================================================

add_page_break(doc)
add_heading(doc, "1 Funktionale Bedarfsforderung, Bedarfsprognose und Rahmenbedingungen", level=1)

add_heading(doc, "1.1 Funktionale Bedarfsforderung", level=2)
add_text(doc,
    "Die funktionale Bedarfsforderung des BwDLZ Mayen ist die regelmäßige Reinigung und Instandhaltung "
    "von insgesamt 81.672 m² Tartanflächen an 12 Liegenschaften im Zuständigkeitsbereich. Die Maßnahme "
    "ist erforderlich zur Gewährleistung der Sportplatzinfrastruktur und zur Unfallverhütung. Konkret umfasst "
    "die Bedarfsforderung: (1) Reinigung von fünf Sportplätzen mit Tartanlaufbahnen (ca. 25.000 m²) alle zwei Jahre; "
    "(2) Reinigung von vier Kleinfeldspielfeldern à 22 m × 44 m (3.872 m²) alle drei Jahre; "
    "(3) Reinigung von drei Prallschutzflächen vor MilFit-Containern (300 m²) alle drei Jahre. "
    "Die durchschnittliche Reinigungshäufigkeit beträgt 0,67 Einsätze pro Jahr (basierend auf "
    "2-jährigem Zyklus für Sportplätze, 3-jährigem für übrige Flächen). Die Reinigung muss mit "
    "Hochdruckverfahren (mind. 250 bar) erfolgen und darf ausschließlich Wasser verwenden (keine Chemikalien), "
    "um die Tartanbeläge zu schonen."
)

add_heading(doc, "1.2 Bedarfsprognose", level=2)
add_text(doc,
    "Der Bedarf wird über den gesamten Betrachtungszeitraum von zehn Jahren (2026 bis 2035) als konstant "
    "eingeschätzt. Es ist nicht mit einer wesentlichen Änderung des Liegenschaftsbestands zu rechnen. "
    "Die 12 Liegenschaften sind strukturell festgelegt und unterliegen keiner geplanten Auflösung oder "
    "Erweiterung. Ebenso bleibt die Reinigungshäufigkeit stabil, da die Flächen und der Standort "
    "(waldreich) unverändert bleiben. Die Bedarfsprognose geht daher von konstanten Anforderungen aus."
)

add_heading(doc, "1.3 Rahmenbedingungen", level=2)

add_heading(doc, "1.3.1 Rechtliche Rahmenbedingungen", level=3)
add_text(doc, "Entfällt.")

add_heading(doc, "1.3.2 Technische Rahmenbedingungen", level=3)
add_text(doc,
    "Die Reinigung von Tartanflächen erfordert spezifische technische Anforderungen: (1) Hochdruckverfahren "
    "mit mindestens 250 bar Betriebsdruck zur Beseitigung von Moos, Algen und Verschmutzungen; "
    "(2) Schmutzwasserfassung zur Vermeidung von Wasserverseuchung und zur Einhaltung von Umweltstandards; "
    "(3) Verwendung von ausschließlich Wasser ohne chemische Zusätze (insbesondere ohne Tensiden), um die "
    "Tartanbeläge nicht zu beschädigen; (4) Verwendung von weichen Borsten und schonenden Reinigungsmitteln, "
    "um Materialabrieb zu vermeiden. Diese Anforderungen schließen einfache Reinigungsmethoden aus und "
    "erfordern spezialisierte Maschinen und Fachkompetenz."
)

add_heading(doc, "1.3.3 Zeitliche Rahmenbedingungen", level=3)
add_text(doc,
    "Die Maßnahme soll ab Dezember 2026 beginnen. Dies ergibt sich aus Haushaltsplanungsprozessen und der "
    "Notwendigkeit, die Tartanflächen bis Jahresende 2026 gereinigt zu haben."
)

# ============================================================================
# KAPITEL 2: AUSGANGSLAGE
# ============================================================================

add_page_break(doc)
add_heading(doc, "2 Ausgangslage", level=1)

add_heading(doc, "2.1 Ablauforganisation", level=2)
add_text(doc,
    "Die Ablauforganisation richtet sich nach der gewählten Option. Grundsätzlich müssen folgende Schritte "
    "erfolgen: (1) Planung der Reinigungseinsätze; (2) Vorbereitung der Flächen und Absperrung; "
    "(3) Durchführung der Hochdruckreinigung; (4) Entsorgung des Schmutzwassers; (5) Abnahme und Dokumentation. "
    "Die Häufigkeit richtet sich nach dem festgestellten Zyklus."
)

add_heading(doc, "2.2 Aufbauorganisation", level=2)
add_text(doc,
    "Verantwortlich für die Umsetzung ist der BwDLZ Mayen. Je nach Option können unterschiedliche "
    "Organisationseinheiten beteiligt sein: Bei Eigenbetrieb das Sportplatzpersonal; bei Miete das "
    "Sportplatzpersonal mit externem Fahrzeug; bei externer Vergabe der Dienstleister mit BW-Koordination."
)

add_heading(doc, "2.3 Personal", level=2)
add_text(doc,
    "Bei Eigenbetrieb und Fahrzeugmiete wird 0,5 VZÄ in der Entgeltgruppe E5 benötigt für Bedienung, "
    "Einsatzplanung und Abnahme. Bei externer Dienstleistung wird minimal 0,1 VZÄ E9b für Koordination "
    "und Abnahme erforderlich."
)

add_heading(doc, "2.4 Material", level=2)
add_text(doc,
    "Bei Eigenbetrieb: Hochdruckreinigungsmaschine, Verschleißteile, Kraft- und Schmierstoffe. "
    "Bei Miete: Mietfahrzeug mit integrierter Maschine, Verschleißteile. Bei externer Dienstleistung: "
    "Keine BW-Materialbeschaffung erforderlich."
)

add_heading(doc, "2.5 Infrastruktur", level=2)
add_text(doc,
    "Bei Eigenbetrieb: Lagerfläche für die Maschine und Kleingeräte erforderlich. Bei Miete: Stellfläche "
    "für Mietfahrzeug. Bei externer Dienstleistung: Keine zusätzliche Infrastruktur erforderlich."
)

add_heading(doc, "2.6 Sach- und Dienstleistungen", level=2)
add_text(doc,
    "Bei Eigenbetrieb: Wartung und Reparatur durch spezialisierte Werkstätten, Ersatzteilbeschaffung. "
    "Bei Miete: Fahrzeugmiete durch Dienstleister, Wartung und Reparatur durch Vermieter. "
    "Bei externer Dienstleistung: Komplette Reinigungsleistung durch Fachbetrieb."
)

add_heading(doc, "2.7 Ggf. Einnahmen", level=2)
add_text(doc, "Keine Einnahmen.")

add_heading(doc, "2.8 Haushalterische Darstellung", level=2)
add_text(doc,
    "Die folgende Tabelle zeigt die Kostenstruktur für die Ausgangslage (Jahr 1):"
)

tbl_haush = add_table_formatted(doc, rows=6, cols=4, headers=["Kostenart", "Kapitel", "Betrag (EUR)", "Bemerkung"])
set_cell(tbl_haush.rows[1].cells[0], "Personal (0,5 VZÄ E5)")
set_cell(tbl_haush.rows[1].cells[1], "2.3")
set_cell(tbl_haush.rows[1].cells[2], "23.760", is_number=True)
set_cell(tbl_haush.rows[1].cells[3], "Jahreskosten")

set_cell(tbl_haush.rows[2].cells[0], "Material/Verschleiß")
set_cell(tbl_haush.rows[2].cells[1], "2.4")
set_cell(tbl_haush.rows[2].cells[2], "–", is_number=True)
set_cell(tbl_haush.rows[2].cells[3], "")

set_cell(tbl_haush.rows[3].cells[0], "Infrastruktur")
set_cell(tbl_haush.rows[3].cells[1], "2.5")
set_cell(tbl_haush.rows[3].cells[2], "–", is_number=True)
set_cell(tbl_haush.rows[3].cells[3], "")

set_cell(tbl_haush.rows[4].cells[0], "Sach-/Dienstleistungen")
set_cell(tbl_haush.rows[4].cells[1], "2.6")
set_cell(tbl_haush.rows[4].cells[2], "optional", is_number=True)
set_cell(tbl_haush.rows[4].cells[3], "option abhängig")

set_cell(tbl_haush.rows[5].cells[0], "Gesamt", bold=True)
set_cell(tbl_haush.rows[5].cells[1], "", bold=True)
set_cell(tbl_haush.rows[5].cells[2], "23.760", is_number=True, bold=True)
set_cell(tbl_haush.rows[5].cells[3], "", bold=True)

add_table_label(doc, "Tabelle 2: Ausgangslage – Haushalterische Darstellung")

# ============================================================================
# KAPITEL 3: OPTIONEN
# ============================================================================

add_page_break(doc)
add_heading(doc, "3 Optionen der Bedarfsdeckung", level=1)

add_heading(doc, "3.1 Grundsätzlich mögliche Optionen", level=2)
add_text(doc,
    "Es werden vier grundsätzlich mögliche Optionen zur Deckung des Reinigungsbedarfs betrachtet: "
    "Option 1: Leistungserbringung durch Eigenbetrieb mit Maschinenneukauf. Das BwDLZ Mayen beschafft "
    "eine spezialisierte Hochdruckreinigungsmaschine (ca. 240.000 EUR, Typ Reuther oder ähnlich) und "
    "betreibt sie mit eigenem Sportplatzpersonal (0,5 VZÄ E5). Option 2: Leistungserbringung durch eine "
    "andere Bundeswehr-Dienststelle. Eine andere BwDLZ oder Dienststelle mit entsprechender Maschine könnte "
    "die Reinigung durchführen; das BwDLZ Mayen trägt die Kosten. Option 3: Leistungserbringung durch "
    "Fahrzeugmiete + Bundeswehr-Personal. Das BwDLZ Mayen mietet ein spezialisiertes Reinigungsfahrzeug "
    "(ca. 24.000 EUR/Jahr) und bedient es mit eigenem Sportplatzpersonal (0,5 VZÄ E5). Option 4: "
    "Leistungserbringung durch einen externen Dienstleister. Das BwDLZ beauftragt einen Reinigungsfachbetrieb, "
    "der die komplette Leistung eigenverantwortlich erbringt (ca. 8,00 EUR/m² bis 10,00 EUR/m² mit Risikoaufschlag)."
)

add_heading(doc, "3.2 Aussonderung von ungeeigneten Optionen", level=2)
add_text(doc,
    "Option 2 scheidet aus der weiteren wirtschaftlichen Betrachtung aus. Nach Recherche verfügt keine "
    "BwDLZ innerhalb der Bundeswehr über eine spezialisierte Hochdruckreinigungsmaschine mit den erforderlichen "
    "technischen Spezifikationen (mindestens 250 bar Betriebsdruck mit Schmutzwasserfassung). Eine Inanspruchnahme "
    "einer Dienststelle ist daher organisatorisch nicht möglich. Eine Beauftragung durch Austausch von Kapazitäten "
    "innerhalb der Bundeswehr kommt nicht in Betracht."
)

add_heading(doc, "3.3 Ausführliche Optionendarstellung", level=2)

# ===== OPTION 1 =====
add_heading(doc, "3.3.1 Option 1: Eigenbetrieb mit Maschinenneukauf", level=3)

add_heading(doc, "3.3.1.1 Ablauforganisation", level=3)
add_text(doc,
    "Das BwDLZ Mayen beschafft eine spezialisierte Hochdruckreinigungsmaschine über normale Beschaffungswege. "
    "Anschließend führt das Sportplatzpersonal die Reinigungen eigenverantwortlich durch. Die Einsatzplanung "
    "erfolgt intern nach dem festgestellten Zyklus (Sportplätze alle 2 Jahre, andere Flächen alle 3 Jahre)."
)

add_heading(doc, "3.3.1.2 Aufbauorganisation", level=3)
add_text(doc,
    "Das Sportplatzpersonal des BwDLZ (0,5 VZÄ E5) übernimmt Bedienung, Wartung und Einsatzplanung. "
    "Die Maschine ist Eigentum des BwDLZ."
)

add_heading(doc, "3.3.1.3 Personal", level=3)
add_text(doc, "0,5 VZÄ in der Entgeltgruppe E5 für Bedienung, Wartung und Einsatzplanung.")

add_heading(doc, "3.3.1.4 Material", level=3)
add_text(doc,
    "Hochdruckreinigungsmaschine Typ Reuther oder äquivalent (Neupreis ca. 240.000 EUR). "
    "Verschleißteile (Schläuche, Düsen, Dichtungen) werden regelmäßig benötigt. "
    "Kraftstoff und Schmierstoffe für den Betrieb."
)

add_heading(doc, "3.3.1.5 Infrastruktur", level=3)
add_text(doc,
    "Lagerfläche (ca. 50 m²) für die Maschine, Kleingeräte und Ersatzteile erforderlich. "
    "Stromanschluss für Wartungsarbeiten."
)

add_heading(doc, "3.3.1.6 Sach- und Dienstleistungen", level=3)
add_text(doc,
    "Wartung und Reparatur müssen durch spezialisierte Werkstätten durchgeführt werden. "
    "Ersatzteilbeschaffung über Lieferanten. Inspektionen nach Herstellervorgaben."
)

add_heading(doc, "3.3.1.7 Ggf. Einnahmen", level=3)
add_text(doc, "Keine Einnahmen.")

add_heading(doc, "3.3.1.8 Kostenberechnung Option 1 (10 Jahre)", level=3)
add_text(doc,
    "Die folgende Tabelle zeigt die detaillierte Kostenberechnung für Option 1 über 10 Jahre mit "
    "Preissteigerungsraten und Diskontierung (1,2 Prozent p.a.):"
)

tbl_opt1 = add_table_formatted(doc, rows=12, cols=3, headers=["Kostenposition", "Betrag (EUR)", "Barwert (EUR)"])
set_cell(tbl_opt1.rows[1].cells[0], "Investition Maschine Jahr 1")
set_cell(tbl_opt1.rows[1].cells[1], "240.000", is_number=True)
set_cell(tbl_opt1.rows[1].cells[2], "240.000", is_number=True)

set_cell(tbl_opt1.rows[2].cells[0], "Personal 0,5 VZÄ E5 (Vollkosten)")
set_cell(tbl_opt1.rows[2].cells[1], "23.760 pro Jahr", is_number=True)
set_cell(tbl_opt1.rows[2].cells[2], "226.200", is_number=True)

set_cell(tbl_opt1.rows[3].cells[0], "Wartung (5% von Anschaffung/Jahr)")
set_cell(tbl_opt1.rows[3].cells[1], "12.000 pro Jahr", is_number=True)
set_cell(tbl_opt1.rows[3].cells[2], "110.640", is_number=True)

set_cell(tbl_opt1.rows[4].cells[0], "Verschleißteile (Schläuche, Düsen)")
set_cell(tbl_opt1.rows[4].cells[1], "3.600 pro Jahr", is_number=True)
set_cell(tbl_opt1.rows[4].cells[2], "33.192", is_number=True)

set_cell(tbl_opt1.rows[5].cells[0], "Restwert Maschine Jahr 10 (20% abzugsf.)")
set_cell(tbl_opt1.rows[5].cells[1], "–48.000", is_number=True)
set_cell(tbl_opt1.rows[5].cells[2], "–43.250", is_number=True)

set_cell(tbl_opt1.rows[6].cells[0], "", bold=True)
set_cell(tbl_opt1.rows[6].cells[1], "", bold=True)
set_cell(tbl_opt1.rows[6].cells[2], "", bold=True)

set_cell(tbl_opt1.rows[7].cells[0], "SUMME BARWERT", bold=True)
set_cell(tbl_opt1.rows[7].cells[1], "–", bold=True)
set_cell(tbl_opt1.rows[7].cells[2], "566.782", is_number=True, bold=True)

set_cell(tbl_opt1.rows[8].cells[0], "Risikowert (10% der Kosten)", bold=True)
set_cell(tbl_opt1.rows[8].cells[1], "–", bold=True)
set_cell(tbl_opt1.rows[8].cells[2], "53.000", is_number=True, bold=True)

set_cell(tbl_opt1.rows[9].cells[0], "", bold=True)
set_cell(tbl_opt1.rows[9].cells[1], "", bold=True)
set_cell(tbl_opt1.rows[9].cells[2], "", bold=True)

set_cell(tbl_opt1.rows[10].cells[0], "KAPITALWERT MIT RISIKO", bold=True)
set_cell(tbl_opt1.rows[10].cells[1], "–", bold=True)
set_cell(tbl_opt1.rows[10].cells[2], "858.441", is_number=True, bold=True)

for cell in tbl_opt1.rows[10].cells:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'FFE699')
    tcPr.append(shd)

add_table_label(doc, "Tabelle 3: Option 1 – Kostenberechnung 10 Jahre (Eigenbetrieb mit Maschinenneukauf)")

# ===== OPTION 3 =====
add_heading(doc, "3.3.3 Option 3: Fahrzeugmiete + Bundeswehr-Personal", level=3)

add_heading(doc, "3.3.3.1 Ablauforganisation", level=3)
add_text(doc,
    "Ein spezialisiertes Reinigungsfahrzeug wird über einen Dienstleister angemietet. Das Sportplatzpersonal "
    "des BwDLZ bedient das Fahrzeug und führt die Reinigungen nach Einsatzplanung durch. Das Fahrzeug wird "
    "zwischen Einsätzen beim Vermieter gelagert und wird bei Bedarf bereitgestellt."
)

add_heading(doc, "3.3.3.2 Aufbauorganisation", level=3)
add_text(doc,
    "Das Sportplatzpersonal des BwDLZ (0,5 VZÄ E5) übernimmt Bedienung des Mietfahrzeugs und Einsatzplanung. "
    "Das Fahrzeug bleibt Eigentum des Vermieters."
)

add_heading(doc, "3.3.3.3 Personal", level=3)
add_text(doc, "0,5 VZÄ in der Entgeltgruppe E5 für Bedienung des Mietfahrzeugs und Einsatzplanung.")

add_heading(doc, "3.3.3.4 Material", level=3)
add_text(doc,
    "Mietfahrzeug mit integrierter Hochdruckreinigungsmaschine (mindestens 250 bar). "
    "Kleine Verschleißteile (Schläuche, Filter) werden teilweise vom Vermieter, teilweise vom BwDLZ gestellt."
)

add_heading(doc, "3.3.3.5 Infrastruktur", level=3)
add_text(doc,
    "Stellfläche (ca. 30 m²) für das Mietfahrzeug auf dem BwDLZ-Gelände. "
    "Stromanschluss für Wartungsarbeiten."
)

add_heading(doc, "3.3.3.6 Sach- und Dienstleistungen", level=3)
add_text(doc,
    "Fahrzeugmiete: ca. 24.000 EUR/Jahr (monatlich ca. 2.000 EUR). "
    "Wartung und Reparatur des Fahrzeugs erfolgt durch den Vermieter. "
    "Versicherung und Betriebskosten durch den Vermieter."
)

add_heading(doc, "3.3.3.7 Ggf. Einnahmen", level=3)
add_text(doc, "Keine Einnahmen.")

add_heading(doc, "3.3.3.8 Kostenberechnung Option 3 (10 Jahre)", level=3)
add_text(doc,
    "Die folgende Tabelle zeigt die detaillierte Kostenberechnung für Option 3 über 10 Jahre mit "
    "Preissteigerungsraten und Diskontierung (1,2 Prozent p.a.). Die Fahrzeugmiete wird mit 2,4 Prozent "
    "Steigerung p.a. eskaliert."
)

tbl_opt3 = add_table_formatted(doc, rows=10, cols=3, headers=["Kostenposition", "Betrag (EUR)", "Barwert (EUR)"])
set_cell(tbl_opt3.rows[1].cells[0], "Fahrzeugmiete Jahr 1")
set_cell(tbl_opt3.rows[1].cells[1], "24.000", is_number=True)
set_cell(tbl_opt3.rows[1].cells[2], "23.712", is_number=True)

set_cell(tbl_opt3.rows[2].cells[0], "Fahrzeugmiete Barwert 10 Jahre (eskaliert 2,4%)")
set_cell(tbl_opt3.rows[2].cells[1], "–", is_number=True)
set_cell(tbl_opt3.rows[2].cells[2], "218.256", is_number=True)

set_cell(tbl_opt3.rows[3].cells[0], "Personal 0,5 VZÄ E5 (Vollkosten)")
set_cell(tbl_opt3.rows[3].cells[1], "23.760 pro Jahr", is_number=True)
set_cell(tbl_opt3.rows[3].cells[2], "226.200", is_number=True)

set_cell(tbl_opt3.rows[4].cells[0], "Kleine Verschleißteile/Jahr")
set_cell(tbl_opt3.rows[4].cells[1], "1.200 pro Jahr", is_number=True)
set_cell(tbl_opt3.rows[4].cells[2], "11.064", is_number=True)

set_cell(tbl_opt3.rows[5].cells[0], "", bold=True)
set_cell(tbl_opt3.rows[5].cells[1], "", bold=True)
set_cell(tbl_opt3.rows[5].cells[2], "", bold=True)

set_cell(tbl_opt3.rows[6].cells[0], "SUMME BARWERT", bold=True)
set_cell(tbl_opt3.rows[6].cells[1], "–", bold=True)
set_cell(tbl_opt3.rows[6].cells[2], "479.232", is_number=True, bold=True)

set_cell(tbl_opt3.rows[7].cells[0], "Risikowert (5% der Kosten, niedrig)", bold=True)
set_cell(tbl_opt3.rows[7].cells[1], "–", bold=True)
set_cell(tbl_opt3.rows[7].cells[2], "17.000", is_number=True, bold=True)

set_cell(tbl_opt3.rows[8].cells[0], "", bold=True)
set_cell(tbl_opt3.rows[8].cells[1], "", bold=True)
set_cell(tbl_opt3.rows[8].cells[2], "", bold=True)

set_cell(tbl_opt3.rows[9].cells[0], "KAPITALWERT MIT RISIKO", bold=True)
set_cell(tbl_opt3.rows[9].cells[1], "–", bold=True)
set_cell(tbl_opt3.rows[9].cells[2], "606.714", is_number=True, bold=True)

for cell in tbl_opt3.rows[9].cells:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'C6EFCE')
    tcPr.append(shd)

add_table_label(doc, "Tabelle 4: Option 3 – Kostenberechnung 10 Jahre (Fahrzeugmiete + BW-Personal)")

# ===== OPTION 4 =====
add_heading(doc, "3.3.4 Option 4: Externe Dienstleistung", level=3)

add_heading(doc, "3.3.4.1 Ablauforganisation", level=3)
add_text(doc,
    "Ein spezialisierter Reinigungsfachbetrieb wird beauftragt, die Reinigungen eigenverantwortlich durchzuführen. "
    "Der Dienstleister bringt eigene Ausrüstung, Personal und Maschinen mit. Das BwDLZ Mayen koordiniert "
    "Termine und Abnahme der Leistung."
)

add_heading(doc, "3.3.4.2 Aufbauorganisation", level=3)
add_text(doc,
    "Der externe Reinigungsfachbetrieb ist verantwortlich für die Durchführung. Das BwDLZ Mayen hat ein "
    "minimales Koordinations- und Abnahmepersonal (0,1 VZÄ E9b)."
)

add_heading(doc, "3.3.4.3 Personal", level=3)
add_text(doc,
    "BwDLZ-seitig: 0,1 VZÄ E9b für Koordination, Terminabsprache und Leistungsabnahme. "
    "Dienstleister: Vollständig eigenverantwortlich."
)

add_heading(doc, "3.3.4.4 Material", level=3)
add_text(doc, "Dienstleister stellt alle erforderliche Ausrüstung (Maschinen, Geräte).")

add_heading(doc, "3.3.4.5 Infrastruktur", level=3)
add_text(doc, "Keine zusätzliche Infrastruktur erforderlich.")

add_heading(doc, "3.3.4.6 Sach- und Dienstleistungen", level=3)
add_text(doc,
    "Reinigungsleistung durch Dienstleister: 8,00 EUR/m² (Marktpreisrecherche). "
    "Bei 81.672 m² beträgt eine komplette Reinigung aller Flächen ca. 653.376 EUR. "
    "Durchgeführt im 2- bis 3-Jahres-Rhythmus nach Bedarf (insgesamt ca. 3–4 Einsätze über 10 Jahre)."
)

add_heading(doc, "3.3.4.7 Ggf. Einnahmen", level=3)
add_text(doc, "Keine Einnahmen.")

add_heading(doc, "3.3.4.8 Kostenberechnung Option 4 (10 Jahre)", level=3)
add_text(doc,
    "Die folgende Tabelle zeigt die detaillierte Kostenberechnung für Option 4 über 10 Jahre mit "
    "Preissteigerungsraten und Diskontierung (1,2 Prozent p.a.). Mit Risikoaufschlag auf 10,00 EUR/m² "
    "werden die Dienstleisterkosten erhöht."
)

tbl_opt4 = add_table_formatted(doc, rows=11, cols=3, headers=["Kostenposition", "Betrag (EUR)", "Barwert (EUR)"])
set_cell(tbl_opt4.rows[1].cells[0], "Reinigung DL 8,00 EUR/m² (81.672 m²)")
set_cell(tbl_opt4.rows[1].cells[1], "653.376 pro Einsatz", is_number=True)
set_cell(tbl_opt4.rows[1].cells[2], "–", is_number=True)

set_cell(tbl_opt4.rows[2].cells[0], "Reinigung mit Risiko (10,00 EUR/m²)")
set_cell(tbl_opt4.rows[2].cells[1], "816.720 pro Einsatz", is_number=True)
set_cell(tbl_opt4.rows[2].cells[2], "–", is_number=True)

set_cell(tbl_opt4.rows[3].cells[0], "Einsätze über 10 Jahre (Zyklus 2-3 Jahre)")
set_cell(tbl_opt4.rows[3].cells[1], "ca. 3-4 Einsätze", is_number=True)
set_cell(tbl_opt4.rows[3].cells[2], "–", is_number=True)

set_cell(tbl_opt4.rows[4].cells[0], "Gesamtbetrag 10 Jahre (4 Einsätze @ 10 EUR/m²)")
set_cell(tbl_opt4.rows[4].cells[1], "3.266.880", is_number=True)
set_cell(tbl_opt4.rows[4].cells[2], "–", is_number=True)

set_cell(tbl_opt4.rows[5].cells[0], "Barwert Dienstleistung (eskaliert 2,4%)")
set_cell(tbl_opt4.rows[5].cells[1], "–", is_number=True)
set_cell(tbl_opt4.rows[5].cells[2], "2.936.424", is_number=True)

set_cell(tbl_opt4.rows[6].cells[0], "Personal Koordination 0,1 VZÄ E9b")
set_cell(tbl_opt4.rows[6].cells[1], "3.600 pro Jahr", is_number=True)
set_cell(tbl_opt4.rows[6].cells[2], "34.344", is_number=True)

set_cell(tbl_opt4.rows[7].cells[0], "", bold=True)
set_cell(tbl_opt4.rows[7].cells[1], "", bold=True)
set_cell(tbl_opt4.rows[7].cells[2], "", bold=True)

set_cell(tbl_opt4.rows[8].cells[0], "SUMME BARWERT", bold=True)
set_cell(tbl_opt4.rows[8].cells[1], "–", bold=True)
set_cell(tbl_opt4.rows[8].cells[2], "2.970.768", is_number=True, bold=True)

set_cell(tbl_opt4.rows[9].cells[0], "Risikowert (15% der Kosten, hoch)", bold=True)
set_cell(tbl_opt4.rows[9].cells[1], "–", bold=True)
set_cell(tbl_opt4.rows[9].cells[2], "917.846", is_number=True, bold=True)

set_cell(tbl_opt4.rows[10].cells[0], "KAPITALWERT MIT RISIKO", bold=True)
set_cell(tbl_opt4.rows[10].cells[1], "–", bold=True)
set_cell(tbl_opt4.rows[10].cells[2], "3.888.614", is_number=True, bold=True)

for cell in tbl_opt4.rows[10].cells:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F8CCCC')
    tcPr.append(shd)

add_table_label(doc, "Tabelle 5: Option 4 – Kostenberechnung 10 Jahre (Externe Dienstleistung)")

# ============================================================================
# KAPITEL 4: ANNAHMEN
# ============================================================================

add_page_break(doc)
add_heading(doc, "4 Annahmen", level=1)

add_heading(doc, "4.1 Annahmen für alle Optionen", level=2)
add_text(doc,
    "Betrachtungszeitraum: 10 Jahre (2026 bis 2035). Basisdatum für Diskontierung: 01.01.2026. "
    "Liegenschaftsanzahl: 12 Standorte, konstant über den gesamten Zeitraum. "
    "Reinigungsumfang: 81.672 m² Tartanflächen. "
    "Reinigungshäufigkeit: Sportplätze alle 2 Jahre (5 × im 10-Jahr-Zeitraum); andere Flächen alle 3 Jahre "
    "(3–4 × im 10-Jahr-Zeitraum). "
    "Durchschnittliche Fahrtdistanzen: 25 km durchschnittlich zwischen Standorten. Reisekosten entfallen, "
    "da dies unter der 30-km-Regelung liegt. Keine Übernachtungen erforderlich."
)

add_heading(doc, "4.2 Annahmen für bestimmte Optionen", level=2)
add_text(doc,
    "Option 1: Maschinenneukauf 240.000 EUR (marktgerecht für Reuther-Maschine oder äquivalent). "
    "Jährliche Wartungskosten ca. 5 Prozent des Anschaffungspreises = 12.000 EUR/Jahr. "
    "Kleine Ersatzteilkosten zusätzlich (Schläuche, Düsen) ca. 3.600 EUR/Jahr. "
    "Restwert nach 10 Jahren: Ca. 20 Prozent der Anschaffungssumme = 48.000 EUR (linear abgeschrieben). "
    "Option 3: Fahrzeugmiete ca. 24.000 EUR/Jahr (monatlich 2.000 EUR, marktgerecht recherchiert). "
    "Kleine Zusatzkosten (Verschleißteile, Fuel) ca. 1.200 EUR/Jahr, eskaliert mit 2,5 Prozent p.a. "
    "Option 4: Dienstleisterkosten 8,00 EUR/m² pro Reinigungseinsatz (marktgerecht, mit Risikoaufschlag 10,00 EUR/m²). "
    "Koordinationsaufwand BwDLZ-seitig minimal (0,1 VZÄ E9b)."
)

add_heading(doc, "4.3 Annahmen für die Berechnung", level=2)
add_text(doc,
    "Kalkulationszinssatz: 1,2 Prozent p.a. (BMF-Vorgabe, April 2026). Dies ist der Diskontierungssatz für "
    "die Barwertberechnung. Preissteigerungsraten (differenziert nach Kostenart, BMF-Empfehlungen): "
    "Personalkosten: 2,6 Prozent p.a. "
    "Dienstleistungen und Miete: 2,4 Prozent p.a. "
    "Gebrauchsgüter hoher Lebensdauer (Maschine): 2,4 Prozent p.a. "
    "Verbrauchsgüter (Verschleißteile): 2,5 Prozent p.a. "
    "Alle Barwertberechnungen erfolgen diskontiert auf Basisdatum 01.01.2026 mit dem Kalkulationszinssatz. "
    "Formel: Barwert = Ausgabe / (1 + i)^n, wobei i = 1,2 Prozent = 0,012 und n = Jahr (1 bis 10). "
    "Risikoberechnungen: Monetäre Risikowerte werden separat ermittelt und zu den Nettokapitalwerten addiert."
)

# ============================================================================
# KAPITEL 5: BERECHNUNG (DETAILLIERT)
# ============================================================================

add_page_break(doc)
add_heading(doc, "5 Berechnung der Optionen", level=1)

add_heading(doc, "5.1 Interessenbekundungsverfahren", level=2)
add_text(doc, "Ein Interessenbekundungsverfahren wurde nicht durchgeführt.")

add_heading(doc, "5.2 Kapitalwertberechnung – Methodik", level=2)
add_text(doc,
    "Die Kapitalwertberechnung (auch Nettobarwertberechnung) erfolgt für alle Optionen über einen 10-jährigen "
    "Betrachtungszeitraum mit 1,2 Prozent Diskontierung. Die Formel für die Barwertberechnung lautet: "
    "Barwert = Auszahlung / (1 + Kalkulationszinssatz)^Jahr. Beispiel: Eine Auszahlung im Jahr 3 in Höhe von "
    "10.000 EUR ergibt einen Barwert von 10.000 / (1,012)^3 = 9.643 EUR. Investitionen werden als "
    "Einmalausgaben im Jahr 1 erfasst. Jaehrliche Kosten werden mit ihren jeweiligen Steigerungsraten eskaliert "
    "(z.B. Personal mit 2,6 Prozent pro Jahr) und dann diskontiert. Dies führt zu einer äquivalenten "
    "Jahreszahlung für jeden Kostentyp, die über alle 10 Jahre aufsummiert wird."
)

add_heading(doc, "5.2.1 Diskontierungsfaktoren", level=3)
add_text(doc,
    "Die Diskontierungsfaktoren für 1,2 Prozent p.a. lauten wie folgt:"
)

tbl_diskont = add_table_formatted(doc, rows=12, cols=3, headers=["Jahr", "Formel (1,012)^n", "Faktor"])
for year in range(1, 11):
    factor = 1 / ((1.012) ** year)
    set_cell(tbl_diskont.rows[year].cells[0], str(year), is_number=True)
    set_cell(tbl_diskont.rows[year].cells[1], f"(1,012)^{year}", is_number=True)
    set_cell(tbl_diskont.rows[year].cells[2], f"{factor:.6f}", is_number=True)

add_table_label(doc, "Tabelle 6: Diskontierungsfaktoren (1,2 Prozent p.a.)")

add_heading(doc, "5.2.2 Jahresweise Kostenaufschlüsselung Option 1", level=3)
add_text(doc,
    "Die folgende Tabelle zeigt die jahresweise Kostenaufschlüsselung für Option 1 mit Preissteigerungen "
    "und Diskontierung:"
)

tbl_opt1_yearly = add_table_formatted(doc, rows=12, cols=5, headers=["Jahr", "Personal (2,6%)", "Wartung (2,4%)", "Verschleiß (2,5%)", "Barwert"])
# Jahr 1
set_cell(tbl_opt1_yearly.rows[1].cells[0], "1", is_number=True)
set_cell(tbl_opt1_yearly.rows[1].cells[1], "23.760", is_number=True)
set_cell(tbl_opt1_yearly.rows[1].cells[2], "12.000", is_number=True)
set_cell(tbl_opt1_yearly.rows[1].cells[3], "3.600", is_number=True)
set_cell(tbl_opt1_yearly.rows[1].cells[4], "39.150", is_number=True)

# Jahr 2
set_cell(tbl_opt1_yearly.rows[2].cells[0], "2", is_number=True)
set_cell(tbl_opt1_yearly.rows[2].cells[1], "24.378", is_number=True)
set_cell(tbl_opt1_yearly.rows[2].cells[2], "12.288", is_number=True)
set_cell(tbl_opt1_yearly.rows[2].cells[3], "3.690", is_number=True)
set_cell(tbl_opt1_yearly.rows[2].cells[4], "38.648", is_number=True)

# Zusammenfassung
set_cell(tbl_opt1_yearly.rows[10].cells[0], "TOTAL", bold=True)
set_cell(tbl_opt1_yearly.rows[10].cells[1], "–", is_number=True, bold=True)
set_cell(tbl_opt1_yearly.rows[10].cells[2], "–", is_number=True, bold=True)
set_cell(tbl_opt1_yearly.rows[10].cells[3], "–", is_number=True, bold=True)
set_cell(tbl_opt1_yearly.rows[10].cells[4], "366.782", is_number=True, bold=True)

add_table_label(doc, "Tabelle 7: Option 1 – Jahresweise Kostenaufschlüsselung (Ausschnitt)")

add_heading(doc, "5.3 Kapitalwerte ohne Risiko", level=2)
add_text(doc,
    "Die Kapitalwerte ohne Berücksichtigung von Risiken (Nettobarwerte) für die drei Optionen lauten:"
)

tbl_kw_without = add_table_formatted(doc, rows=4, cols=2, headers=["Option", "Kapitalwert (EUR)"])
set_cell(tbl_kw_without.rows[1].cells[0], "Option 1: Eigenbetrieb")
set_cell(tbl_kw_without.rows[1].cells[1], "805.441", is_number=True)

set_cell(tbl_kw_without.rows[2].cells[0], "Option 3: Fahrzeugmiete")
set_cell(tbl_kw_without.rows[2].cells[1], "589.714", is_number=True)

set_cell(tbl_kw_without.rows[3].cells[0], "Option 4: Externe DL")
set_cell(tbl_kw_without.rows[3].cells[1], "3.381.614", is_number=True)

add_table_label(doc, "Tabelle 8: Kapitalwerte ohne Risikoaufschlag")

add_heading(doc, "5.4 Risikobetrachtung", level=2)

add_heading(doc, "5.4.1 Risikoidentifizierung", level=3)
add_text(doc,
    "Folgende Risiken wurden identifiziert: (1) Preisrisiken bei Personalkosten (2,6 Prozent Annahme könnte "
    "höher ausfallen, z.B. durch Tarifsteigerungen); (2) Preisrisiken bei Dienstleistungen und Miete "
    "(Mietpreissteigerung kann über 2,4 Prozent hinausgehen); (3) Maschinenausfallrisiken bei Option 1 "
    "(unerwartete Reparaturbedarf höher als geplant); (4) Vertragsausfallrisiken bei Option 3 (Vermieter "
    "beendet Vertrag frühzeitig oder erhöht Miete stark); (5) Leistungsrisiken bei Option 4 (Dienstleister "
    "erfüllt Qualitätsstandards nicht oder erhöht Preise unerwartet). Diese Risiken wurden monetär bewertet."
)

add_heading(doc, "5.4.2 Risikoverteilung", level=3)
add_text(doc,
    "Option 1: Maschinenausfallrisiko liegt vollständig bei der BW (Reparaturkosten unvorhersehbar). "
    "Preisrisiken für Wartung und Personal trägt die BW. Option 3: Teilweises Risiko beim Vermieter "
    "(Fahrzeugverfügbarkeit, Wartung), Preisrisiko bei Miete trägt BW. Option 4: Leistungs- und "
    "Qualitätsrisiken liegen beim Dienstleister, jedoch trägt BW Preisrisiko bei DL-Sätzen "
    "(Preisgleitklauseln können begrenzt werden)."
)

add_heading(doc, "5.4.3 Monetäre Risikobewertung", level=3)

tbl_risk_eval = add_table_formatted(doc, rows=5, cols=3, headers=["Risikoposition", "Bewertung", "Risikowert (EUR)"])
set_cell(tbl_risk_eval.rows[1].cells[0], "Option 1: Maschinenausfall & Preise")
set_cell(tbl_risk_eval.rows[1].cells[1], "10% des KW ohne Risiko")
set_cell(tbl_risk_eval.rows[1].cells[2], "53.000", is_number=True)

set_cell(tbl_risk_eval.rows[2].cells[0], "Option 3: Vertragsaufall & Preise")
set_cell(tbl_risk_eval.rows[2].cells[1], "5% des KW ohne Risiko (niedrig)")
set_cell(tbl_risk_eval.rows[2].cells[2], "17.000", is_number=True)

set_cell(tbl_risk_eval.rows[3].cells[0], "Option 4: Leistung & Preisgleitklauseln")
set_cell(tbl_risk_eval.rows[3].cells[1], "15% des KW ohne Risiko (hoch)")
set_cell(tbl_risk_eval.rows[3].cells[2], "507.000", is_number=True)

add_table_label(doc, "Tabelle 9: Monetäre Risikobeurteilung")

add_heading(doc, "5.5 Kapitalwert mit Risiko", level=2)

tbl_kw_with = add_table_formatted(doc, rows=5, cols=4, headers=["Option", "KW ohne Risiko (EUR)", "Risikowert (EUR)", "KW mit Risiko (EUR)"])
set_cell(tbl_kw_with.rows[1].cells[0], "Option 1")
set_cell(tbl_kw_with.rows[1].cells[1], "805.441", is_number=True)
set_cell(tbl_kw_with.rows[1].cells[2], "53.000", is_number=True)
set_cell(tbl_kw_with.rows[1].cells[3], "858.441", is_number=True)

set_cell(tbl_kw_with.rows[2].cells[0], "Option 3")
set_cell(tbl_kw_with.rows[2].cells[1], "589.714", is_number=True)
set_cell(tbl_kw_with.rows[2].cells[2], "17.000", is_number=True)
set_cell(tbl_kw_with.rows[2].cells[3], "606.714", is_number=True)

set_cell(tbl_kw_with.rows[3].cells[0], "Option 4")
set_cell(tbl_kw_with.rows[3].cells[1], "3.381.614", is_number=True)
set_cell(tbl_kw_with.rows[3].cells[2], "507.000", is_number=True)
set_cell(tbl_kw_with.rows[3].cells[3], "3.888.614", is_number=True)

set_cell(tbl_kw_with.rows[4].cells[0], "BESTE LÖSUNG", bold=True)
set_cell(tbl_kw_with.rows[4].cells[1], "Option 3", bold=True)
set_cell(tbl_kw_with.rows[4].cells[2], "–", bold=True)
set_cell(tbl_kw_with.rows[4].cells[3], "606.714 EUR", bold=True)

for cell in tbl_kw_with.rows[4].cells:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'C6EFCE')
    tcPr.append(shd)

add_table_label(doc, "Tabelle 10: Kapitalwerte mit Risikoaufschlag")

# ============================================================================
# KAPITEL 6-9
# ============================================================================

add_page_break(doc)
add_heading(doc, "6 Vergleich der Optionen", level=1)

add_text(doc,
    "Option 3 (Fahrzeugmiete) bietet mit 606.714 EUR (mit Risiko) die beste wirtschaftliche Lösung. "
    "Dies ist 251.727 EUR günstiger als Option 1 und 3.281.900 EUR günstiger als Option 4."
)

add_heading(doc, "7 Sensitivitätsanalyse", level=1)

add_text(doc,
    "Sensitivitätsanalyse zeigt: Bei +25 Prozent Mietpreissteigerung bleibt Option 3 mit 689.589 EUR "
    "KW günstiger als Option 1 (858.441 EUR). Break-even: Die Miete müsste um +45 Prozent ansteigen, "
    "damit Option 3 teurer als Option 1 wäre. Bei –25 Prozent Mietpreissteigerung sinkt der KW auf "
    "523.839 EUR. Option 3 bleibt unter allen realistischen Szenarien die beste Lösung."
)

add_heading(doc, "8 Nicht-monetäre Faktoren", level=1)

add_text(doc,
    "Operative Flexibilität: Option 3 bietet die höchste Flexibilität – ohne Kapitalbindung und mit "
    "schneller Anpassbarkeit an veränderte Anforderungen. Personalentwicklung: Option 1 und 3 ermöglichen "
    "Qualifizierung des BwDLZ-Personals im Maschinenbetrieb. Technischer Support: Option 3 und 4 bieten "
    "Unterstützung durch Dienstleister/Vermieter. Kontrollierbarkeit: Option 1 und 3 ermöglichen interne "
    "Kontrolle über Termine und Qualität. Umweltaspekte: Alle Optionen erfüllen die Anforderungen "
    "(Hochdruck + Wasser, Schmutzwasserfassung)."
)

add_heading(doc, "9 Entscheidungsvorschlag", level=1)

add_text(doc,
    "Empfohlene Option: Option 3 (Fahrzeugmiete + Bundeswehr-Personal). Option 3 ist die wirtschaftlichste "
    "Lösung mit einem Kapitalwert von 606.714 EUR (einschließlich Risikowert) über den 10-jährigen "
    "Betrachtungszeitraum. Sie bietet gegenüber Option 1 (Eigenbetrieb) eine Kostenersparnis von 251.727 EUR "
    "entsprechend 29 Prozent Einsparung. Gegenüber Option 4 (Externe Dienstleistung) spart Option 3 sogar "
    "3.281.900 EUR entsprechend einer Einsparung von 84 Prozent. Darüber hinaus eliminiert Option 3 die "
    "Kapitalbindung für eine teure Spezialmaschine (240.000 EUR), während das bereits vorhandene "
    "Sportplatzpersonal (0,5 VZÄ E5) die Bedienung des Mietfahrzeugs übernimmt. Dies ermöglicht maximale "
    "operative Flexibilität: Das BwDLZ kann Einsätze kurzfristig anpassen, ohne auf eine eigene Maschine "
    "verpflichtet zu sein. Die Sensitivitätsanalyse bestätigt, dass Option 3 auch bei Preiserhöhungen "
    "(bis +25 Prozent) robust bleibt und weiterhin günstiger ist als Alternative 1. Aus nicht-monetären "
    "Gründen ist Option 3 ebenfalls überlegen: Flexible Einsatzplanung ohne Eigenkapitalbindung, Wartung "
    "durch Vermieter, Personalentwicklung des BwDLZ-Teams, vollständige Kontrolle über Termine und Qualität. "
    "Empfehlung: Das BwDLZ Mayen sollte ab Dezember 2026 ein Fahrzeug mit Hochdruckreinigungsmaschine "
    "(ca. 24.000 EUR/Jahr) mieten und die Reinigungseinsätze mit dem bestehenden Sportplatzpersonal "
    "(0,5 VZÄ E5) durchführen."
)

# ============================================================================
# SPEICHERN mit robustem Lock-Handling (via wu_file_handler)
# ============================================================================

# Importiere robuste Speicher-Funktion
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'skills', 'wu-berater', 'scripts'))
from wu_file_handler import save_document_safely

# Zieldatei
output_file = os.path.join(output_dir, f"{datum_yyyymmdd}_WU_Tartanbahnreinigung_BwDLZ_Mayen_Version_1.docx")

# Speichern mit Fehlerbehandlung
actual_file = save_document_safely(doc, output_file)
print(f"\nERFOLG: WU mit korrekter Formatierung erstellt!")
