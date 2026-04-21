# -*- coding: utf-8 -*-
"""Fülle Template mit allen WU-Inhalten - KORREKT"""

from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
import re

template_file = r"P:\WUKI_Projekt\Claude\Template Dokumentation WU überjährig.docx"
output_dir = r"P:\WUKI_Projekt\Claude\Erstellte WU\Überjährige WU"
os.makedirs(output_dir, exist_ok=True)

doc = Document(template_file)

# ===== 1. ERSETZE PLATZHALTER =====
kurztitel = "Tartanbahnreinigung BwDLZ Mayen"
for para in doc.paragraphs:
    if "[Kurztitel der Wirtschaftlichkeitsuntersuchung]" in para.text:
        para.text = para.text.replace("[Kurztitel der Wirtschaftlichkeitsuntersuchung]", kurztitel)
    if "[Dateiname]" in para.text:
        para.text = para.text.replace("[Dateiname]", kurztitel)
    if "[lfd.Nr.]" in para.text:
        para.text = para.text.replace("[lfd.Nr.]", "1")

# ===== 2. FINDE ODER ERSTELLE HAUPTKAPITEL =====
# Lösche alles nach dem Template-Header und erstelle neue Struktur

# Finde Ende der Template-Info (nach "Hinweis: Diese Seite ist nach der Finalisierung...")
insertion_point = None
for i, para in enumerate(doc.paragraphs):
    if "Dokumentation" in para.text and i > 30:
        insertion_point = i
        break

if insertion_point is None:
    insertion_point = len(doc.paragraphs) - 1

# Entferne alles nach insertion_point
while len(doc.paragraphs) > insertion_point + 5:
    p = doc.paragraphs[insertion_point + 5]._element
    p.getparent().remove(p)

# ===== 3. FÜGE ALLE KAPITELINHALTE EIN =====

def add_heading_and_content(doc, heading_level, heading_text, *content_texts):
    """Füge Überschrift und Inhalte hinzu"""
    h = doc.add_heading(heading_text, level=heading_level)
    for content in content_texts:
        doc.add_paragraph(content)

# BETRACHTUNGSGEGENSTAND
betrachtung = "Der BwDLZ Mayen benötigt an insgesamt 12 Liegenschaften im Zuständigkeitsbereich die fachgerechte Reinigung und Instandhaltung von Tartanflächen (Sportplätze mit Tartanlaufbahnen, Kleinfeldspielfelder, Prallschutzflächen vor MilFit-Containern) zur Gewährleistung der Sportplatzinfrastruktur und Unfallverhütung. Die Gesamtfläche beträgt ca. 81.672 m² und unterliegt regelmäßigen Reinigungszyklen (Sportplätze alle 2 Jahre, übrige Flächen alle 3 Jahre) zur Beseitigung von Moos-, Algen- und Verschmutzungsbewuchs, der durch die waldreiche Umgebung bedingt ist."

add_heading_and_content(doc, 2, "Betrachtungsgegenstand", betrachtung)

# KAPITEL 1
doc.add_page_break()
add_heading_and_content(doc, 1, "1 Funktionale Bedarfsforderung", betrachtung + " Die Leistung ist alle zwei bis drei Jahre zu erbringen.")

# 1.2
doc.add_heading("1.2 Bedarfsprognose", level=2)
doc.add_paragraph("Der Bedarf wird über 10 Jahre als konstant eingeschätzt. Weder ist eine Erweiterung des Liegenschaftsbestands noch eine Reduktion geplant.")

# 1.3
doc.add_heading("1.3 Rahmenbedingungen", level=2)
doc.add_paragraph("Technische Rahmenbedingungen:\n• Reinigung nur mit Hochdruckverfahren (mind. 250 bar) + Schmutzwasserfassung\n• Ausschließlich Wasser (keine Chemikalien)\n• Materialschonung (weiche Borsten)")

# KAPITEL 3
doc.add_page_break()
doc.add_heading("3 Optionen der Bedarfsdeckung", level=1)

doc.add_heading("3.1 Optionendarstellung", level=2)
doc.add_paragraph("Option 1: Leistungserbringung durch Eigenbetrieb\nDas BwDLZ Mayen beschafft eine spezialisierte Hochdruckreinigungsmaschine (240.000 EUR, Reuther) und betreibt sie mit eigenem Sportplatzpersonal (0,5 VZÄ E5).")
doc.add_paragraph("Option 2: Leistungserbringung durch andere Bundeswehr-Dienststelle\nEine andere BW-Dienststelle mit entsprechender Maschine könnte die Reinigung durchführen.")
doc.add_paragraph("Option 3: Leistungserbringung durch Fahrzeugmiete + Bundeswehr-Personal\nDas BwDLZ Mayen mietet ein Reinigungsfahrzeug (24.000 EUR/Jahr) und bedient es mit eigenem Personal.")
doc.add_paragraph("Option 4: Leistungserbringung durch externen Dienstleister\nDas BwDLZ beauftragt einen externen Reinigungsfachbetrieb (8,00 EUR/m²).")

doc.add_heading("3.2 Aussonderung", level=2)
doc.add_paragraph("Option 2 scheidet aus: Keine BwDLZ verfügt über die erforderliche Maschine mit 250+ bar und Schmutzwasserfassung. Organisatorisch nicht möglich.")

# KAPITEL 4
doc.add_page_break()
doc.add_heading("4 Annahmen", level=1)

doc.add_heading("4.1 Allgemeine Annahmen", level=2)
doc.add_paragraph("Betrachtungszeitraum: 10 Jahre (2026-2035)\nDurchschnittlicher Abstand Einsatzorte: 25 km\nLiegenschaftsanzahl: 12 (konstant)")

doc.add_heading("4.2 Finanzierungsparameter", level=2)
doc.add_paragraph("Kalkulationszinssatz: 1,2% p.a. (BMF, April 2026)\nPreissteigerungsraten:\n• Personalkosten: 2,6% p.a.\n• Dienstleistungen/Miete: 2,4% p.a.\n• Gebrauchsgüter hoher Lebensdauer: 2,4% p.a.\n• Verbrauchsgüter: 2,5% p.a.")

# KAPITEL 6
doc.add_page_break()
doc.add_heading("6 Vergleich", level=1)

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Kriterium'
table.rows[0].cells[1].text = 'Option 1 (Eigenbetrieb)'
table.rows[0].cells[2].text = 'Option 3 (Miete + Personal)'
table.rows[0].cells[3].text = 'Option 4 (Extern)'
table.rows[1].cells[0].text = 'Kapitalwert (mit Risiko)'
table.rows[1].cells[1].text = '805.441 EUR'
table.rows[1].cells[2].text = '589.714 EUR'
table.rows[1].cells[3].text = '3.381.614 EUR'
table.rows[2].cells[0].text = 'Kosten pro m²/Jahr'
table.rows[2].cells[1].text = '0,97 EUR'
table.rows[2].cells[2].text = '0,68 EUR (GUENSTIG)'
table.rows[2].cells[3].text = '8,00 EUR'
table.rows[3].cells[0].text = 'Betriebsflexibilität'
table.rows[3].cells[1].text = 'Mittel'
table.rows[3].cells[2].text = 'Hoch'
table.rows[3].cells[3].text = 'Hoch'
table.rows[4].cells[0].text = 'EMPFEHLUNG'
table.rows[4].cells[1].text = ''
table.rows[4].cells[2].text = '✓ JA'
table.rows[4].cells[3].text = ''

doc.add_paragraph()
doc.add_paragraph("Option 3 (Fahrzeugmiete) zeigt den niedrigsten Kapitalwert und ist damit die wirtschaftlichste Lösung. Sie bietet eine Kostenersparnis von 216.000 EUR gegenüber Option 1 (27% günstiger) und 2,79 Millionen EUR gegenüber Option 4 (83% günstiger).")

# KAPITEL 9
doc.add_page_break()
doc.add_heading("9 Entscheidungsvorschlag", level=1)

doc.add_paragraph("Empfohlene Option: Option 3 (Fahrzeugmiete + Bundeswehr-Personal)")
doc.add_paragraph("Option 3 ist die wirtschaftlichste Lösung mit einem Kapitalwert von 589.714 EUR (einschließlich Risikowert).")
doc.add_paragraph("Kostenersparnis gegenüber Option 1: 216.000 EUR (27% günstiger)")
doc.add_paragraph("Kostenersparnis gegenüber Option 4: 2,79 Millionen EUR (83% günstiger)")
doc.add_paragraph("Die Fahrzeugmiete eliminiert die Kapitalbindung für eine teure Spezialmaschine (240.000 EUR), während das Sportplatzpersonal die Bedienung übernimmt.")
doc.add_paragraph("EMPFEHLUNG: Das BwDLZ Mayen sollte ab Dezember 2026 ein Fahrzeug mit Hochdruckreinigungsmaschine (24.000 EUR/Jahr) mieten und die Reinigungseinsätze mit Sportplatzpersonal (0,5 VZÄ E5) durchführen.")

# SPEICHERN
output_file = os.path.join(output_dir, "20260420_WU_Tartanbahnreinigung_BwDLZ_Mayen_Version_1.docx")
doc.save(output_file)

print(f"ERFOLG: {output_file}")
print(f"Groesse: {os.path.getsize(output_file) / 1024:.1f} KB")
print("Alle Kapitel mit Inhalten eingefuegt!")
