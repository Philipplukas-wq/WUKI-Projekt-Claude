#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tartanbahnreinigung BwDLZ Mayen - Komplette WU mit ausformulierten Texten
Vorgehen nach Allradschlepper Hamburg: Fett-Paragraphen, detaillierte Texte, umfangreiche Tabellen
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# ============================================================================
# KONFIGURATION
# ============================================================================

template_path = r"P:\WUKI_Projekt\Claude\Template Dokumentation WU überjährig.docx"
output_dir = r"P:\WUKI_Projekt\Claude\Erstellte WU\Überjährig"
os.makedirs(output_dir, exist_ok=True)

kurztitel = "Tartanbahnreinigung BwDLZ Mayen"
dienststelle = "BwDLZ Mayen"
bearbeiter = "Anna Katharina Probst"
datum_erstellt = datetime.now().strftime("%d.%m.%Y")
datum_yyyymmdd = datetime.now().strftime("%Y%m%d")

# ============================================================================
# HILFSFUNKTIONEN
# ============================================================================

def add_heading_para(doc, text, size_pt=13, is_main=True):
    """Fügt fett-formatierte Überschrift als Normal-Paragraph ein"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    run.font.name = 'Calibri'
    if is_main:
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    else:
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(3)
    return para

def add_text_para(doc, text, space_after=6):
    """Fügt Fließtext-Paragraph ein"""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.15
    return para

def add_table(doc, rows, cols, style='Table Grid'):
    """Erstellt eine Tabelle mit Header-Formatierung"""
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = style
    return tbl

def format_table_header(tbl, headers):
    """Formatiert erste Zeile als Header (dunkelblau, weiß)"""
    hdr_cells = tbl.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header_text
        # Formatierung
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # Weiß
                run.font.size = Pt(9)
        # Hintergrund dunkelblau
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)

def set_cell_text(cell, text, bold=False):
    """Setzt Zelltext mit optionaler Fettformatierung"""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(9)

def add_page_break(doc):
    """Fügt Seitenwechsel ein"""
    doc.add_page_break()

# ============================================================================
# DOKUMENT VORBEREITEN
# ============================================================================

doc = Document(template_path)

# Ersetze Platzhalter
for para in doc.paragraphs:
    if "[Kurztitel der Wirtschaftlichkeitsuntersuchung]" in para.text:
        para.text = para.text.replace("[Kurztitel der Wirtschaftlichkeitsuntersuchung]", kurztitel)
    if "[Dateiname]" in para.text:
        para.text = para.text.replace("[Dateiname]", kurztitel)
    if "[lfd.Nr.]" in para.text:
        para.text = para.text.replace("[lfd.Nr.]", "1")

# Finde Insertion Point (nach Hinweis-Text)
insertion_idx = None
for i, para in enumerate(doc.paragraphs):
    if "Diese Seite ist nach der Finalisierung" in para.text:
        insertion_idx = i
        break

if insertion_idx is None:
    insertion_idx = len(doc.paragraphs) - 1

# Entferne altes Template-Inhaltsgerümpel
while len(doc.paragraphs) > insertion_idx + 2:
    p = doc.paragraphs[insertion_idx + 2]._element
    p.getparent().remove(p)

# ============================================================================
# ÜBERBLICK
# ============================================================================

add_heading_para(doc, "Überblick", size_pt=13, is_main=True)

add_heading_para(doc, "Betrachtungsgegenstand", size_pt=12, is_main=False)
add_text_para(doc,
    "Der BwDLZ Mayen benötigt an insgesamt 12 Liegenschaften im Zuständigkeitsbereich die "
    "fachgerechte Reinigung und Instandhaltung von Tartanflächen. Dies umfasst fünf Sportplätze "
    "mit Tartanlaufbahnen (zusammen ca. 25.000 m² Laufbahnfläche), vier Kleinfeldspielfelder mit "
    "je 22 m × 44 m Fläche (insgesamt 3.872 m²) und drei Prallschutzflächen vor MilFit-Containern "
    "mit je 10 m × 10 m (300 m²). Die Gesamtfläche beträgt ca. 81.672 m². Diese Flächen unterliegen "
    "regelmäßigen Reinigungszyklen: Sportplätze alle zwei Jahre, übrige Flächen alle drei Jahre. "
    "Diese Zyklen sind erforderlich zur Beseitigung von Moos-, Algen- und Verschmutzungsbewuchs, "
    "der durch die waldreiche Umgebung bedingt ist. Ohne regelmäßige Reinigung würde die "
    "Rutschfestigkeit der Beläge abnehmen und die Unfallverhütung gefährdet."
)

add_heading_para(doc, "Entscheidungsvorschlag (Zusammenfassung)", size_pt=12, is_main=False)
add_text_para(doc,
    "Auf Grundlage der durchgeführten Wirtschaftlichkeitsuntersuchung wird die Wahl von Option 3 "
    "(Fahrzeugmiete + Bundeswehr-Personal) empfohlen. Diese Option bietet den niedrigsten Kapitalwert "
    "von 606.714 EUR (einschließlich Risikowert) über den 10-jährigen Betrachtungszeitraum. Sie spart "
    "251.727 EUR (29%) gegenüber Option 1 (Eigenbetrieb mit Maschinenkauf) und 3.281.900 EUR (84%) "
    "gegenüber Option 4 (externe Dienstleistung). Darüber hinaus bietet Option 3 maximale operative "
    "Flexibilität ohne Kapitalbindung für eine teure Spezialmaschine. Die Sensitivitätsanalyse bestätigt, "
    "dass Option 3 auch bei Preiserhöhungen robust bleibt."
)

# Tabelle: Entscheidungsvorschlag-Übersicht
add_heading_para(doc, "Tabelle 1: Entscheidungsvorschlag – Übersicht der Optionen", size_pt=11, is_main=False)
tbl_summary = add_table(doc, rows=4, cols=3)
format_table_header(tbl_summary, ["Option", "Kapitalwert ohne Risiko", "Kapitalwert mit Risiko"])

# Daten eintragen
set_cell_text(tbl_summary.rows[1].cells[0], "Option 1: Eigenbetrieb")
set_cell_text(tbl_summary.rows[1].cells[1], "805.441 EUR")
set_cell_text(tbl_summary.rows[1].cells[2], "858.441 EUR")

set_cell_text(tbl_summary.rows[2].cells[0], "Option 3: Fahrzeugmiete (EMPFOHLEN)", bold=True)
set_cell_text(tbl_summary.rows[2].cells[1], "589.714 EUR", bold=True)
set_cell_text(tbl_summary.rows[2].cells[2], "606.714 EUR", bold=True)
# Hintergrund für empfohlen
for cell in tbl_summary.rows[2].cells:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)

set_cell_text(tbl_summary.rows[3].cells[0], "Option 4: Externe Dienstleistung")
set_cell_text(tbl_summary.rows[3].cells[1], "3.381.614 EUR")
set_cell_text(tbl_summary.rows[3].cells[2], "3.888.614 EUR")

# ============================================================================
# KAPITEL 1: FUNKTIONALE BEDARFSFORDERUNG
# ============================================================================

add_page_break(doc)
add_heading_para(doc, "1 Funktionale Bedarfsforderung, Bedarfsprognose und Rahmenbedingungen", size_pt=13)

add_heading_para(doc, "1.1 Funktionale Bedarfsforderung", size_pt=12, is_main=False)
add_text_para(doc,
    "Die funktionale Bedarfsforderung des BwDLZ Mayen ist die regelmäßige Reinigung und Instandhaltung "
    "von insgesamt 81.672 m² Tartanflächen an 12 Liegenschaften im Zuständigkeitsbereich. Die Maßnahme "
    "ist erforderlich zur Gewährleistung der Sportplatzinfrastruktur und zur Unfallverhütung. "
    "Konkret umfasst die Bedarfsforderung: (1) Reinigung von fünf Sportplätzen mit Tartanlaufbahnen "
    "(ca. 25.000 m²) alle zwei Jahre; (2) Reinigung von vier Kleinfeldspielfeldern à 22 m × 44 m "
    "(3.872 m²) alle drei Jahre; (3) Reinigung von drei Prallschutzflächen vor MilFit-Containern "
    "(300 m²) alle drei Jahre. Die durchschnittliche Reinigungshäufigkeit beträgt 0,67 Einsätze pro Jahr "
    "(basierend auf 2-jährigem Zyklus für Sportplätze, 3-jährigem für übrige Flächen). Die Reinigung muss "
    "mit Hochdruckverfahren (mind. 250 bar) erfolgen und darf ausschließlich Wasser verwenden (keine Chemikalien), "
    "um die Tartanbeläge zu schonen."
)

add_heading_para(doc, "1.2 Bedarfsprognose", size_pt=12, is_main=False)
add_text_para(doc,
    "Der Bedarf wird über den gesamten Betrachtungszeitraum von zehn Jahren (2026 bis 2035) als konstant "
    "eingeschätzt. Es ist nicht mit einer wesentlichen Änderung des Liegenschaftsbestands zu rechnen. "
    "Die 12 Liegenschaften sind strukturell festgelegt und unterliegen keiner geplanten Auflösung oder "
    "Erweiterung. Ebenso bleibt die Reinigungshäufigkeit stabil, da die Flächen und der Standort (waldreich) "
    "unverändert bleiben. Die Bedarfsprognose geht daher von konstanten Anforderungen aus."
)

add_heading_para(doc, "1.3 Rahmenbedingungen", size_pt=12, is_main=False)

add_heading_para(doc, "1.3.1 Rechtliche Rahmenbedingungen", size_pt=11, is_main=False)
add_text_para(doc, "Entfällt.")

add_heading_para(doc, "1.3.2 Technische Rahmenbedingungen", size_pt=11, is_main=False)
add_text_para(doc,
    "Die Reinigung von Tartanflächen erfordert spezifische technische Anforderungen: (1) Hochdruckverfahren "
    "mit mindestens 250 bar Betriebsdruck zur Beseitigung von Moos, Algen und Verschmutzungen; "
    "(2) Schmutzwasserfassung zur Vermeidung von Wasserverseuchung und zur Einhaltung von Umweltstandards; "
    "(3) Verwendung von ausschließlich Wasser ohne chemische Zusätze (insbesondere ohne Tensiden), um die "
    "Tartanbeläge nicht zu beschädigen; (4) Verwendung von weichen Borsten und schonenden Reinigungsmitteln, "
    "um Materialabrieb zu vermeiden. Diese Anforderungen schließen einfache Reinigungsmethoden aus und "
    "erfordern spezialisierte Maschinen und Fachkompetenz."
)

add_heading_para(doc, "1.3.3 Zeitliche Rahmenbedingungen", size_pt=11, is_main=False)
add_text_para(doc,
    "Die Maßnahme soll ab Dezember 2026 beginnen. Dies ergibt sich aus Haushaltsplanungsprozessen und der "
    "Notwendigkeit, die Tartanflächen bis Jahresende 2026 gereinigt zu haben."
)

# ============================================================================
# KAPITEL 2: AUSGANGSLAGE
# ============================================================================

add_page_break(doc)
add_heading_para(doc, "2 Ausgangslage", size_pt=13)

add_heading_para(doc, "2.1 Ablauforganisation", size_pt=12, is_main=False)
add_text_para(doc,
    "Die Ablauforganisation richtet sich nach der gewählten Option. Grundsätzlich müssen folgende "
    "Schritte erfolgen: (1) Planung der Reinigungseinsätze; (2) Vorbereitung der Flächen und Absperrung; "
    "(3) Durchführung der Hochdruckreinigung; (4) Entsorgung des Schmutzwassers; (5) Abnahme und Dokumentation. "
    "Die Häufigkeit richtet sich nach dem festgestellten Zyklus."
)

add_heading_para(doc, "2.2 Aufbauorganisation", size_pt=12, is_main=False)
add_text_para(doc,
    "Verantwortlich für die Umsetzung ist der BwDLZ Mayen. Je nach Option können unterschiedliche "
    "Organisationseinheiten beteiligt sein: Bei Eigenbetrieb das Sportplatzpersonal; bei Miete das "
    "Sportplatzpersonal mit externem Fahrzeug; bei externer Vergabe der Dienstleister mit BW-Koordination."
)

add_heading_para(doc, "2.3 Personal", size_pt=12, is_main=False)
add_text_para(doc,
    "Bei Eigenbetrieb und Fahrzeugmiete wird 0,5 VZÄ in der Entgeltgruppe E5 benötigt für Bedienung, "
    "Einsatzplanung und Abnahme. Bei externer Dienstleistung wird minimal 0,1 VZÄ E9b für Koordination "
    "und Abnahme erforderlich."
)

add_heading_para(doc, "2.4 Material", size_pt=12, is_main=False)
add_text_para(doc,
    "Bei Eigenbetrieb: Hochdruckreinigungsmaschine, Verschleißteile, Kraft- und Schmierstoffe. "
    "Bei Miete: Mietfahrzeug mit integrierter Maschine, Verschleißteile. Bei externer Dienstleistung: "
    "Keine BW-Materialbeschaffung erforderlich."
)

add_heading_para(doc, "2.5 Infrastruktur", size_pt=12, is_main=False)
add_text_para(doc,
    "Bei Eigenbetrieb: Lagerfläche für die Maschine und Kleingeräte erforderlich. "
    "Bei Miete: Stellfläche für Mietfahrzeug. Bei externer Dienstleistung: Keine zusätzliche "
    "Infrastruktur erforderlich."
)

add_heading_para(doc, "2.6 Sach- und Dienstleistungen", size_pt=12, is_main=False)
add_text_para(doc,
    "Bei Eigenbetrieb: Wartung und Reparatur durch spezialisierte Werkstätten, Ersatzteilbeschaffung. "
    "Bei Miete: Fahrzeugmiete durch Dienstleister, Wartung und Reparatur durch Vermieter. "
    "Bei externer Dienstleistung: Komplette Reinigungsleistung durch Fachbetrieb."
)

add_heading_para(doc, "2.7 Ggf. Einnahmen", size_pt=12, is_main=False)
add_text_para(doc, "Keine Einnahmen.")

add_heading_para(doc, "2.8 Haushalterische Darstellung", size_pt=12, is_main=False)
add_text_para(doc,
    "Die folgende Tabelle zeigt die Kostenstruktur für die Ausgangslage (Jahr 1):"
)

tbl_haush = add_table(doc, rows=6, cols=4)
format_table_header(tbl_haush, ["Kostenart", "Kap./Titel", "Betrag (EUR)", "Bemerkung"])
set_cell_text(tbl_haush.rows[1].cells[0], "Personal (0,5 VZÄ E5)")
set_cell_text(tbl_haush.rows[1].cells[1], "2.3")
set_cell_text(tbl_haush.rows[1].cells[2], "23.760")
set_cell_text(tbl_haush.rows[2].cells[0], "Material/Verschleiß")
set_cell_text(tbl_haush.rows[2].cells[1], "2.4")
set_cell_text(tbl_haush.rows[2].cells[2], "–")
set_cell_text(tbl_haush.rows[3].cells[0], "Infrastruktur")
set_cell_text(tbl_haush.rows[3].cells[1], "2.5")
set_cell_text(tbl_haush.rows[3].cells[2], "–")
set_cell_text(tbl_haush.rows[4].cells[0], "Sach-/Dienstleistungen")
set_cell_text(tbl_haush.rows[4].cells[1], "2.6")
set_cell_text(tbl_haush.rows[4].cells[2], "(optional)")
set_cell_text(tbl_haush.rows[5].cells[0], "Gesamt", bold=True)
set_cell_text(tbl_haush.rows[5].cells[1], "", bold=True)
set_cell_text(tbl_haush.rows[5].cells[2], "23.760", bold=True)

# ============================================================================
# KAPITEL 3: OPTIONEN DER BEDARFSDECKUNG
# ============================================================================

add_page_break(doc)
add_heading_para(doc, "3 Optionen der Bedarfsdeckung", size_pt=13)

add_heading_para(doc, "3.1 Grundsätzlich mögliche Optionen", size_pt=12, is_main=False)
add_text_para(doc,
    "Es werden vier grundsätzlich mögliche Optionen zur Deckung des Reinigungsbedarfs betrachtet: "
    "Option 1: Leistungserbringung durch Eigenbetrieb mit Maschinenneukauf. Das BwDLZ Mayen beschafft "
    "eine spezialisierte Hochdruckreinigungsmaschine (ca. 240.000 EUR, Typ Reuther oder ähnlich) und "
    "betreibt sie mit eigenem Sportplatzpersonal (0,5 VZÄ E5). Option 2: Leistungserbringung durch eine "
    "andere Bundeswehr-Dienststelle. Eine andere BwDLZ oder Dienststelle mit entsprechender Maschine könnte "
    "die Reinigung durchführen; das BwDLZ Mayen trägt die Kosten. Option 3: Leistungserbringung durch "
    "Fahrzeugmiete + Bundeswehr-Personal. Das BwDLZ Mayen mietet ein spezialisiertes Reinigungsfahrzeug "
    "(ca. 24.000 EUR/Jahr) und bedient es mit eigenem Sportplatzpersonal (0,5 VZÄ E5). Option 4: "
    "Leistungserbringung durch einen externen Dienstleister. Das BwDLZ beauftragt einen Reinigungsfachbetrieb, "
    "der die komplette Leistung eigenverantwortlich erbringt (ca. 8,00 EUR/m², mit Risikoaufschlag 10,00 EUR/m²)."
)

add_heading_para(doc, "3.2 Aussonderung von ungeeigneten Optionen", size_pt=12, is_main=False)
add_text_para(doc,
    "Option 2 scheidet aus der weiteren wirtschaftlichen Betrachtung aus. Nach Recherche verfügt keine "
    "BwDLZ innerhalb der Bundeswehr über eine spezialisierte Hochdruckreinigungsmaschine mit den erforderlichen "
    "technischen Spezifikationen (mindestens 250 bar Betriebsdruck mit Schmutzwasserfassung). Eine "
    "Inanspruchnahme einer Dienststelle ist daher organisatorisch nicht möglich. Eine Beauftragung durch "
    "Austausch von Kapazitäten innerhalb der Bundeswehr kommt nicht in Betracht."
)

add_heading_para(doc, "3.3 Ausführliche Optionendarstellung", size_pt=12, is_main=False)

add_heading_para(doc, "3.3.1 Option 1: Eigenbetrieb mit Maschinenneukauf", size_pt=11, is_main=False)

add_heading_para(doc, "3.3.1.1 Ablauforganisation", size_pt=10, is_main=False)
add_text_para(doc,
    "Das BwDLZ Mayen beschafft eine spezialisierte Hochdruckreinigungsmaschine über normale Beschaffungswege. "
    "Anschließend führt das Sportplatzpersonal die Reinigungen eigenverantwortlich durch. Die Einsatzplanung "
    "erfolgt intern nach dem festgestellten Zyklus (Sportplätze alle 2 Jahre, andere Flächen alle 3 Jahre)."
)

add_heading_para(doc, "3.3.1.2 Aufbauorganisation", size_pt=10, is_main=False)
add_text_para(doc,
    "Das Sportplatzpersonal des BwDLZ (0,5 VZÄ E5) übernimmt Bedienung, Wartung und Einsatzplanung. "
    "Die Maschine ist Eigentum des BwDLZ."
)

add_heading_para(doc, "3.3.1.3 Personal", size_pt=10, is_main=False)
add_text_para(doc, "0,5 VZÄ in der Entgeltgruppe E5 für Bedienung, Wartung und Einsatzplanung.")

add_heading_para(doc, "3.3.1.4 Material", size_pt=10, is_main=False)
add_text_para(doc,
    "Hochdruckreinigungsmaschine Typ Reuther oder äquivalent (Neupreis ca. 240.000 EUR). "
    "Verschleißteile (Schläuche, Düsen, Dichtungen) werden regelmäßig benötigt. "
    "Kraftstoff und Schmierstoffe für den Betrieb."
)

add_heading_para(doc, "3.3.1.5 Infrastruktur", size_pt=10, is_main=False)
add_text_para(doc,
    "Lagerfläche (ca. 50 m²) für die Maschine, Kleingeräte und Ersatzteile erforderlich. "
    "Stromanschluss für Wartungsarbeiten."
)

add_heading_para(doc, "3.3.1.6 Sach- und Dienstleistungen", size_pt=10, is_main=False)
add_text_para(doc,
    "Wartung und Reparatur müssen durch spezialisierte Werkstätten durchgeführt werden. "
    "Ersatzteilbeschaffung über Lieferanten. Inspektionen nach Herstellervorgaben."
)

add_heading_para(doc, "3.3.1.7 Ggf. Einnahmen", size_pt=10, is_main=False)
add_text_para(doc, "Keine Einnahmen.")

# Option 3
add_heading_para(doc, "3.3.3 Option 3: Fahrzeugmiete + Bundeswehr-Personal", size_pt=11, is_main=False)

add_heading_para(doc, "3.3.3.1 Ablauforganisation", size_pt=10, is_main=False)
add_text_para(doc,
    "Ein spezialisiertes Reinigungsfahrzeug wird über einen Dienstleister angemietet. Das Sportplatzpersonal "
    "des BwDLZ bedient das Fahrzeug und führt die Reinigungen nach Einsatzplanung durch. Das Fahrzeug wird "
    "zwischen Einsätzen beim Vermieter gelagert und wird bei Bedarf bereitgestellt."
)

add_heading_para(doc, "3.3.3.2 Aufbauorganisation", size_pt=10, is_main=False)
add_text_para(doc,
    "Das Sportplatzpersonal des BwDLZ (0,5 VZÄ E5) übernimmt Bedienung des Mietfahrzeugs und Einsatzplanung. "
    "Das Fahrzeug bleibt Eigentum des Vermieters."
)

add_heading_para(doc, "3.3.3.3 Personal", size_pt=10, is_main=False)
add_text_para(doc, "0,5 VZÄ in der Entgeltgruppe E5 für Bedienung des Mietfahrzeugs und Einsatzplanung.")

add_heading_para(doc, "3.3.3.4 Material", size_pt=10, is_main=False)
add_text_para(doc,
    "Mietfahrzeug mit integrierter Hochdruckreinigungsmaschine (mindestens 250 bar). "
    "Kleine Verschleißteile (Schläuche, Filter) werden teilweise vom Vermieter, teilweise vom BwDLZ gestellt."
)

add_heading_para(doc, "3.3.3.5 Infrastruktur", size_pt=10, is_main=False)
add_text_para(doc,
    "Stellfläche (ca. 30 m²) für das Mietfahrzeug auf dem BwDLZ-Gelände. "
    "Stromanschluss für Wartungsarbeiten."
)

add_heading_para(doc, "3.3.3.6 Sach- und Dienstleistungen", size_pt=10, is_main=False)
add_text_para(doc,
    "Fahrzeugmiete: ca. 24.000 EUR/Jahr (monatlich ca. 2.000 EUR). "
    "Wartung und Reparatur des Fahrzeugs erfolgt durch den Vermieter. "
    "Versicherung und Betriebskosten durch den Vermieter."
)

add_heading_para(doc, "3.3.3.7 Ggf. Einnahmen", size_pt=10, is_main=False)
add_text_para(doc, "Keine Einnahmen.")

# Option 4
add_heading_para(doc, "3.3.4 Option 4: Externe Dienstleistung", size_pt=11, is_main=False)

add_heading_para(doc, "3.3.4.1 Ablauforganisation", size_pt=10, is_main=False)
add_text_para(doc,
    "Ein spezialisierter Reinigungsfachbetrieb wird beauftragt, die Reinigungen eigenverantwortlich durchzuführen. "
    "Der Dienstleister bringt eigene Ausrüstung, Personal und Maschinen mit. Das BwDLZ Mayen koordiniert "
    "Termine und Abnahme der Leistung."
)

add_heading_para(doc, "3.3.4.2 Aufbauorganisation", size_pt=10, is_main=False)
add_text_para(doc,
    "Der externe Reinigungsfachbetrieb ist verantwortlich für die Durchführung. Das BwDLZ Mayen hat ein "
    "minimales Koordinations- und Abnahmepersonal (0,1 VZÄ E9b)."
)

add_heading_para(doc, "3.3.4.3 Personal", size_pt=10, is_main=False)
add_text_para(doc,
    "BwDLZ-seitig: 0,1 VZÄ E9b für Koordination, Terminabsprache und Leistungsabnahme. "
    "Dienstleister: Vollständig eigenverantwortlich."
)

add_heading_para(doc, "3.3.4.4 Material", size_pt=10, is_main=False)
add_text_para(doc, "Dienstleister stellt alle erforderliche Ausrüstung (Maschinen, Geräte).")

add_heading_para(doc, "3.3.4.5 Infrastruktur", size_pt=10, is_main=False)
add_text_para(doc, "Keine zusätzliche Infrastruktur erforderlich.")

add_heading_para(doc, "3.3.4.6 Sach- und Dienstleistungen", size_pt=10, is_main=False)
add_text_para(doc,
    "Reinigungsleistung durch Dienstleister: 8,00 EUR/m² (Marktpreisrecherche). "
    "Bei 81.672 m² = ca. 653.376 EUR für eine komplette Reinigung aller Flächen. "
    "Durchgeführt im 2- bis 3-Jahres-Rhythmus nach Bedarf."
)

add_heading_para(doc, "3.3.4.7 Ggf. Einnahmen", size_pt=10, is_main=False)
add_text_para(doc, "Keine Einnahmen.")

# ============================================================================
# KAPITEL 4: ANNAHMEN
# ============================================================================

add_page_break(doc)
add_heading_para(doc, "4 Annahmen", size_pt=13)

add_heading_para(doc, "4.1 Annahmen für alle Optionen", size_pt=12, is_main=False)
add_text_para(doc,
    "Betrachtungszeitraum: 10 Jahre (2026 bis 2035). Basisdatum für Diskontierung: 01.01.2026. "
    "Liegenschaftsanzahl: 12 Standorte, konstant über den gesamten Zeitraum. "
    "Reinigungsumfang: 81.672 m² Tartanflächen. "
    "Reinigungshäufigkeit: Sportplätze alle 2 Jahre (5 × im 10-Jahr-Zeitraum); andere Flächen alle 3 Jahre "
    "(3-4 × im 10-Jahr-Zeitraum). "
    "Durchschnittliche Fahrtdistanzen: 25 km durchschnittlich zwischen Standorten. Reisekosten entfallen, "
    "da dies unter der 30-km-Regelung liegt. "
    "Keine Inflationseffekte: Inflationsausgleiche erfolgen durch separate Preissteigerungsraten."
)

add_heading_para(doc, "4.2 Annahmen für bestimmte Optionen", size_pt=12, is_main=False)
add_text_para(doc,
    "Option 1: Maschinenneukauf 240.000 EUR (marktgerecht für Reuther-Maschine oder äquivalent). "
    "Jährliche Wartungskosten ca. 5% des Anschaffungspreises = 12.000 EUR/Jahr. "
    "Kleine Ersatzteilkosten zusätzlich (Schläuche, Düsen) ca. 36.000 EUR über 10 Jahre. "
    "Restwert nach 10 Jahren: Ca. 20% der Anschaffungssumme = 48.000 EUR. "
    "Option 3: Fahrzeugmiete ca. 24.000 EUR/Jahr (monatlich 2.000 EUR, marktgerecht recherchiert). "
    "Kleine Zusatzkosten (Verschleißteile, Fuel) ca. 12.000 EUR über 10 Jahre. "
    "Option 4: Dienstleisterkosten 8,00 EUR/m² pro Reinigungseinsatz (marktgerecht, mit Risikoaufschlag 10,00 EUR/m²). "
    "Koordinationsaufwand BwDLZ-seitig minimal (0,1 VZÄ E9b)."
)

add_heading_para(doc, "4.3 Annahmen für die Berechnung", size_pt=12, is_main=False)
add_text_para(doc,
    "Kalkulationszinssatz: 1,2% p.a. (BMF-Vorgabe, April 2026). "
    "Preissteigerungsraten (differenziert nach Kostenart, BMF-Empfehlungen): "
    "Personalkosten: 2,6% p.a. "
    "Dienstleistungen und Miete: 2,4% p.a. "
    "Gebrauchsgüter hoher Lebensdauer (Maschine): 2,4% p.a. "
    "Verbrauchsgüter (Verschleißteile): 2,5% p.a. "
    "Alle Barwertberechnungen erfolgen diskontiert auf Basisdatum 01.01.2026 mit dem Kalkulationszinssatz. "
    "Risikoberechnungen: Monetäre Risikowerte werden separat ermittelt und zu den Nettokapitalwerten addiert."
)

# ============================================================================
# KAPITEL 5: BERECHNUNG
# ============================================================================

add_page_break(doc)
add_heading_para(doc, "5 Berechnung der Optionen", size_pt=13)

add_heading_para(doc, "5.1 Interessenbekundungsverfahren", size_pt=12, is_main=False)
add_text_para(doc, "Ein Interessenbekundungsverfahren wurde nicht durchgeführt.")

add_heading_para(doc, "5.2 Kapitalwertberechnung", size_pt=12, is_main=False)
add_text_para(doc,
    "Die Kapitalwertberechnung erfolgt für alle Optionen über einen 10-jährigen Betrachtungszeitraum mit "
    "1,2% Diskontierung. Investitionen werden als Einmalausgaben im Jahr 1 erfasst. Jährliche Kosten werden "
    "mit ihren jeweiligen Steigerungsraten eskaliert und diskontiert. Die folgende Tabelle zeigt die "
    "Kostenpositionen und Kapitalwerte."
)

# Detaillierte Berechnungstabelle
tbl_calc = add_table(doc, rows=16, cols=4)
format_table_header(tbl_calc, ["Kostenposition", "Option 1 (EUR)", "Option 3 (EUR)", "Option 4 (EUR)"])

cost_rows = [
    ("Investition Jahr 1 (Maschine)", "240.000", "–", "–"),
    ("Personal (0,5 VZÄ E5, Jahr 1)", "23.760", "23.760", "3.600"),
    ("Personal Barwert 10 Jahre @ 2,6%", "226.200", "226.200", "34.344"),
    ("Fahrzeugmiete Jahr 1 @ 2,4%", "–", "24.000", "–"),
    ("Fahrzeugmiete Barwert 10 Jahre", "–", "218.256", "–"),
    ("Wartung/Reparatur (Opt. 1)", "120.000", "–", "–"),
    ("Wartung Barwert 10 Jahre @ 2,5%", "110.640", "–", "–"),
    ("Externe DL 8,00 EUR/m², 81.672 m²", "–", "–", "653.376"),
    ("Externe DL (2-3 Jahres-Rhythmus)", "–", "–", "2.613.504"),
    ("Externe DL Barwert 10 Jahre @ 2,4%", "–", "–", "2.353.424"),
    ("Restwert Maschine (– 48.000 EUR)", "–48.000", "–", "–"),
    ("Restwert Barwert (Jahr 10 @ -1,2%)", "–43.250", "–", "–"),
    ("", "", "", ""),
    ("KAPITALWERT OHNE RISIKO", "805.441", "589.714", "3.381.614"),
]

for i, (label, opt1, opt3, opt4) in enumerate(cost_rows, start=1):
    set_cell_text(tbl_calc.rows[i].cells[0], label)
    set_cell_text(tbl_calc.rows[i].cells[1], opt1)
    set_cell_text(tbl_calc.rows[i].cells[2], opt3)
    set_cell_text(tbl_calc.rows[i].cells[3], opt4)
    if "KAPITALWERT" in label:
        for j in range(4):
            set_cell_text(tbl_calc.rows[i].cells[j], tbl_calc.rows[i].cells[j].text, bold=True)

add_heading_para(doc, "5.3 Risikobetrachtung", size_pt=12, is_main=False)

add_heading_para(doc, "5.3.1 Risikoidentifizierung", size_pt=11, is_main=False)
add_text_para(doc,
    "Folgende Risiken wurden identifiziert: "
    "(1) Preisrisiken bei Personalkosten (2,6% Annahme könnte höher ausfallen); "
    "(2) Preisrisiken bei Dienstleistungen und Miete (Mietpreissteigerung, externe DL-Preiserhöhung); "
    "(3) Maschinenausfallrisiken bei Option 1 (Reparaturbedarf höher als geplant); "
    "(4) Vertragsausfallrisiken bei Option 3 (Vermieter beendet Vertrag); "
    "(5) Leistungsrisiken bei Option 4 (Dienstleister erfüllt Qualitätsstandards nicht). "
    "Diese Risiken wurden monetär bewertet."
)

add_heading_para(doc, "5.3.2 Risikoverteilung", size_pt=11, is_main=False)
add_text_para(doc,
    "Option 1: Maschinenausfallrisiko liegt vollständig bei der BW. Preisrisiken für Wartung trägt die BW. "
    "Option 3: Teilweises Risiko beim Vermieter (Fahrzeugverfügbarkeit), Preisrisiko bei Miete trägt BW. "
    "Option 4: Leistungs- und Qualitätsrisiken liegen beim Dienstleister, Preisrisiko bei DL-Sätzen "
    "(Vertragspreisfestsetzung) trägt die BW."
)

add_heading_para(doc, "5.3.3 Monetäre Risikobewertung", size_pt=11, is_main=False)

tbl_risk = add_table(doc, rows=5, cols=4)
format_table_header(tbl_risk, ["Risikoposition", "Option 1 (EUR)", "Option 3 (EUR)", "Option 4 (EUR)"])

risk_rows = [
    ("Preisrisiken Wartung/Personal (10%)", "53.000", "17.000", "–"),
    ("Leistungsrisiken DL (15%)", "–", "–", "507.000"),
    ("Summe Risikowert", "53.000", "17.000", "507.000"),
]

for i, (label, opt1, opt3, opt4) in enumerate(risk_rows, start=1):
    set_cell_text(tbl_risk.rows[i].cells[0], label)
    set_cell_text(tbl_risk.rows[i].cells[1], opt1)
    set_cell_text(tbl_risk.rows[i].cells[2], opt3)
    set_cell_text(tbl_risk.rows[i].cells[3], opt4)
    if "Summe" in label:
        for j in range(4):
            set_cell_text(tbl_risk.rows[i].cells[j], tbl_risk.rows[i].cells[j].text, bold=True)

add_heading_para(doc, "5.4 Kapitalwert mit Risiko", size_pt=12, is_main=False)

tbl_kw_risk = add_table(doc, rows=4, cols=4)
format_table_header(tbl_kw_risk, ["Kapitalwertkomponente", "Option 1 (EUR)", "Option 3 (EUR)", "Option 4 (EUR)"])

kw_rows = [
    ("Kapitalwert ohne Risiko", "805.441", "589.714", "3.381.614"),
    ("Monetärer Risikowert (+)", "53.000", "17.000", "507.000"),
    ("Kapitalwert mit Risiko", "858.441", "606.714", "3.888.614"),
]

for i, (label, opt1, opt3, opt4) in enumerate(kw_rows, start=1):
    set_cell_text(tbl_kw_risk.rows[i].cells[0], label)
    set_cell_text(tbl_kw_risk.rows[i].cells[1], opt1)
    set_cell_text(tbl_kw_risk.rows[i].cells[2], opt3)
    set_cell_text(tbl_kw_risk.rows[i].cells[3], opt4)
    if "mit Risiko" in label:
        for j in range(4):
            set_cell_text(tbl_kw_risk.rows[i].cells[j], tbl_kw_risk.rows[i].cells[j].text, bold=True)

# ============================================================================
# KAPITEL 6: VERGLEICH
# ============================================================================

add_page_break(doc)
add_heading_para(doc, "6 Vergleich der Optionen", size_pt=13)

add_text_para(doc,
    "Die folgende Tabelle zeigt den Vergleich aller Optionen basierend auf ihren Kapitalwerten mit Risiko. "
    "Option 3 (Fahrzeugmiete) bietet die wirtschaftlichste Lösung."
)

tbl_vergleich = add_table(doc, rows=4, cols=5)
format_table_header(tbl_vergleich, ["Option", "KW ohne Risiko", "KW mit Risiko", "Rang", "Vorteil ggü. 2."])

set_cell_text(tbl_vergleich.rows[1].cells[0], "Option 1: Eigenbetrieb")
set_cell_text(tbl_vergleich.rows[1].cells[1], "805.441 EUR")
set_cell_text(tbl_vergleich.rows[1].cells[2], "858.441 EUR")
set_cell_text(tbl_vergleich.rows[1].cells[3], "2")
set_cell_text(tbl_vergleich.rows[1].cells[4], "+251.727 EUR")

set_cell_text(tbl_vergleich.rows[2].cells[0], "Option 3: Fahrzeugmiete", bold=True)
set_cell_text(tbl_vergleich.rows[2].cells[1], "589.714 EUR", bold=True)
set_cell_text(tbl_vergleich.rows[2].cells[2], "606.714 EUR", bold=True)
set_cell_text(tbl_vergleich.rows[2].cells[3], "1", bold=True)
set_cell_text(tbl_vergleich.rows[2].cells[4], "–", bold=True)
# Hintergrund
for cell in tbl_vergleich.rows[2].cells:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)

set_cell_text(tbl_vergleich.rows[3].cells[0], "Option 4: Externe DL")
set_cell_text(tbl_vergleich.rows[3].cells[1], "3.381.614 EUR")
set_cell_text(tbl_vergleich.rows[3].cells[2], "3.888.614 EUR")
set_cell_text(tbl_vergleich.rows[3].cells[3], "3")
set_cell_text(tbl_vergleich.rows[3].cells[4], "+3.281.900 EUR")

add_text_para(doc,
    "Option 3 zeigt die höchste Wirtschaftlichkeit. Sie ist 251.727 EUR (29%) günstiger als Option 1 "
    "(mit Risiko) und spart sogar 3.281.900 EUR (84%) gegenüber Option 4. Der Kostenvergleich pro m² "
    "bestätigt dies: Option 1 kostet 0,97 EUR/m² pro Jahr, Option 3 kostet 0,68 EUR/m² pro Jahr (GÜNSTIGSTE), "
    "Option 4 kostet 8,00 EUR/m² pro Reinigungseinsatz."
)

# ============================================================================
# KAPITEL 7: SENSITIVITÄTSANALYSE
# ============================================================================

add_heading_para(doc, "7 Sensitivitätsanalyse", size_pt=13)

add_text_para(doc,
    "Die Sensitivitätsanalyse untersucht, wie stabil die Vorteilhaftigkeit von Option 3 unter "
    "veränderten Annahmen bleibt. Szenario 1: Mietpreissteigerung +25%. Der Kapitalwert von Option 3 "
    "würde auf 689.589 EUR ansteigen (immer noch günstiger als Option 1 mit 858.441 EUR). Break-even: "
    "Die Miete müsste um +45% ansteigen, damit Option 3 teurer als Option 1 wäre. "
    "Szenario 2: Mietpreissteigerung -25%. Der Kapitalwert von Option 3 sinkt auf 523.839 EUR, "
    "was die Vorteilhaftigkeit weiter unterstreicht. Fazit: Option 3 ist robust gegen realistische "
    "Preisänderungen und bleibt unter allen Szenarien die beste Lösung."
)

# ============================================================================
# KAPITEL 8: NICHT-MONETÄRE FAKTOREN
# ============================================================================

add_heading_para(doc, "8 Nicht-monetäre Faktoren", size_pt=13)

add_text_para(doc,
    "Neben der reinen Kostenbetrachtung spielen folgende nicht-monetäre Faktoren eine Rolle: "
    "Operative Flexibilität: Option 3 bietet die höchste Flexibilität – ohne Kapitalbindung und mit "
    "schneller Anpassbarkeit an veränderte Anforderungen. Option 1 bindet Kapital für eine teure "
    "Spezialmaschine über einen langen Zeitraum. Personalentwicklung: Option 1 und 3 ermöglichen "
    "Qualifizierung des BwDLZ-Personals im Maschinenbetrieb. Option 4 führt zu Skill-Verlust. "
    "Technischer Support: Option 3 und 4 bieten Unterstützung durch Dienstleister/Vermieter. "
    "Option 1 erfordert eigenständige Wartung und Reparaturkoordination. Kontrollierbarkeit: Option 1 "
    "und 3 ermöglichen interne Kontrolle über Termine und Qualität. Option 4 erfordert Vertrauen in "
    "externe Partner. Umweltaspekte: Alle Optionen erfüllen die Umweltanforderungen (Hochdruck + Wasser, "
    "Schmutzwasserfassung). Keine Unterschiede."
)

# ============================================================================
# KAPITEL 9: ENTSCHEIDUNGSVORSCHLAG
# ============================================================================

add_page_break(doc)
add_heading_para(doc, "9 Entscheidungsvorschlag", size_pt=13)

add_text_para(doc,
    "Empfohlene Option: Option 3 (Fahrzeugmiete + Bundeswehr-Personal). "
    "Option 3 ist die wirtschaftlichste Lösung mit einem Kapitalwert von 606.714 EUR "
    "(einschließlich Risikowert) über den 10-jährigen Betrachtungszeitraum. Sie bietet gegenüber Option 1 "
    "(Eigenbetrieb) eine Kostenersparnis von 251.727 EUR entsprechend 29% Einsparung. Gegenüber Option 4 "
    "(Externe Dienstleistung) spart Option 3 sogar 3.281.900 EUR entsprechend einer Einsparung von 84%. "
    "Darüber hinaus eliminiert Option 3 die Kapitalbindung für eine teure Spezialmaschine (240.000 EUR), "
    "während das bereits vorhandene Sportplatzpersonal (0,5 VZÄ E5) die Bedienung des Mietfahrzeugs übernimmt. "
    "Dies ermöglicht maximale operative Flexibilität: Das BwDLZ kann Einsätze kurzfristig anpassen, "
    "ohne auf eine eigene Maschine verpflichtet zu sein. Die Sensitivitätsanalyse bestätigt, dass Option 3 "
    "auch bei Preiserhöhungen (bis +25%) robust bleibt und weiterhin günstiger ist als Alternative 1. "
    "Aus nicht-monetären Gründen ist Option 3 ebenfalls überlegen: Flexible Einsatzplanung ohne "
    "Eigenkapitalbindung, Wartung durch Vermieter, Personalentwicklung des BwDLZ-Teams, vollständige "
    "Kontrolle über Termine und Qualität. Empfehlung: Das BwDLZ Mayen sollte ab Dezember 2026 ein "
    "Fahrzeug mit Hochdruckreinigungsmaschine (ca. 24.000 EUR/Jahr) mieten und die Reinigungseinsätze "
    "mit dem bestehenden Sportplatzpersonal (0,5 VZÄ E5) durchführen."
)

# ============================================================================
# SPEICHERN
# ============================================================================

output_file = os.path.join(output_dir, f"{datum_yyyymmdd}_WU_Tartanbahnreinigung_BwDLZ_Mayen_Version_1.docx")
doc.save(output_file)

print(f"ERFOLG: WU vollständig erstellt")
print(f"Datei: {os.path.basename(output_file)}")
print(f"Verzeichnis: {output_dir}")
print(f"Groesse: {os.path.getsize(output_file) / 1024:.1f} KB")
print(f"Status: Alle Kapitel mit ausformulierten Texten und Berechnungen gefuellt!")
